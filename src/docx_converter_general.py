# # # # # from docx import Document
# # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # from docx.oxml.ns import qn
# # # # # from docx.oxml import OxmlElement
# # # # # from bs4 import BeautifulSoup, Tag, NavigableString
# # # # # import os
# # # # # import re

# # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # #         html_content = file.read()
    
# # # # #     soup = BeautifulSoup(html_content, 'html.parser')
    
# # # # #     doc = Document()
    
# # # # #     sections = doc.sections
# # # # #     for section in sections:
# # # # #         section.top_margin = Inches(0.5)
# # # # #         section.bottom_margin = Inches(0.5)
# # # # #         section.left_margin = Inches(0.5)
# # # # #         section.right_margin = Inches(0.5)
    
# # # # #     try:
# # # # #         logo_path = "download.png" 
# # # # #         if os.path.exists(logo_path):
# # # # #             logo_para = doc.add_paragraph()
# # # # #             logo_para.paragraph_format.space_before = Pt(12)
# # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # #             run = logo_para.add_run()
# # # # #             run.add_picture(logo_path, width=Inches(3.0))
# # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # #     except Exception as e:
# # # # #         print(f"Could not add logo: {e}")
    
# # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # #     if main_table:
# # # # #         doc_table = doc.add_table(rows=0, cols=2)
# # # # #         doc_table.style = 'Table Grid'
# # # # #         doc_table.autofit = False
# # # # #         doc_table.allow_autofit = False
# # # # #         doc_table.columns[0].width = Inches(1.875)
# # # # #         doc_table.columns[1].width = Inches(5.625)
# # # # #         set_table_borders(doc_table)
        
# # # # #         rows = main_table.find_all('tr')
# # # # #         for row in rows:
# # # # #             cells = row.find_all('td')
# # # # #             if len(cells) == 2:
# # # # #                 doc_row = doc_table.add_row().cells
# # # # #                 left_cell = doc_row[0]
# # # # #                 left_text = cells[0].get_text(strip=True)
# # # # #                 set_cell_content(left_cell, left_text, is_header=True, center_align=True)
                
# # # # #                 right_cell = doc_row[1]
# # # # #                 set_cell_content(right_cell, "", is_header=False, center_align=False) 
# # # # #                 process_html_content(right_cell, cells[1])
        
# # # # #     project_title_div = soup.find('div', {'class': 'project-title-bar'})
# # # # #     projects_table_html = soup.find('table', {'class': 'projects-table'})
    
# # # # #     if projects_table_html and project_title_div:
# # # # #         proj_table = doc.add_table(rows=0, cols=2)
# # # # #         proj_table.style = 'Table Grid'
# # # # #         proj_table.autofit = False
# # # # #         proj_table.allow_autofit = False
# # # # #         proj_table.columns[0].width = Inches(1.875)
# # # # #         proj_table.columns[1].width = Inches(5.625)
# # # # #         set_table_borders(proj_table)

# # # # #         title_row_cells = proj_table.add_row().cells
# # # # #         merged_cell = title_row_cells[0].merge(title_row_cells[1])
# # # # #         title_text = project_title_div.get_text(strip=True)
# # # # #         set_cell_content(merged_cell, title_text, is_header=True, center_align=False, font_size=14)
        
# # # # #         rows = projects_table_html.find_all('tr')
# # # # #         for row in rows:
# # # # #             cells = row.find_all('td')
# # # # #             if len(cells) == 2:
# # # # #                 doc_row = proj_table.add_row().cells
# # # # #                 left_cell = doc_row[0]
# # # # #                 left_text = cells[0].get_text(strip=True)
# # # # #                 set_cell_content(left_cell, left_text, is_header=True, center_align=True)
                
# # # # #                 right_cell = doc_row[1]
# # # # #                 set_cell_content(right_cell, "", is_header=False, center_align=False)
# # # # #                 process_project_content(right_cell, cells[1])
    
# # # # #     doc.save(docx_file_path)

# # # # # def set_table_borders(table):
# # # # #     tbl = table._tbl
# # # # #     tblPr = tbl.tblPr
# # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # #         border = OxmlElement(f'w:{border_name}')
# # # # #         border.set(qn('w:val'), 'single')
# # # # #         border.set(qn('w:sz'), '4')
# # # # #         border.set(qn('w:space'), '0')
# # # # #         border.set(qn('w:color'), '000000')
# # # # #         tblBorders.append(border)
# # # # #     tblPr.append(tblBorders)

# # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12):
# # # # #     cell.text = ""
# # # # #     paragraph = cell.paragraphs[0]
# # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # #     paragraph.paragraph_format.space_after = Pt(0)
    
# # # # #     if center_align: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# # # # #     else: paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
# # # # #     if is_header:
# # # # #         set_cell_background(cell, "0d3d62")
# # # # #         run = paragraph.add_run(text)
# # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # #         run.font.bold = True
# # # # #         run.font.size = Pt(font_size)
# # # # #         run.font.name = 'Arial'
# # # # #         cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# # # # #     else:
# # # # #         if text:
# # # # #             run = paragraph.add_run(text)
# # # # #             run.font.size = Pt(font_size)
# # # # #             run.font.name = 'Arial'
# # # # #             run.font.bold = False

# # # # # def set_cell_background(cell, color_hex):
# # # # #     shading_elm = OxmlElement('w:shd')
# # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # def process_html_content(cell, html_element):
# # # # #     for element in html_element.contents:
# # # # #         if isinstance(element, Tag):
# # # # #             if element.name == 'div' and 'section-title' in element.get('class', []):
# # # # #                 para = cell.add_paragraph()
# # # # #                 para.paragraph_format.space_before = Pt(8)
# # # # #                 para.paragraph_format.space_after = Pt(4)
# # # # #                 run = para.add_run(element.get_text(strip=True))
# # # # #                 run.bold = True
# # # # #                 run.font.size = Pt(14)
# # # # #                 run.font.name = 'Arial'
# # # # #                 run.font.color.rgb = RGBColor(0, 0, 0)
# # # # #             elif element.name == 'ul':
# # # # #                 for li in element.find_all('li'):
# # # # #                     para = cell.add_paragraph()
# # # # #                     para.style = 'List Bullet'
# # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # #                     para.paragraph_format.space_before = Pt(0)
# # # # #                     para.paragraph_format.space_after = Pt(6)
                    
# # # # #                     strong = li.find('strong')
# # # # #                     if strong:
# # # # #                         strong_run = para.add_run(strong.get_text(strip=True))
# # # # #                         strong_run.bold = True
# # # # #                         strong_run.font.size = Pt(12)
# # # # #                         strong_run.font.name = 'Arial'
# # # # #                         remaining_text = li.get_text().replace(strong.get_text(), '', 1).strip()
# # # # #                         if remaining_text:
# # # # #                             para.add_run(" ").font.size = Pt(12)
# # # # #                             regular_run = para.add_run(remaining_text)
# # # # #                             regular_run.font.size = Pt(12)
# # # # #                             regular_run.font.name = 'Arial'
# # # # #                     else:
# # # # #                         run = para.add_run(li.get_text(strip=True))
# # # # #                         run.font.size = Pt(12)
# # # # #                         run.font.name = 'Arial'
# # # # #             else:
# # # # #                 process_html_content(cell, element)
# # # # #         elif isinstance(element, NavigableString):
# # # # #             text = str(element).strip()
# # # # #             if text:
# # # # #                 para = cell.add_paragraph()
# # # # #                 para.paragraph_format.space_after = Pt(4)
# # # # #                 run = para.add_run(text)
# # # # #                 run.font.size = Pt(12)
# # # # #                 run.font.name = 'Arial'
# # # # #                 run.font.bold = 'strong' in str(html_element)
# # # # #                 run.font.color.rgb = RGBColor(0, 0, 0)

# # # # # def process_project_content(cell, html_element):
# # # # #     cell.text = ""
# # # # #     project_heading = html_element.find('div', {'class': 'project-heading'})
# # # # #     if project_heading:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_after = Pt(3)
# # # # #         run = para.add_run(project_heading.get_text(strip=True))
# # # # #         run.bold = True
# # # # #         run.font.size = Pt(13)
# # # # #         run.font.name = 'Arial'
    
# # # # #     role_heading = html_element.find('div', {'class': 'role-heading'})
# # # # #     if role_heading:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_after = Pt(2)
# # # # #         role_label = para.add_run("Role: ")
# # # # #         role_label.bold = True
# # # # #         role_label.font.size = Pt(12)
# # # # #         role_label.font.name = 'Arial'
# # # # #         text = role_heading.get_text(strip=True)
# # # # #         val = text.replace("Role:", "", 1).strip()
# # # # #         role_val = para.add_run(val)
# # # # #         role_val.font.size = Pt(12)
# # # # #         role_val.font.name = 'Arial'

# # # # #     duration_heading = html_element.find('div', {'class': 'duration-heading'})
# # # # #     if duration_heading:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_after = Pt(2)
# # # # #         dur_label = para.add_run("Duration: ")
# # # # #         dur_label.bold = True
# # # # #         dur_label.font.size = Pt(12)
# # # # #         dur_label.font.name = 'Arial'
# # # # #         text = duration_heading.get_text(strip=True)
# # # # #         val = text.replace("Duration:", "", 1).strip()
# # # # #         dur_val = para.add_run(val)
# # # # #         dur_val.font.size = Pt(12)
# # # # #         dur_val.font.name = 'Arial'

# # # # #     link = html_element.find('a')
# # # # #     if link:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_after = Pt(3)
# # # # #         para.add_run("Link- ").font.size = Pt(12)
# # # # #         run = para.add_run(link.get('href', ''))
# # # # #         run.font.size = Pt(12)
# # # # #         run.font.name = 'Arial'
    
# # # # #     description = html_element.find('p')
# # # # #     if description:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_before = Pt(5)
# # # # #         para.paragraph_format.space_after = Pt(3)
# # # # #         run = para.add_run(description.get_text(strip=True))
# # # # #         run.font.size = Pt(12)
# # # # #         run.font.name = 'Arial'
    
# # # # #     techstack = html_element.find('div', {'class': 'techstack'})
# # # # #     if techstack:
# # # # #         para = cell.add_paragraph()
# # # # #         para.paragraph_format.space_before = Pt(8)
# # # # #         strong = techstack.find('strong')
# # # # #         if strong:
# # # # #             strong_run = para.add_run(strong.get_text(strip=True))
# # # # #             strong_run.bold = True
# # # # #             strong_run.font.size = Pt(12)
# # # # #             strong_run.font.name = 'Arial'
# # # # #             remaining_text = techstack.get_text().replace(strong.get_text(), '', 1).strip()
# # # # #             if remaining_text:
# # # # #                 para.add_run("\n")
# # # # #                 regular_run = para.add_run(remaining_text)
# # # # #                 regular_run.font.size = Pt(12)
# # # # #                 regular_run.font.name = 'Arial'

# # # # # def convert_resume_to_docx(html_file_path, docx_file_path=None):
# # # # #     if docx_file_path is None:
# # # # #         docx_file_path = os.path.splitext(html_file_path)[0] + '.docx'
# # # # #     try:
# # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # #     except Exception as e:
# # # # #         print(f"Conversion failed: {e}")
# # # # #         raise



# # # # from docx import Document
# # # # from docx.shared import Inches, Pt, RGBColor
# # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # from docx.oxml.ns import qn
# # # # from docx.oxml import OxmlElement
# # # # from bs4 import BeautifulSoup, Tag, NavigableString
# # # # import os
# # # # import re

# # # # def html_to_docx(html_file_path, docx_file_path):
# # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # #         html_content = file.read()
    
# # # #     soup = BeautifulSoup(html_content, 'html.parser')
    
# # # #     doc = Document()
    
# # # #     sections = doc.sections
# # # #     for section in sections:
# # # #         section.top_margin = Inches(0.5)
# # # #         section.bottom_margin = Inches(0.5)
# # # #         section.left_margin = Inches(0.5)
# # # #         section.right_margin = Inches(0.5)
    
# # # #     try:
# # # #         logo_path = "download.png" # Updated path to look in root
# # # #         if os.path.exists(logo_path):
# # # #             logo_para = doc.add_paragraph()
# # # #             logo_para.paragraph_format.space_before = Pt(12)
# # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # #             run = logo_para.add_run()
# # # #             run.add_picture(logo_path, width=Inches(3.0))
# # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # #     except Exception as e:
# # # #         print(f"Could not add logo: {e}")
    
# # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # #     if main_table:
# # # #         doc_table = doc.add_table(rows=0, cols=2)
# # # #         doc_table.style = 'Table Grid'
# # # #         doc_table.autofit = False
# # # #         doc_table.allow_autofit = False
        
# # # #         doc_table.columns[0].width = Inches(1.875)
# # # #         doc_table.columns[1].width = Inches(5.625)
        
# # # #         set_table_borders(doc_table)
        
# # # #         rows = main_table.find_all('tr')
# # # #         for row in rows:
# # # #             cells = row.find_all('td')
# # # #             if len(cells) == 2:
# # # #                 doc_row = doc_table.add_row().cells
            
# # # #                 left_cell = doc_row[0]
# # # #                 left_text = cells[0].get_text(strip=True)
# # # #                 set_cell_content(left_cell, left_text, is_header=True, center_align=True)
                
# # # #                 right_cell = doc_row[1]
# # # #                 set_cell_content(right_cell, "", is_header=False, center_align=False) 
# # # #                 process_html_content(right_cell, cells[1])
        
# # # #     project_title_div = soup.find('div', {'class': 'project-title-bar'})
# # # #     projects_table_html = soup.find('table', {'class': 'projects-table'})
    
# # # #     if projects_table_html and project_title_div:
# # # #         proj_table = doc.add_table(rows=0, cols=2)
# # # #         proj_table.style = 'Table Grid'
# # # #         proj_table.autofit = False
# # # #         proj_table.allow_autofit = False
        
# # # #         proj_table.columns[0].width = Inches(1.875)
# # # #         proj_table.columns[1].width = Inches(5.625)
        
# # # #         set_table_borders(proj_table)

# # # #         title_row_cells = proj_table.add_row().cells
# # # #         merged_cell = title_row_cells[0].merge(title_row_cells[1])
# # # #         title_text = project_title_div.get_text(strip=True)
# # # #         set_cell_content(merged_cell, title_text, is_header=True, center_align=False, font_size=14)
        
# # # #         rows = projects_table_html.find_all('tr')
# # # #         for row in rows:
# # # #             cells = row.find_all('td')
# # # #             if len(cells) == 2:
# # # #                 doc_row = proj_table.add_row().cells

# # # #                 left_cell = doc_row[0]
# # # #                 left_text = cells[0].get_text(strip=True)
# # # #                 set_cell_content(left_cell, left_text, is_header=True, center_align=True)
                
# # # #                 right_cell = doc_row[1]
# # # #                 set_cell_content(right_cell, "", is_header=False, center_align=False)
# # # #                 process_project_content(right_cell, cells[1])
    
# # # #     doc.save(docx_file_path)

# # # # def set_table_borders(table):
# # # #     tbl = table._tbl
# # # #     tblPr = tbl.tblPr
# # # #     tblBorders = OxmlElement('w:tblBorders')
# # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # #         border = OxmlElement(f'w:{border_name}')
# # # #         border.set(qn('w:val'), 'single')
# # # #         border.set(qn('w:sz'), '4')
# # # #         border.set(qn('w:space'), '0')
# # # #         border.set(qn('w:color'), '000000')
# # # #         tblBorders.append(border)
# # # #     tblPr.append(tblBorders)

# # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12):
# # # #     cell.text = ""
# # # #     paragraph = cell.paragraphs[0]
# # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # #     paragraph.paragraph_format.space_after = Pt(0)
    
# # # #     if center_align:
# # # #         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# # # #     else:
# # # #         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
# # # #     if is_header:
# # # #         set_cell_background(cell, "0d3d62")
# # # #         run = paragraph.add_run(text)
# # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # #         run.font.bold = True
# # # #         run.font.size = Pt(font_size)
# # # #         run.font.name = 'Arial'
# # # #         cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# # # #     else:
# # # #         if text:
# # # #             run = paragraph.add_run(text)
# # # #             run.font.size = Pt(font_size)
# # # #             run.font.name = 'Arial'
# # # #             run.font.bold = False

# # # # def set_cell_background(cell, color_hex):
# # # #     shading_elm = OxmlElement('w:shd')
# # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # def process_html_content(cell, html_element):
# # # #     for element in html_element.contents:
# # # #         if isinstance(element, Tag):
# # # #             if element.name == 'div' and 'section-title' in element.get('class', []):
# # # #                 para = cell.add_paragraph()
# # # #                 para.paragraph_format.space_before = Pt(8)
# # # #                 para.paragraph_format.space_after = Pt(4)
# # # #                 run = para.add_run(element.get_text(strip=True))
# # # #                 run.bold = True
# # # #                 run.font.size = Pt(14)
# # # #                 run.font.name = 'Arial'
# # # #                 run.font.color.rgb = RGBColor(0, 0, 0)
# # # #             elif element.name == 'ul':
# # # #                 for li in element.find_all('li'):
# # # #                     para = cell.add_paragraph()
# # # #                     para.style = 'List Bullet'
# # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # #                     para.paragraph_format.space_before = Pt(0)
# # # #                     para.paragraph_format.space_after = Pt(6)
                    
# # # #                     strong = li.find('strong')
# # # #                     if strong:
# # # #                         strong_run = para.add_run(strong.get_text(strip=True))
# # # #                         strong_run.bold = True
# # # #                         strong_run.font.size = Pt(12)
# # # #                         strong_run.font.name = 'Arial'
                        
# # # #                         remaining_text = li.get_text().replace(strong.get_text(), '', 1).strip()
# # # #                         if remaining_text:
# # # #                             para.add_run(" ").font.size = Pt(12)
# # # #                             regular_run = para.add_run(remaining_text)
# # # #                             regular_run.font.size = Pt(12)
# # # #                             regular_run.font.name = 'Arial'
# # # #                     else:
# # # #                         run = para.add_run(li.get_text(strip=True))
# # # #                         run.font.size = Pt(12)
# # # #                         run.font.name = 'Arial'
# # # #             else:
# # # #                 process_html_content(cell, element)
# # # #         elif isinstance(element, NavigableString):
# # # #             text = str(element).strip()
# # # #             if text:
# # # #                 para = cell.add_paragraph()
# # # #                 para.paragraph_format.space_after = Pt(4)
# # # #                 run = para.add_run(text)
# # # #                 run.font.size = Pt(12)
# # # #                 run.font.name = 'Arial'
# # # #                 run.font.bold = 'strong' in str(html_element)
# # # #                 run.font.color.rgb = RGBColor(0, 0, 0)

# # # # def process_project_content(cell, html_element):
# # # #     cell.text = ""
    
# # # #     # 1. Project Heading
# # # #     project_heading = html_element.find('div', {'class': 'project-heading'})
# # # #     if project_heading:
# # # #         para = cell.add_paragraph()
# # # #         para.paragraph_format.space_after = Pt(3)
# # # #         run = para.add_run(project_heading.get_text(strip=True))
# # # #         run.bold = True
# # # #         run.font.size = Pt(13)
# # # #         run.font.name = 'Arial'
    
# # # #     # 2. Extract Role and Duration
# # # #     role_heading = html_element.find('div', {'class': 'role-heading'})
# # # #     if role_heading:
# # # #         para = cell.add_paragraph()
# # # #         para.paragraph_format.space_after = Pt(2)
            
# # # #         role_label = para.add_run("Role: ")
# # # #         role_label.bold = True
# # # #         role_label.font.size = Pt(12)
# # # #         role_label.font.name = 'Arial'
            
# # # #         text = role_heading.get_text(strip=True)
# # # #         val = text.replace("Role:", "", 1).strip()
# # # #         role_val = para.add_run(val)
# # # #         role_val.font.size = Pt(12)
# # # #         role_val.font.name = 'Arial'

# # # #     duration_heading = html_element.find('div', {'class': 'duration-heading'})
# # # #     if duration_heading:
# # # #         para = cell.add_paragraph()
# # # #         para.paragraph_format.space_after = Pt(2)
            
# # # #         dur_label = para.add_run("Duration: ")
# # # #         dur_label.bold = True
# # # #         dur_label.font.size = Pt(12)
# # # #         dur_label.font.name = 'Arial'
            
# # # #         text = duration_heading.get_text(strip=True)
# # # #         val = text.replace("Duration:", "", 1).strip()
# # # #         dur_val = para.add_run(val)
# # # #         dur_val.font.size = Pt(12)
# # # #         dur_val.font.name = 'Arial'

# # # #     # 3. Link
# # # #     link = html_element.find('a')
# # # #     if link:
# # # #         para = cell.add_paragraph()
# # # #         para.paragraph_format.space_after = Pt(3)
# # # #         para.add_run("Link- ").font.size = Pt(12)
# # # #         run = para.add_run(link.get('href', ''))
# # # #         run.font.size = Pt(12)
# # # #         run.font.name = 'Arial'
    
# # # #     # 4. Description (UPDATED to handle bullet points)
# # # #     # First check if there is a list
# # # #     ul_element = html_element.find('ul')
# # # #     if ul_element:
# # # #         for li in ul_element.find_all('li'):
# # # #             para = cell.add_paragraph()
# # # #             para.style = 'List Bullet'
# # # #             para.paragraph_format.left_indent = Inches(0.25)
# # # #             para.paragraph_format.space_after = Pt(3)
            
# # # #             run = para.add_run(li.get_text(strip=True))
# # # #             run.font.size = Pt(12)
# # # #             run.font.name = 'Arial'
# # # #     else:
# # # #         # Fallback to paragraph if no list is found
# # # #         description = html_element.find('p')
# # # #         if description:
# # # #             para = cell.add_paragraph()
# # # #             para.paragraph_format.space_before = Pt(5)
# # # #             para.paragraph_format.space_after = Pt(3)
# # # #             run = para.add_run(description.get_text(strip=True))
# # # #             run.font.size = Pt(12)
# # # #             run.font.name = 'Arial'
    
# # # #     # 5. Tech Stack
# # # #     techstack = html_element.find('div', {'class': 'techstack'})
# # # #     if techstack:
# # # #         para = cell.add_paragraph()
# # # #         para.paragraph_format.space_before = Pt(8)
# # # #         strong = techstack.find('strong')
# # # #         if strong:
# # # #             strong_run = para.add_run(strong.get_text(strip=True))
# # # #             strong_run.bold = True
# # # #             strong_run.font.size = Pt(12)
# # # #             strong_run.font.name = 'Arial'

# # # #             remaining_text = techstack.get_text().replace(strong.get_text(), '', 1).strip()
# # # #             if remaining_text:
# # # #                 para.add_run("\n")
# # # #                 regular_run = para.add_run(remaining_text)
# # # #                 regular_run.font.size = Pt(12)
# # # #                 regular_run.font.name = 'Arial'

# # # # def convert_resume_to_docx(html_file_path, docx_file_path=None):
# # # #     if docx_file_path is None:
# # # #         docx_file_path = os.path.splitext(html_file_path)[0] + '.docx'
# # # #     try:
# # # #         html_to_docx(html_file_path, docx_file_path)
# # # #     except Exception as e:
# # # #         print(f"Conversion failed: {e}")
# # # #         raise


# # # # new


# # # from docx import Document
# # # from docx.shared import Inches, Pt, RGBColor
# # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # from docx.oxml.ns import qn
# # # from docx.oxml import OxmlElement
# # # from bs4 import BeautifulSoup, Tag
# # # import os
# # # import re

# # # # === CONFIGURATION ===
# # # THEME_COLOR = "0d3d62" # Navy Blue

# # # def html_to_docx(html_file_path, docx_file_path):
# # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # #         html_content = file.read()
    
# # #     soup = BeautifulSoup(html_content, 'html.parser')
    
# # #     doc = Document()
    
# # #     # Set Margins
# # #     for section in doc.sections:
# # #         section.top_margin = Inches(0.5)
# # #         section.bottom_margin = Inches(0.5)
# # #         section.left_margin = Inches(0.5)
# # #         section.right_margin = Inches(0.5)
    
# # #     try:
# # #         # === LOGO ===
# # #         logo_path = "/home/ca/Projects/resume_converter/download.png"
# # #         if os.path.exists(logo_path):
# # #             logo_para = doc.add_paragraph()
# # #             logo_para.paragraph_format.space_before = Pt(12)
# # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # #             run = logo_para.add_run()
# # #             run.add_picture(logo_path, width=Inches(3.0))
# # #             logo_para.paragraph_format.space_after = Pt(12)
# # #         elif os.path.exists("download.png"):
# # #             logo_para = doc.add_paragraph()
# # #             run = logo_para.add_run()
# # #             run.add_picture("download.png", width=Inches(3.0))
# # #     except Exception as e:
# # #         print(f"Could not add logo: {e}")
    
# # #     # === MAIN TABLE ===
# # #     main_table = soup.find('table', {'class': 'resume-table'})
# # #     if main_table:
# # #         doc_table = doc.add_table(rows=0, cols=2)
# # #         doc_table.style = 'Table Grid'
# # #         doc_table.autofit = False
# # #         doc_table.allow_autofit = False
        
# # #         doc_table.columns[0].width = Inches(1.875)
# # #         doc_table.columns[1].width = Inches(5.625)
        
# # #         set_table_borders(doc_table)
        
# # #         rows = main_table.find_all('tr')
# # #         for row in rows:
# # #             cells = row.find_all('td')
# # #             if len(cells) == 2:
# # #                 doc_row = doc_table.add_row().cells
            
# # #                 left_cell = doc_row[0]
# # #                 left_text = cells[0].get_text(strip=True)
# # #                 set_cell_content(left_cell, left_text, is_header=True, center_align=False)
                
# # #                 right_cell = doc_row[1]
# # #                 set_cell_content(right_cell, "", is_header=False, center_align=True) 
# # #                 process_html_content(right_cell, cells[1])
        
# # #     # === PROJECTS SECTION ===
# # #     project_title_div = soup.find('div', {'class': 'project-title-bar'})
# # #     if not project_title_div: 
# # #         project_title_div = soup.find('div', {'class': 'project-title-bar-text-only'})

# # #     projects_table_html = soup.find('table', {'class': 'projects-table'})
    
# # #     if projects_table_html:
# # #         # 1. Add Project Header
# # #         if project_title_div:
# # #             p = doc.add_paragraph()
# # #             p.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # #             p.paragraph_format.space_before = Pt(12)
# # #             p.paragraph_format.space_after = Pt(4)
            
# # #             run = p.add_run(project_title_div.get_text(strip=True))
# # #             run.font.name = 'Arial'
# # #             run.bold = True
# # #             run.font.size = Pt(14)
# # #             run.font.color.rgb = RGBColor(255, 255, 255)
            
# # #             set_run_background(run, THEME_COLOR)

# # #         # 2. Projects Table
# # #         proj_table = doc.add_table(rows=0, cols=2)
# # #         proj_table.style = 'Table Grid'
# # #         proj_table.autofit = False
# # #         proj_table.allow_autofit = False
        
# # #         proj_table.columns[0].width = Inches(1.875)
# # #         proj_table.columns[1].width = Inches(5.625)
        
# # #         set_table_borders(proj_table)

# # #         rows = projects_table_html.find_all('tr')
# # #         for row in rows:
# # #             cells = row.find_all('td')
# # #             if len(cells) == 2:
# # #                 doc_row = proj_table.add_row().cells

# # #                 # Left Cell: "Project 1", etc.
# # #                 left_cell = doc_row[0]
# # #                 left_text = cells[0].get_text(strip=True)
# # #                 set_cell_content(left_cell, left_text, is_header=False, center_align=True, bold=True)
                
# # #                 # Right Cell: Details
# # #                 right_cell = doc_row[1]
# # #                 set_cell_content(right_cell, "", is_header=False, center_align=False)
# # #                 process_project_content(right_cell, cells[1])
    
# # #     doc.save(docx_file_path)

# # # def set_run_background(run, color_hex):
# # #     rPr = run._r.get_or_add_rPr()
# # #     shd = OxmlElement('w:shd')
# # #     shd.set(qn('w:val'), 'clear')
# # #     shd.set(qn('w:color'), 'auto')
# # #     shd.set(qn('w:fill'), color_hex)
# # #     rPr.append(shd)

# # # def set_table_borders(table):
# # #     tbl = table._tbl
# # #     tblPr = tbl.tblPr
# # #     tblBorders = OxmlElement('w:tblBorders')
# # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # #         border = OxmlElement(f'w:{border_name}')
# # #         border.set(qn('w:val'), 'single')
# # #         border.set(qn('w:sz'), '4')
# # #         border.set(qn('w:space'), '0')
# # #         border.set(qn('w:color'), '000000')
# # #         tblBorders.append(border)
# # #     tblPr.append(tblBorders)

# # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, bold=False):
# # #     cell.text = ""
# # #     paragraph = cell.paragraphs[0]
    
# # #     paragraph.paragraph_format.space_before = Pt(8)
# # #     paragraph.paragraph_format.space_after = Pt(8)
    
# # #     if center_align:
# # #         paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# # #     else:
# # #         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
# # #     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # #     if is_header:
# # #         set_cell_background(cell, THEME_COLOR)
# # #         run = paragraph.add_run(text)
# # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # #         run.font.bold = True
# # #         run.font.size = Pt(font_size)
# # #         run.font.name = 'Arial'
# # #     else:
# # #         if text:
# # #             run = paragraph.add_run(text)
# # #             run.font.size = Pt(font_size)
# # #             run.font.name = 'Arial'
# # #             run.font.bold = bold

# # # def set_cell_background(cell, color_hex):
# # #     shading_elm = OxmlElement('w:shd')
# # #     shading_elm.set(qn('w:fill'), color_hex)
# # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # def process_html_content(cell, html_element):
# # #     """Processes content for the Main Info table"""
# # #     text = html_element.get_text(strip=True)
# # #     if text:
# # #         para = cell.paragraphs[0]
# # #         # Just append text to the formatted paragraph
# # #         run = para.add_run(text)
# # #         run.font.size = Pt(12)
# # #         run.font.name = 'Arial'

# # # def process_project_content(cell, html_element):
# # #     """Processes content for the Projects table"""
# # #     cell.text = ""
    
# # #     # 1. Project Heading -> Format: "Project : [Name]"
# # #     project_heading = html_element.find('div', {'class': 'project-heading'})
# # #     if project_heading:
# # #         para = cell.add_paragraph()
# # #         para.paragraph_format.space_after = Pt(3)
        
# # #         raw_title = project_heading.get_text(strip=True)
# # #         clean_title = raw_title.replace("Project :", "").replace("Project:", "").strip()
        
# # #         run = para.add_run(f"Project : {clean_title}")
# # #         run.bold = True
# # #         run.font.size = Pt(13)
# # #         run.font.name = 'Arial'
    
# # #     # 2. Role
# # #     role_heading = html_element.find('div', {'class': 'role-heading'})
# # #     if role_heading:
# # #         para = cell.add_paragraph()
# # #         para.paragraph_format.space_after = Pt(2)
        
# # #         text = role_heading.get_text(strip=True)
# # #         val = text.replace("Role:", "", 1).strip()
        
# # #         lbl = para.add_run("Role: ")
# # #         lbl.bold = True
# # #         lbl.font.name = 'Arial'
# # #         lbl.font.size = Pt(12)
        
# # #         v = para.add_run(val)
# # #         v.font.name = 'Arial'
# # #         v.font.size = Pt(12)

# # #     # 3. Duration
# # #     duration_heading = html_element.find('div', {'class': 'duration-heading'})
# # #     if duration_heading:
# # #         para = cell.add_paragraph()
# # #         para.paragraph_format.space_after = Pt(2)
        
# # #         text = duration_heading.get_text(strip=True)
# # #         val = text.replace("Duration:", "", 1).strip()
        
# # #         lbl = para.add_run("Duration: ")
# # #         lbl.bold = True
# # #         lbl.font.name = 'Arial'
# # #         lbl.font.size = Pt(12)
        
# # #         v = para.add_run(val)
# # #         v.font.name = 'Arial'
# # #         v.font.size = Pt(12)

# # #     # 4. Link
# # #     link = html_element.find('a')
# # #     if link:
# # #         para = cell.add_paragraph()
# # #         para.paragraph_format.space_after = Pt(3)
# # #         para.add_run("Link- ").font.size = Pt(12)
# # #         run = para.add_run(link.get('href', ''))
# # #         run.font.size = Pt(12)
# # #         run.font.name = 'Arial'
    
# # #     # 5. Description (Bullet Points)
# # #     ul_element = html_element.find('ul')
# # #     if ul_element:
# # #         for li in ul_element.find_all('li'):
# # #             para = cell.add_paragraph()
# # #             para.style = 'List Bullet' # Native Word Bullet Point
# # #             para.paragraph_format.left_indent = Inches(0.25)
# # #             para.paragraph_format.space_after = Pt(3)
            
# # #             run = para.add_run(li.get_text(strip=True))
# # #             run.font.size = Pt(12)
# # #             run.font.name = 'Arial'
# # #     else:
# # #         # Fallback to paragraph if no list is found
# # #         description = html_element.find('p')
# # #         if description:
# # #             para = cell.add_paragraph()
# # #             para.paragraph_format.space_before = Pt(5)
# # #             para.paragraph_format.space_after = Pt(3)
# # #             run = para.add_run(description.get_text(strip=True))
# # #             run.font.size = Pt(12)
# # #             run.font.name = 'Arial'
    
# # #     # 6. Tech Stack
# # #     techstack = html_element.find('div', {'class': 'techstack'})
# # #     if techstack:
# # #         para = cell.add_paragraph()
# # #         para.paragraph_format.space_before = Pt(8)
# # #         strong = techstack.find('strong')
# # #         if strong:
# # #             strong_run = para.add_run(strong.get_text(strip=True))
# # #             strong_run.bold = True
# # #             strong_run.font.size = Pt(12)
# # #             strong_run.font.name = 'Arial'

# # #             remaining_text = techstack.get_text().replace(strong.get_text(), '', 1).strip()
# # #             if remaining_text:
# # #                 para.add_run("\n")
# # #                 regular_run = para.add_run(remaining_text)
# # #                 regular_run.font.size = Pt(12)
# # #                 regular_run.font.name = 'Arial'

# # # def convert_resume_to_docx(html_file_path, docx_file_path=None):
# # #     if docx_file_path is None:
# # #         docx_file_path = os.path.splitext(html_file_path)[0] + '.docx'
# # #     try:
# # #         html_to_docx(html_file_path, docx_file_path)
# # #     except Exception as e:
# # #         print(f"Conversion failed: {e}")
# # #         raise

# from docx import Document
# from docx.shared import Inches, Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from bs4 import BeautifulSoup, Tag
# import os

# THEME_COLOR = "0d3d62"

# def html_to_docx(html_file_path, docx_file_path):
#     with open(html_file_path, 'r', encoding='utf-8') as file:
#         html_content = file.read()
    
#     soup = BeautifulSoup(html_content, 'html.parser')
#     doc = Document()
    
#     for section in doc.sections:
#         section.top_margin = Inches(0.5)
#         section.bottom_margin = Inches(0.5)
#         section.left_margin = Inches(0.5)
#         section.right_margin = Inches(0.5)
    
#     try:
#         logo_path = "/home/ca/Projects/resume_converter/download.png"
#         if os.path.exists(logo_path):
#             logo_para = doc.add_paragraph()
#             logo_para.paragraph_format.space_before = Pt(12)
#             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
#             run = logo_para.add_run()
#             run.add_picture(logo_path, width=Inches(3.0))
#             logo_para.paragraph_format.space_after = Pt(12)
#     except Exception: pass
    
#     # === MAIN TABLE ===
#     main_table = soup.find('table', {'class': 'resume-table'})
#     if main_table:
#         doc_table = doc.add_table(rows=0, cols=2)
#         doc_table.style = 'Table Grid'
#         doc_table.columns[0].width = Inches(1.875)
#         doc_table.columns[1].width = Inches(5.625)
#         set_table_borders(doc_table)
        
#         for row in main_table.find_all('tr'):
#             cells = row.find_all('td')
#             if len(cells) == 2:
#                 doc_row = doc_table.add_row().cells
#                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=True)
#                 set_cell_content(doc_row[1], "", is_header=False, center_align=True) 
                
#                 # Use SPECIAL processing for the Right Cell (Expertise etc.)
#                 process_right_col_content(doc_row[1], cells[1])
        
#     # === PROJECTS SECTION ===
#     project_title_div = soup.find('div', {'class': 'project-title-bar'})
#     if not project_title_div: project_title_div = soup.find('div', {'class': 'project-title-bar-text-only'})

#     projects_table_html = soup.find('table', {'class': 'projects-table'})
    
#     if projects_table_html:
#         if project_title_div:
#             p = doc.add_paragraph()
#             p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#             p.paragraph_format.space_before = Pt(12)
#             p.paragraph_format.space_after = Pt(4)
#             run = p.add_run(project_title_div.get_text(strip=True))
#             run.font.name = 'Arial'
#             run.bold = True
#             run.font.size = Pt(14)
#             run.font.color.rgb = RGBColor(255, 255, 255)
#             set_run_background(run, THEME_COLOR)

#         proj_table = doc.add_table(rows=0, cols=2)
#         proj_table.style = 'Table Grid'
#         proj_table.columns[0].width = Inches(1.875)
#         proj_table.columns[1].width = Inches(5.625)
#         set_table_borders(proj_table)

#         for row in projects_table_html.find_all('tr'):
#             cells = row.find_all('td')
#             if len(cells) == 2:
#                 doc_row = proj_table.add_row().cells
#                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, bold=True)
#                 set_cell_content(doc_row[1], "", is_header=False, center_align=False)
#                 process_project_content(doc_row[1], cells[1])
    
#     doc.save(docx_file_path)

# def set_run_background(run, color_hex):
#     rPr = run._r.get_or_add_rPr()
#     shd = OxmlElement('w:shd')
#     shd.set(qn('w:val'), 'clear')
#     shd.set(qn('w:color'), 'auto')
#     shd.set(qn('w:fill'), color_hex)
#     rPr.append(shd)

# def set_table_borders(table):
#     tbl = table._tbl
#     tblPr = tbl.tblPr
#     tblBorders = OxmlElement('w:tblBorders')
#     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
#         border = OxmlElement(f'w:{border_name}')
#         border.set(qn('w:val'), 'single')
#         border.set(qn('w:sz'), '4')
#         border.set(qn('w:space'), '0')
#         border.set(qn('w:color'), '000000')
#         tblBorders.append(border)
#     tblPr.append(tblBorders)

# def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, bold=False):
#     cell.text = ""
#     paragraph = cell.paragraphs[0]
#     paragraph.paragraph_format.space_before = Pt(8)
#     paragraph.paragraph_format.space_after = Pt(8)
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
#     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

#     if is_header:
#         set_cell_background(cell, THEME_COLOR)
#         run = paragraph.add_run(text)
#         run.font.color.rgb = RGBColor(255, 255, 255)
#         run.font.bold = True
#         run.font.size = Pt(font_size)
#         run.font.name = 'Arial'
#     else:
#         if text:
#             run = paragraph.add_run(text)
#             run.font.size = Pt(font_size)
#             run.font.name = 'Arial'
#             run.font.bold = bold

# def set_cell_background(cell, color_hex):
#     shading_elm = OxmlElement('w:shd')
#     shading_elm.set(qn('w:fill'), color_hex)
#     cell._tc.get_or_add_tcPr().append(shading_elm)

# # === NEW FUNCTION: INTELLIGENT EXPERTISE PARSING ===
# def process_right_col_content(cell, html_element):
#     """Parses Expertise section to maintain bolding and structure."""
    
#     # 1. Check for Expertise Containers
#     expertise_categories = html_element.find_all('div', {'class': 'expertise-category'})
    
#     if expertise_categories:
#         # If we have structured expertise (from the template update)
#         for cat in expertise_categories:
#             # Category Title (e.g. "Salesforce Development")
#             title_div = cat.find('div', {'class': 'section-title'})
#             if title_div:
#                 para = cell.add_paragraph()
#                 para.paragraph_format.space_before = Pt(12)
#                 run = para.add_run(title_div.get_text(strip=True))
#                 run.bold = True
#                 run.underline = True
#                 run.font.name = 'Arial'
#                 run.font.size = Pt(12)

#             # Skill Items (e.g. "Core Salesforce: Apex...")
#             skills = cat.find_all('div', {'class': 'skill-item'})
#             for skill in skills:
#                 para = cell.add_paragraph()
#                 para.paragraph_format.space_after = Pt(2)
                
#                 # Extract Strong Label (Sub-category)
#                 strong_tag = skill.find('strong')
#                 if strong_tag:
#                     label = strong_tag.get_text(strip=True)
#                     run = para.add_run(label + " ")
#                     run.bold = True
#                     run.font.name = 'Arial'
#                     run.font.size = Pt(12)
                    
#                     # Remaining text
#                     val = skill.get_text().replace(label, '', 1).strip()
#                     run2 = para.add_run(val)
#                     run2.font.name = 'Arial'
#                     run2.font.size = Pt(12)
#                 else:
#                     run = para.add_run(skill.get_text(strip=True))
#                     run.font.name = 'Arial'
#                     run.font.size = Pt(12)
#     else:
#         # Fallback for simple text (Name, Experience fields)
#         text = html_element.get_text(strip=True)
#         if text:
#             para = cell.paragraphs[0]
#             run = para.add_run(text)
#             run.font.size = Pt(12)
#             run.font.name = 'Arial'

# def process_project_content(cell, html_element):
#     cell.text = ""
    
#     # 1. Project Heading
#     heading = html_element.find('div', {'class': 'project-heading'})
#     if heading:
#         para = cell.add_paragraph()
#         para.paragraph_format.space_after = Pt(3)
#         raw = heading.get_text(strip=True).replace("Project :", "").replace("Project:", "").strip()
#         run = para.add_run(f"Project : {raw}")
#         run.bold = True
#         run.font.size = Pt(13)
#         run.font.name = 'Arial'
    
#     # 2. Role
#     role = html_element.find('div', {'class': 'role-heading'})
#     if role:
#         para = cell.add_paragraph()
#         para.paragraph_format.space_after = Pt(2)
#         val = role.get_text(strip=True).replace("Role:", "", 1).strip()
#         para.add_run("Role: ").bold = True
#         para.add_run(val)

#     # 3. Duration
#     dur = html_element.find('div', {'class': 'duration-heading'})
#     if dur:
#         para = cell.add_paragraph()
#         para.paragraph_format.space_after = Pt(2)
#         val = dur.get_text(strip=True).replace("Duration:", "", 1).strip()
#         para.add_run("Duration: ").bold = True
#         para.add_run(val)

#     # 4. Links & Description
#     link = html_element.find('a')
#     if link:
#         para = cell.add_paragraph()
#         para.add_run("Link- ")
#         para.add_run(link.get('href', ''))

#     ul = html_element.find('ul')
#     if ul:
#         for li in ul.find_all('li'):
#             para = cell.add_paragraph()
#             para.style = 'List Bullet'
#             para.paragraph_format.left_indent = Inches(0.25)
#             para.add_run(li.get_text(strip=True))

#     tech = html_element.find('div', {'class': 'techstack'})
#     if tech:
#         para = cell.add_paragraph()
#         para.paragraph_format.space_before = Pt(8)
#         strong = tech.find('strong')
#         if strong:
#             para.add_run(strong.get_text(strip=True)).bold = True
#             val = tech.get_text().replace(strong.get_text(), '', 1).strip()
#             if val: para.add_run("\n" + val)

# def convert_resume_to_docx(html_file_path, docx_file_path=None):
#     if docx_file_path is None:
#         docx_file_path = os.path.splitext(html_file_path)[0] + '.docx'
#     try:
#         html_to_docx(html_file_path, docx_file_path)
#     except Exception as e:
#         print(f"Conversion failed: {e}")
#         raise

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, Tag
import os

THEME_COLOR = "0d3d62"

def html_to_docx(html_file_path, docx_file_path):
    with open(html_file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    try:
        logo_path = "/home/ca/Projects/resume_converter/download.png"
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.paragraph_format.space_before = Pt(12)
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(3.0))
            logo_para.paragraph_format.space_after = Pt(12)
    except Exception: pass
    
    # === MAIN TABLE ===
    main_table = soup.find('table', {'class': 'resume-table'})
    if main_table:
        doc_table = doc.add_table(rows=0, cols=2)
        doc_table.style = 'Table Grid'
        doc_table.columns[0].width = Inches(1.875)
        doc_table.columns[1].width = Inches(5.625)
        set_table_borders(doc_table)
        
        for row in main_table.find_all('tr'):
            cells = row.find_all('td')
            if len(cells) == 2:
                doc_row = doc_table.add_row().cells
                set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=True)
                set_cell_content(doc_row[1], "", is_header=False, center_align=True) 
                process_right_col_content(doc_row[1], cells[1])
        
    # === PROJECTS SECTION ===
    project_title_div = soup.find('div', {'class': 'project-title-bar'})
    if not project_title_div: project_title_div = soup.find('div', {'class': 'project-title-bar-text-only'})

    projects_table_html = soup.find('table', {'class': 'projects-table'})
    
    if projects_table_html:
        if project_title_div:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(project_title_div.get_text(strip=True))
            run.font.name = 'Arial'
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_run_background(run, THEME_COLOR)

        proj_table = doc.add_table(rows=0, cols=2)
        proj_table.style = 'Table Grid'
        proj_table.columns[0].width = Inches(1.875)
        proj_table.columns[1].width = Inches(5.625)
        set_table_borders(proj_table)

        for row in projects_table_html.find_all('tr'):
            cells = row.find_all('td')
            if len(cells) == 2:
                doc_row = proj_table.add_row().cells
                
                # === FIX: Project Left Column -> BLUE HEADER STYLE ===
                set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=True, center_align=True, bold=True)
                
                set_cell_content(doc_row[1], "", is_header=False, center_align=False)
                process_project_content(doc_row[1], cells[1])
    
    doc.save(docx_file_path)

def set_run_background(run, color_hex):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    rPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if is_header:
        set_cell_background(cell, THEME_COLOR)
        run = paragraph.add_run(text)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
    else:
        if text:
            run = paragraph.add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
            run.font.bold = bold

def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def process_right_col_content(cell, html_element):
    expertise_categories = html_element.find_all('div', {'class': 'expertise-category'})
    
    if expertise_categories:
        for cat in expertise_categories:
            title_div = cat.find('div', {'class': 'section-title'})
            if title_div:
                para = cell.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                run = para.add_run(title_div.get_text(strip=True))
                run.bold = True
                
                # === FIX: NO UNDERLINE ===
                run.underline = False
                
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            skills = cat.find_all('div', {'class': 'skill-item'})
            for skill in skills:
                para = cell.add_paragraph()
                para.paragraph_format.space_after = Pt(2)
                
                strong_tag = skill.find('strong')
                if strong_tag:
                    label = strong_tag.get_text(strip=True)
                    run = para.add_run(label + " ")
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    
                    val = skill.get_text().replace(label, '', 1).strip()
                    run2 = para.add_run(val)
                    run2.font.name = 'Arial'
                    run2.font.size = Pt(12)
                else:
                    run = para.add_run(skill.get_text(strip=True))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
    else:
        text = html_element.get_text(strip=True)
        if text:
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.name = 'Arial'

def process_project_content(cell, html_element):
    cell.text = ""
    
    # 1. Project Heading
    heading = html_element.find('div', {'class': 'project-heading'})
    if heading:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        raw = heading.get_text(strip=True).replace("Project :", "").replace("Project:", "").strip()
        run = para.add_run(f"Project : {raw}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Arial'
    
    # 2. Role
    role = html_element.find('div', {'class': 'role-heading'})
    if role:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        val = role.get_text(strip=True).replace("Role:", "", 1).strip()
        para.add_run("Role: ").bold = True
        para.add_run(val)

    # 3. Duration
    dur = html_element.find('div', {'class': 'duration-heading'})
    if dur:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        val = dur.get_text(strip=True).replace("Duration:", "", 1).strip()
        para.add_run("Duration: ").bold = True
        para.add_run(val)

    # 4. Links & Description
    link = html_element.find('a')
    if link:
        para = cell.add_paragraph()
        para.add_run("Link- ")
        para.add_run(link.get('href', ''))

    ul = html_element.find('ul')
    if ul:
        for li in ul.find_all('li'):
            para = cell.add_paragraph()
            para.style = 'List Bullet'
            para.paragraph_format.left_indent = Inches(0.25)
            para.add_run(li.get_text(strip=True))

    tech = html_element.find('div', {'class': 'techstack'})
    if tech:
        para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        strong = tech.find('strong')
        if strong:
            para.add_run(strong.get_text(strip=True)).bold = True
            val = tech.get_text().replace(strong.get_text(), '', 1).strip()
            if val: para.add_run("\n" + val)

def convert_resume_to_docx(html_file_path, docx_file_path=None):
    if docx_file_path is None:
        docx_file_path = os.path.splitext(html_file_path)[0] + '.docx'
    try:
        html_to_docx(html_file_path, docx_file_path)
    except Exception as e:
        print(f"Conversion failed: {e}")
        raise