# # # # # # # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # # # # # # import os
# # # # # # # # # # # # # # # # # import json
# # # # # # # # # # # # # # # # # import time
# # # # # # # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # # # # # # import re
# # # # # # # # # # # # # # # # # import docx
# # # # # # # # # # # # # # # # # import traceback

# # # # # # # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # # # # # # #     return genai

# # # # # # # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # # # # # # def transform_expertise_data(expertise_list):
# # # # # # # # # # # # # # # # #     transformed_expertise = []
# # # # # # # # # # # # # # # # #     if not expertise_list: return []
# # # # # # # # # # # # # # # # #     for category in expertise_list:
# # # # # # # # # # # # # # # # #         transformed_skills = []
# # # # # # # # # # # # # # # # #         for skill in category.get("skills", []):
# # # # # # # # # # # # # # # # #             details = skill.get("details", [])
# # # # # # # # # # # # # # # # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # # # # # # # # # # # # # # # #             transformed_skills.append({"area": skill.get("skill_name", ""), "technologies": technologies})
# # # # # # # # # # # # # # # # #         transformed_expertise.append({"category_name": category.get("category_name", ""), "skills": transformed_skills})
# # # # # # # # # # # # # # # # #     return transformed_expertise

# # # # # # # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # # # # # # #         # Format Title: "Project Name (Clouds)"
# # # # # # # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # # # # # # #         clouds = project.get("Clouds", "")
# # # # # # # # # # # # # # # # #         if clouds and clouds not in ["NA", "None"]:
# # # # # # # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # # # # # # #         else:
# # # # # # # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # # # # # # #         # Ensure description is a list for numbering/bullets
# # # # # # # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # # # # # # #             desc = [desc]
            
# # # # # # # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # # # # # # #             "description": desc,
# # # # # # # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # # # # # # #         })
# # # # # # # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # # # # # # #     full_text = []
# # # # # # # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. Extract information EXACTLY as written.
# # # # # # # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # # # # # # #     1. For 'description' in projects, break the text into distinct bullet points.
# # # # # # # # # # # # # # # # #     2. Extract 'Industry' and 'Clouds' for each project specifically.
# # # # # # # # # # # # # # # # #     3. Output ONLY valid JSON."""

# # # # # # # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.
    
# # # # # # # # # # # # # # # # #     Look specifically for these headers found in the sidebar:
# # # # # # # # # # # # # # # # #     - CRM Administration
# # # # # # # # # # # # # # # # #     - Certifications
# # # # # # # # # # # # # # # # #     - Salesforce Expertise
# # # # # # # # # # # # # # # # #     - Languages
# # # # # # # # # # # # # # # # #     - Salesforce Components
# # # # # # # # # # # # # # # # #     - Ticketing & Case Management
# # # # # # # # # # # # # # # # #     - Database
# # # # # # # # # # # # # # # # #     - Salesforce Clouds
# # # # # # # # # # # # # # # # #     - Soft Skills

# # # # # # # # # # # # # # # # # {
# # # # # # # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # # # # # # #       {
# # # # # # # # # # # # # # # # #         "project_name": "Project Title",
# # # # # # # # # # # # # # # # #         "Clouds": "Clouds used (e.g. Sales Cloud, Service Cloud)",
# # # # # # # # # # # # # # # # #         "Role": "Role",
# # # # # # # # # # # # # # # # #         "Industry": "Industry Name (e.g. Healthcare, Finance)",
# # # # # # # # # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # # # # # # # # #         "description": ["Point 1", "Point 2"], 
# # # # # # # # # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # # # # # # # # #         "link": "Link"
# # # # # # # # # # # # # # # # #       }
# # # # # # # # # # # # # # # # #     ]
# # # # # # # # # # # # # # # # #   }
# # # # # # # # # # # # # # # # # }"""
# # # # # # # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # # # # # # #     # === UPDATED MODEL LIST ===
# # # # # # # # # # # # # # # # #     models_to_try = [
# # # # # # # # # # # # # # # # #         "gemini-2.5-flash-lite", # Worked for General Parser
# # # # # # # # # # # # # # # # #         "gemini-1.5-flash",
# # # # # # # # # # # # # # # # #         "gemini-1.5-pro",
# # # # # # # # # # # # # # # # #         "gemini-pro"
# # # # # # # # # # # # # # # # #     ]
    
# # # # # # # # # # # # # # # # #     response = None
# # # # # # # # # # # # # # # # #     success = False

# # # # # # # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # # # # # # #         if success: break
# # # # # # # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # # # # # # # #         try:
# # # # # # # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # # # # # # #         except Exception:
# # # # # # # # # # # # # # # # #             continue
        
# # # # # # # # # # # # # # # # #         for attempt in range(2):
# # # # # # # # # # # # # # # # #             try:
# # # # # # # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # # # # # # #                 else:
# # # # # # # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # # # # # # #                 success = True
# # # # # # # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # # # # # # #                 break
# # # # # # # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # # # # # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # # # # # # # # # # # # # # #                 break 
# # # # # # # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting...")
# # # # # # # # # # # # # # # # #                 time.sleep(2)
# # # # # # # # # # # # # # # # #                 break
# # # # # # # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # # # # # # #                 break

# # # # # # # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # # # # # # #         if isinstance(val, list):
# # # # # # # # # # # # # # # # #             return ", ".join(val) if val else "NA"
# # # # # # # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # # # # # # #     template_data = {
# # # # # # # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # # # # # # #     }
    
# # # # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # # # # # # #         return True
# # # # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # # # # # # #         return False



# # # # # # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # # # # # import os
# # # # # # # # # # # # # # # # import json
# # # # # # # # # # # # # # # # import time
# # # # # # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # # # # # import re
# # # # # # # # # # # # # # # # import docx
# # # # # # # # # # # # # # # # import traceback

# # # # # # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # # # # # #     return genai

# # # # # # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # # # # # def transform_expertise_data(expertise_list):
# # # # # # # # # # # # # # # #     transformed_expertise = []
# # # # # # # # # # # # # # # #     if not expertise_list: return []
# # # # # # # # # # # # # # # #     for category in expertise_list:
# # # # # # # # # # # # # # # #         transformed_skills = []
# # # # # # # # # # # # # # # #         for skill in category.get("skills", []):
# # # # # # # # # # # # # # # #             details = skill.get("details", [])
# # # # # # # # # # # # # # # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # # # # # # # # # # # # # # #             transformed_skills.append({"area": skill.get("skill_name", ""), "technologies": technologies})
# # # # # # # # # # # # # # # #         transformed_expertise.append({"category_name": category.get("category_name", ""), "skills": transformed_skills})
# # # # # # # # # # # # # # # #     return transformed_expertise

# # # # # # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # # # # # #         clouds = project.get("Clouds", "")
# # # # # # # # # # # # # # # #         if clouds and clouds not in ["NA", "None"]:
# # # # # # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # # # # # #         else:
# # # # # # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # # # # # #             desc = [desc]
            
# # # # # # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # # # # # #             "description": desc,
# # # # # # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # # # # # #         })
# # # # # # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # # # # # #     full_text = []
# # # # # # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. Extract information EXACTLY as written.
# # # # # # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # # # # # #     1. For 'description' in projects, break the text into distinct bullet points.
# # # # # # # # # # # # # # # #     2. Extract 'Industry' and 'Clouds' for each project specifically.
# # # # # # # # # # # # # # # #     3. Output ONLY valid JSON."""

# # # # # # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.
    
# # # # # # # # # # # # # # # #     Look specifically for these headers found in the sidebar:
# # # # # # # # # # # # # # # #     - CRM Administration
# # # # # # # # # # # # # # # #     - Certifications
# # # # # # # # # # # # # # # #     - Salesforce Expertise
# # # # # # # # # # # # # # # #     - Languages
# # # # # # # # # # # # # # # #     - Salesforce Components
# # # # # # # # # # # # # # # #     - Ticketing & Case Management
# # # # # # # # # # # # # # # #     - Database
# # # # # # # # # # # # # # # #     - Salesforce Clouds
# # # # # # # # # # # # # # # #     - Soft Skills

# # # # # # # # # # # # # # # # {
# # # # # # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # # # # # #       {
# # # # # # # # # # # # # # # #         "project_name": "Project Title",
# # # # # # # # # # # # # # # #         "Clouds": "Clouds used (e.g. Sales Cloud, Service Cloud)",
# # # # # # # # # # # # # # # #         "Role": "Role",
# # # # # # # # # # # # # # # #         "Industry": "Industry Name (e.g. Healthcare, Finance)",
# # # # # # # # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # # # # # # # #         "description": ["Point 1", "Point 2"], 
# # # # # # # # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # # # # # # # #         "link": "Link"
# # # # # # # # # # # # # # # #       }
# # # # # # # # # # # # # # # #     ]
# # # # # # # # # # # # # # # #   }
# # # # # # # # # # # # # # # # }"""
# # # # # # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # # # # # #     # === UPDATED: VALID MODELS ONLY (From your check_models.py output) ===
# # # # # # # # # # # # # # # #     models_to_try = [
# # # # # # # # # # # # # # # #         "gemini-2.5-flash-lite", 
# # # # # # # # # # # # # # # #         "gemini-2.0-flash-lite",
# # # # # # # # # # # # # # # #         "gemini-2.5-flash",
# # # # # # # # # # # # # # # #         "gemini-2.0-flash",
# # # # # # # # # # # # # # # #         "gemini-2.5-pro",
# # # # # # # # # # # # # # # #     ]
    
# # # # # # # # # # # # # # # #     response = None
# # # # # # # # # # # # # # # #     success = False

# # # # # # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # # # # # #         if success: break
# # # # # # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
        
# # # # # # # # # # # # # # # #         try:
# # # # # # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # # # # # #         except Exception:
# # # # # # # # # # # # # # # #             continue
        
# # # # # # # # # # # # # # # #         # Retry loop for Quota Limits
# # # # # # # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # # # # # # #             try:
# # # # # # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # # # # # #                 else:
# # # # # # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # # # # # #                 success = True
# # # # # # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # # # # # #                 break
# # # # # # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # # # # # #                 # If we hit a limit, wait 10s and try again on the SAME model
# # # # # # # # # # # # # # # #                 # or let it fail and loop to the NEXT valid model
# # # # # # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 10s...")
# # # # # # # # # # # # # # # #                 time.sleep(10)
# # # # # # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # # # # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # # # # # # # # # # # # # #                 break 
# # # # # # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # # # # # #                 break

# # # # # # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # # # # # #         if isinstance(val, list):
# # # # # # # # # # # # # # # #             return ", ".join(val) if val else "NA"
# # # # # # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # # # # # #     template_data = {
# # # # # # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # # # # # #     }
    
# # # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # # # # # #         return True
# # # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # # # # import os
# # # # # # # # # # # # # # # import json
# # # # # # # # # # # # # # # import time
# # # # # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # # # # import re
# # # # # # # # # # # # # # # import docx
# # # # # # # # # # # # # # # import traceback

# # # # # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # # # # #     return genai

# # # # # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # # # # def transform_expertise_data(expertise_list):
# # # # # # # # # # # # # # #     transformed_expertise = []
# # # # # # # # # # # # # # #     if not expertise_list: return []
# # # # # # # # # # # # # # #     for category in expertise_list:
# # # # # # # # # # # # # # #         transformed_skills = []
# # # # # # # # # # # # # # #         for skill in category.get("skills", []):
# # # # # # # # # # # # # # #             details = skill.get("details", [])
# # # # # # # # # # # # # # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # # # # # # # # # # # # # #             transformed_skills.append({"area": skill.get("skill_name", ""), "technologies": technologies})
# # # # # # # # # # # # # # #         transformed_expertise.append({"category_name": category.get("category_name", ""), "skills": transformed_skills})
# # # # # # # # # # # # # # #     return transformed_expertise

# # # # # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # # # # #         clouds = project.get("Clouds", "")
# # # # # # # # # # # # # # #         # Use cloud name if found, otherwise just project name
# # # # # # # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # # # # #         else:
# # # # # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # # # # #             desc = [desc]
            
# # # # # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # # # # #             "description": desc,
# # # # # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # # # # #         })
# # # # # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # # # # #     full_text = []
# # # # # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array. Do not leave it empty.
# # # # # # # # # # # # # # #     2. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., Test Engineer, QA), map their technical skills (like Selenium, Java, Testing) to 'salesforce_expertise' and 'crm_administration' where appropriate, instead of leaving them 'NA'.
# # # # # # # # # # # # # # #     3. **Descriptions**: Break project/job descriptions into distinct bullet points.
# # # # # # # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # # # # # #     STRUCTURE GUIDANCE:
# # # # # # # # # # # # # # #     - **salesforce_expertise**: If no direct Salesforce skills, list primary technical skills (e.g., "Automation Testing", "Manual Testing").
# # # # # # # # # # # # # # #     - **crm_administration**: If no CRM skills, list Management/Admin tools used (e.g., "Jira", "TestRail").
# # # # # # # # # # # # # # #     - **projects_descriptions**: Extract EACH job role from the 'Experience' section as a separate project.

# # # # # # # # # # # # # # # {
# # # # # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # # # # #     "crm_administration": ["Skill 1", "Skill 2"],
# # # # # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1", "Expertise 2"],
# # # # # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1 (or 'Web/Mobile' if QA)"],
# # # # # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # # # # #       {
# # # # # # # # # # # # # # #         "project_name": "Company Name OR Project Name",
# # # # # # # # # # # # # # #         "Clouds": "Domain/Cloud used (e.g. Finance, Healthcare, Sales Cloud)",
# # # # # # # # # # # # # # #         "Role": "Job Title / Role",
# # # # # # # # # # # # # # #         "Industry": "Industry Name",
# # # # # # # # # # # # # # #         "Duration": "Dates / Duration",
# # # # # # # # # # # # # # #         "description": ["Responsibility 1", "Responsibility 2"], 
# # # # # # # # # # # # # # #         "tech_stack": ["Tool 1", "Tool 2"],
# # # # # # # # # # # # # # #         "link": ""
# # # # # # # # # # # # # # #       }
# # # # # # # # # # # # # # #     ]
# # # # # # # # # # # # # # #   }
# # # # # # # # # # # # # # # }"""
# # # # # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # # # # #     # Updated Model List for robustness
# # # # # # # # # # # # # # #     models_to_try = [
# # # # # # # # # # # # # # #         "gemini-2.5-flash-lite", 
# # # # # # # # # # # # # # #         "gemini-2.0-flash-lite",
# # # # # # # # # # # # # # #         "gemini-2.5-flash",
# # # # # # # # # # # # # # #         "gemini-2.0-flash",
# # # # # # # # # # # # # # #         "gemini-1.5-flash",
# # # # # # # # # # # # # # #         "gemini-pro"
# # # # # # # # # # # # # # #     ]
    
# # # # # # # # # # # # # # #     response = None
# # # # # # # # # # # # # # #     success = False

# # # # # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # # # # #         if success: break
# # # # # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # # # # # #         try:
# # # # # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # # # # #         except Exception:
# # # # # # # # # # # # # # #             continue
        
# # # # # # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # # # # # #             try:
# # # # # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # # # # #                 else:
# # # # # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # # # # #                 success = True
# # # # # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # # # # #                 break
# # # # # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 5s...")
# # # # # # # # # # # # # # #                 time.sleep(5)
# # # # # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # # # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # # # # # # # # # # # # #                 break 
# # # # # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # # # # #                 break

# # # # # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # # # # #         if isinstance(val, list):
# # # # # # # # # # # # # # #             return ", ".join(val) if val else "NA"
# # # # # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # # # # #     template_data = {
# # # # # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # # # # #     }
    
# # # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # # # # #         return True
# # # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # # # import os
# # # # # # # # # # # # # # import json
# # # # # # # # # # # # # # import time
# # # # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # # # import re
# # # # # # # # # # # # # # import docx
# # # # # # # # # # # # # # import traceback

# # # # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # # # #     return genai

# # # # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # # # def transform_expertise_data(expertise_list):
# # # # # # # # # # # # # #     """Transform expertise data from LLM format to template format"""
# # # # # # # # # # # # # #     transformed_expertise = []
# # # # # # # # # # # # # #     if not expertise_list: return []
# # # # # # # # # # # # # #     for category in expertise_list:
# # # # # # # # # # # # # #         transformed_skills = []
# # # # # # # # # # # # # #         for skill in category.get("skills", []):
# # # # # # # # # # # # # #             details = skill.get("details", [])
# # # # # # # # # # # # # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # # # # # # # # # # # # #             transformed_skills.append({"area": skill.get("skill_name", ""), "technologies": technologies})
# # # # # # # # # # # # # #         transformed_expertise.append({"category_name": category.get("category_name", ""), "skills": transformed_skills})
# # # # # # # # # # # # # #     return transformed_expertise

# # # # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # # # # # # #         # Use cloud name if found, otherwise just project name
# # # # # # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # # # #         else:
# # # # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # # # #         # Ensure description is a LIST for bullet points
# # # # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # # # #             desc = [desc]
            
# # # # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # # # #             "description": desc, # Guaranteed to be a list
# # # # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # # # #         })
# # # # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # # # #     full_text = []
# # # # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array.
# # # # # # # # # # # # # #     2. **Bullet Points**: Break project descriptions into a LIST of strings (bullet points). Do NOT return a single paragraph string.
# # # # # # # # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., Test Engineer, QA), map their technical skills (like Selenium, Java, Testing) to 'salesforce_expertise' and 'crm_administration' where appropriate.
# # # # # # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # # # # #     STRUCTURE GUIDANCE:
# # # # # # # # # # # # # #     - **salesforce_expertise**: If no direct Salesforce skills, list primary technical skills (e.g., "Automation Testing", "Manual Testing").
# # # # # # # # # # # # # #     - **crm_administration**: If no CRM skills, list Management/Admin tools used (e.g., "Jira", "TestRail").
# # # # # # # # # # # # # #     - **projects_descriptions**: Extract EACH job role from the 'Experience' section as a separate project.

# # # # # # # # # # # # # # {
# # # # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # # # #     "crm_administration": ["Skill 1", "Skill 2"],
# # # # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1", "Expertise 2"],
# # # # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1 (or 'Web/Mobile' if QA)"],
# # # # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # # # #       {
# # # # # # # # # # # # # #         "project_name": "Company Name OR Project Name",
# # # # # # # # # # # # # #         "Clouds": "Domain/Cloud used (e.g. Finance, Healthcare, Sales Cloud)",
# # # # # # # # # # # # # #         "Role": "Job Title / Role",
# # # # # # # # # # # # # #         "Industry": "Industry Name",
# # # # # # # # # # # # # #         "Duration": "Dates / Duration",
# # # # # # # # # # # # # #         "description": [
# # # # # # # # # # # # # #             "Bullet point 1 - Responsibility or achievement", 
# # # # # # # # # # # # # #             "Bullet point 2 - Tool used or metric"
# # # # # # # # # # # # # #         ], 
# # # # # # # # # # # # # #         "tech_stack": ["Tool 1", "Tool 2"],
# # # # # # # # # # # # # #         "link": ""
# # # # # # # # # # # # # #       }
# # # # # # # # # # # # # #     ]
# # # # # # # # # # # # # #   }
# # # # # # # # # # # # # # }"""
# # # # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # # # #     # === VALID MODELS ONLY ===
# # # # # # # # # # # # # #     models_to_try = [
# # # # # # # # # # # # # #         "gemini-2.5-flash-lite", 
# # # # # # # # # # # # # #         "gemini-2.0-flash-lite",
# # # # # # # # # # # # # #         "gemini-2.5-flash",
# # # # # # # # # # # # # #         "gemini-2.0-flash",
# # # # # # # # # # # # # #         "gemini-2.5-pro",
# # # # # # # # # # # # # #     ]
    
# # # # # # # # # # # # # #     response = None
# # # # # # # # # # # # # #     success = False

# # # # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # # # #         if success: break
# # # # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # # # # #         try:
# # # # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # # # #         except Exception:
# # # # # # # # # # # # # #             continue
        
# # # # # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # # # # #             try:
# # # # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # # # #                 else:
# # # # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # # # #                 success = True
# # # # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # # # #                 break
# # # # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 10s...")
# # # # # # # # # # # # # #                 time.sleep(10)
# # # # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # # # # # # # # # # # #                 break 
# # # # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # # # #                 break

# # # # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # # # #         return False

# # # # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # # # #         if isinstance(val, list):
# # # # # # # # # # # # # #             return ", ".join(val) if val else "NA"
# # # # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # # # #     template_data = {
# # # # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # # # #     }
    
# # # # # # # # # # # # # #     try:
# # # # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # # # #         return True
# # # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # # # #         return False 


# # # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # # import os
# # # # # # # # # # # # # import json
# # # # # # # # # # # # # import time
# # # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # # import re
# # # # # # # # # # # # # import docx
# # # # # # # # # # # # # import traceback

# # # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # # #     return genai

# # # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # # #         else:
# # # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # # #         # Ensure description is a LIST (for numbered list generation)
# # # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # # #             if '\n' in desc:
# # # # # # # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # # # # # # #             else:
# # # # # # # # # # # # #                 desc = [desc]
            
# # # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # # #             "description": desc,
# # # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # # #         })
# # # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # # #     full_text = []
# # # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # # #     1. **Experience = Projects**: Extract 'Work Experience' into the 'projects_descriptions' array.
# # # # # # # # # # # # #     2. **Numbered List**: Break project descriptions into a LIST of distinct strings so they can be numbered.
# # # # # # # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., QA), map technical skills to 'salesforce_expertise'.
# # # # # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # # # #     STRUCTURE GUIDANCE:
# # # # # # # # # # # # #     - **salesforce_expertise**: If no direct Salesforce skills, list primary technical skills.
# # # # # # # # # # # # #     - **crm_administration**: If no CRM skills, list Management/Admin tools (Jira, etc).
# # # # # # # # # # # # #     - **projects_descriptions**: Extract EACH job role as a separate project.

# # # # # # # # # # # # # {
# # # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # # #       {
# # # # # # # # # # # # #         "project_name": "Project Name",
# # # # # # # # # # # # #         "Clouds": "Domain/Cloud",
# # # # # # # # # # # # #         "Role": "Role",
# # # # # # # # # # # # #         "Industry": "Industry",
# # # # # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # # # # #         "description": [
# # # # # # # # # # # # #             "Point 1", 
# # # # # # # # # # # # #             "Point 2"
# # # # # # # # # # # # #         ], 
# # # # # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # # # # #         "link": ""
# # # # # # # # # # # # #       }
# # # # # # # # # # # # #     ]
# # # # # # # # # # # # #   }
# # # # # # # # # # # # # }"""
# # # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # # #     models_to_try = [
# # # # # # # # # # # # #         "gemini-2.5-flash-lite", 
# # # # # # # # # # # # #         "gemini-2.0-flash-lite",
# # # # # # # # # # # # #         "gemini-2.5-flash",
# # # # # # # # # # # # #         "gemini-2.0-flash",
# # # # # # # # # # # # #         "gemini-2.5-pro",
# # # # # # # # # # # # #     ]
    
# # # # # # # # # # # # #     response = None
# # # # # # # # # # # # #     success = False

# # # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # # #         if success: break
# # # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # # # #         try:
# # # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # # #         except Exception:
# # # # # # # # # # # # #             continue
        
# # # # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # # # #             try:
# # # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # # #                 else:
# # # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # # #                 success = True
# # # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # # #                 break
# # # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 10s...")
# # # # # # # # # # # # #                 time.sleep(10)
# # # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # # # # # # # # # # #                 break 
# # # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # # #                 break

# # # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # # #         return False

# # # # # # # # # # # # #     try:
# # # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # # #         return False

# # # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # # #         if isinstance(val, list):
# # # # # # # # # # # # #             return ", ".join(val) if val else "NA"
# # # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # # #     template_data = {
# # # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # # #     }
    
# # # # # # # # # # # # #     try:
# # # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # # #         return True
# # # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # # #         return False


# # # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # # import os
# # # # # # # # # # # # import json
# # # # # # # # # # # # import time
# # # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # # import re
# # # # # # # # # # # # import docx
# # # # # # # # # # # # import traceback

# # # # # # # # # # # # load_dotenv()

# # # # # # # # # # # # def configure_gemini():
# # # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # # #     return genai

# # # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # # #         else:
# # # # # # # # # # # #             formatted_title = name

# # # # # # # # # # # #         # Ensure description is a list for numbering
# # # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # # #             if '\n' in desc:
# # # # # # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # # # # # #             else:
# # # # # # # # # # # #                 desc = [desc]
            
# # # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # # #             "description": desc,
# # # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # # #         })
# # # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # # #     full_text = []
# # # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # # #     1. **Experience = Projects**: Extract 'Work Experience' into 'projects_descriptions'.
# # # # # # # # # # # #     2. **List Format**: Break descriptions into a list of strings so they can be numbered.
# # # # # # # # # # # #     3. **Skill Mapping**: Map technical skills to 'salesforce_expertise' if the candidate is not a Salesforce Dev.
# # # # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # # # {
# # # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # # #       {
# # # # # # # # # # # #         "project_name": "Project Name",
# # # # # # # # # # # #         "Clouds": "Cloud/Domain",
# # # # # # # # # # # #         "Role": "Role",
# # # # # # # # # # # #         "Industry": "Industry",
# # # # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # # # #         "description": ["Point 1", "Point 2"], 
# # # # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # # # #         "link": ""
# # # # # # # # # # # #       }
# # # # # # # # # # # #     ]
# # # # # # # # # # # #   }
# # # # # # # # # # # # }"""
# # # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # # # # # # #     response = None
# # # # # # # # # # # #     success = False

# # # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # # #         if success: break
# # # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # # #         try:
# # # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # # #         except Exception: continue
        
# # # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # # #             try:
# # # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # # #                 else:
# # # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # # #                 success = True
# # # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # # #                 break
# # # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 10s...")
# # # # # # # # # # # #                 time.sleep(10)
# # # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # # #                 break 
# # # # # # # # # # # #             except Exception as e:
# # # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # # #                 break

# # # # # # # # # # # #     if not success or not response:
# # # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # # #         return False

# # # # # # # # # # # #     try:
# # # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # # #         return False

# # # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # # #     def format_list(key):
# # # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # # #     template_data = {
# # # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # # #     }
    
# # # # # # # # # # # #     try:
# # # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # # #         return True
# # # # # # # # # # # #     except Exception as e:
# # # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # # #         return False



# # # # # # # # # # # #changed


# # # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # # import os
# # # # # # # # # # # import json
# # # # # # # # # # # import time
# # # # # # # # # # # from jinja2 import Template
# # # # # # # # # # # import re
# # # # # # # # # # # import docx
# # # # # # # # # # # import traceback

# # # # # # # # # # # load_dotenv()

# # # # # # # # # # # def configure_gemini():
# # # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # # #     return genai

# # # # # # # # # # # def clean_json_response(text):
# # # # # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # # # # #     transformed_projects = []
# # # # # # # # # # #     if not projects_list: return []
# # # # # # # # # # #     for project in projects_list:
# # # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # # #         else:
# # # # # # # # # # #             formatted_title = name

# # # # # # # # # # #         # Ensure description is a LIST (for numbered list generation)
# # # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # # #             if '\n' in desc:
# # # # # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # # # # #             else:
# # # # # # # # # # #                 desc = [desc]
            
# # # # # # # # # # #         transformed_projects.append({
# # # # # # # # # # #             "title": formatted_title,
# # # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # # #             "description": desc,
# # # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # # #         })
# # # # # # # # # # #     return transformed_projects

# # # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # # #     full_text = []
# # # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # # #     for table in doc.tables:
# # # # # # # # # # #         for row in table.rows:
# # # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array.
# # # # # # # # # # #     2. **List Format**: Break descriptions into a list of distinct strings so they can be numbered.
# # # # # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., QA, Tester), map their technical skills (Selenium, Java, etc.) to 'salesforce_expertise'.
# # # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # # {
# # # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # # #     "name": "Full Name",
# # # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # # #       {
# # # # # # # # # # #         "project_name": "Project Name",
# # # # # # # # # # #         "Clouds": "Domain/Cloud",
# # # # # # # # # # #         "Role": "Role",
# # # # # # # # # # #         "Industry": "Industry",
# # # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # # #         "description": ["Point 1", "Point 2"], 
# # # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # # #         "link": ""
# # # # # # # # # # #       }
# # # # # # # # # # #     ]
# # # # # # # # # # #   }
# # # # # # # # # # # }"""
# # # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # # # # # #     response = None
# # # # # # # # # # #     success = False

# # # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # # #         if success: break
# # # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # # #         try:
# # # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # # #         except Exception: continue
        
# # # # # # # # # # #         for attempt in range(3):
# # # # # # # # # # #             try:
# # # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # # #                 else:
# # # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # # #                 success = True
# # # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # # #                 break
# # # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # # #                 time.sleep(10)
# # # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # # #                 break 
# # # # # # # # # # #             except Exception as e:
# # # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # # #                 break

# # # # # # # # # # #     if not success or not response:
# # # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # # #         return False

# # # # # # # # # # #     try:
# # # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # # #     except Exception as e:
# # # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # # #         return False

# # # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # # #     def format_list(key):
# # # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # # #     template_data = {
# # # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # # #     }
    
# # # # # # # # # # #     try:
# # # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # # #             template = Template(f.read())
# # # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # # #         return True
# # # # # # # # # # #     except Exception as e:
# # # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # # #         return False

# # # # # # # # # # #font changes

# # # # # # # # # # import google.generativeai as genai
# # # # # # # # # # from google.api_core import exceptions
# # # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # # import os
# # # # # # # # # # import json
# # # # # # # # # # import time
# # # # # # # # # # from jinja2 import Template
# # # # # # # # # # import re
# # # # # # # # # # import docx
# # # # # # # # # # import traceback

# # # # # # # # # # load_dotenv()

# # # # # # # # # # def configure_gemini():
# # # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # # #     if not api_key_gemini:
# # # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # # #     return genai

# # # # # # # # # # def clean_json_response(text):
# # # # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # # #     return cleaned.strip()

# # # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # # # #     transformed_projects = []
# # # # # # # # # #     if not projects_list: return []
# # # # # # # # # #     for project in projects_list:
# # # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # # #         else:
# # # # # # # # # #             formatted_title = name

# # # # # # # # # #         # Ensure description is a LIST (for numbered list generation)
# # # # # # # # # #         desc = project.get("description", [])
# # # # # # # # # #         if isinstance(desc, str):
# # # # # # # # # #             # Clean up newlines if present
# # # # # # # # # #             if '\n' in desc:
# # # # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # # # #             else:
# # # # # # # # # #                 desc = [desc]
            
# # # # # # # # # #         transformed_projects.append({
# # # # # # # # # #             "title": formatted_title,
# # # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # # #             "description": desc,
# # # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # # #         })
# # # # # # # # # #     return transformed_projects

# # # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # # #     full_text = []
# # # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # # #     for table in doc.tables:
# # # # # # # # # #         for row in table.rows:
# # # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # # #     CRITICAL RULES:
# # # # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array.
# # # # # # # # # #     2. **List Format**: Break descriptions into a list of distinct strings so they can be numbered. Remove any existing bullet points (like '•' or '-') from the start of the sentences.
# # # # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., QA, Tester), map their technical skills (Selenium, Java, etc.) to 'salesforce_expertise'.
# # # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # # {
# # # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # # #     "name": "Full Name",
# # # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # # #     "database": ["DB 1"],
# # # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # # #     "projects_descriptions": [
# # # # # # # # # #       {
# # # # # # # # # #         "project_name": "Project Name",
# # # # # # # # # #         "Clouds": "Domain/Cloud",
# # # # # # # # # #         "Role": "Role",
# # # # # # # # # #         "Industry": "Industry",
# # # # # # # # # #         "Duration": "Duration",
# # # # # # # # # #         "description": ["Action verb task...", "Achieved result..."], 
# # # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # # #         "link": ""
# # # # # # # # # #       }
# # # # # # # # # #     ]
# # # # # # # # # #   }
# # # # # # # # # # }"""
# # # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # # === THIS IS THE FUNCTION YOUR API IS LOOKING FOR ===
# # # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # # #     genai = configure_gemini()
# # # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # # # # #     response = None
# # # # # # # # # #     success = False

# # # # # # # # # #     for model_name in models_to_try:
# # # # # # # # # #         if success: break
# # # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # # #         try:
# # # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # # #         except Exception: continue
        
# # # # # # # # # #         for attempt in range(3):
# # # # # # # # # #             try:
# # # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # # #                 else:
# # # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # # #                 success = True
# # # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # # #                 break
# # # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # # #                 time.sleep(10)
# # # # # # # # # #             except exceptions.NotFound:
# # # # # # # # # #                 break 
# # # # # # # # # #             except Exception as e:
# # # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # # #                 break

# # # # # # # # # #     if not success or not response:
# # # # # # # # # #         print("❌ All models failed.")
# # # # # # # # # #         return False

# # # # # # # # # #     try:
# # # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # # #     except Exception as e:
# # # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # # #         return False

# # # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # # #     def format_list(key):
# # # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # # # # #         return val if val else "NA"
    
# # # # # # # # # #     template_data = {
# # # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # # #         "database": format_list("database"),
# # # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # # #     }
    
# # # # # # # # # #     try:
# # # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # # #             template = Template(f.read())
# # # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # # #         return True
# # # # # # # # # #     except Exception as e:
# # # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # # #         return False


# # # # # # # # # import google.generativeai as genai
# # # # # # # # # from google.api_core import exceptions
# # # # # # # # # from dotenv import load_dotenv
# # # # # # # # # import os
# # # # # # # # # import json
# # # # # # # # # import time
# # # # # # # # # from jinja2 import Template
# # # # # # # # # import re
# # # # # # # # # import docx
# # # # # # # # # import traceback

# # # # # # # # # load_dotenv()

# # # # # # # # # def configure_gemini():
# # # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # # #     if not api_key_gemini:
# # # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # # #     return genai

# # # # # # # # # def clean_json_response(text):
# # # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # # #     return cleaned.strip()

# # # # # # # # # def transform_projects_data(projects_list):
# # # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # # #     transformed_projects = []
# # # # # # # # #     if not projects_list: return []
# # # # # # # # #     for project in projects_list:
# # # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # # #         else:
# # # # # # # # #             formatted_title = name

# # # # # # # # #         # Ensure description is a LIST (for numbered list generation)
# # # # # # # # #         desc = project.get("description", [])
# # # # # # # # #         if isinstance(desc, str):
# # # # # # # # #             # Clean up newlines if present
# # # # # # # # #             if '\n' in desc:
# # # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # # #             else:
# # # # # # # # #                 desc = [desc]
            
# # # # # # # # #         transformed_projects.append({
# # # # # # # # #             "title": formatted_title,
# # # # # # # # #             "role": project.get("Role", ""),
# # # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # # #             "link": project.get("link", ""),
# # # # # # # # #             "description": desc,
# # # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # # #         })
# # # # # # # # #     return transformed_projects

# # # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # # #     full_text = []
# # # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # # #     for table in doc.tables:
# # # # # # # # #         for row in table.rows:
# # # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # # #     return '\n'.join(full_text)

# # # # # # # # # def get_improved_extraction_prompt():
# # # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # # #     CRITICAL RULES:
# # # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array.
# # # # # # # # #     2. **List Format**: Break descriptions into a list of distinct strings so they can be numbered. Remove any existing bullet points (like '•' or '-') from the start of the sentences.
# # # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., QA, Tester), map their technical skills (Selenium, Java, etc.) to 'salesforce_expertise'.
# # # # # # # # #     4. Output ONLY valid JSON."""

# # # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # # {
# # # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # # #     "name": "Full Name",
# # # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # # #     "experience": "Total Experience",
    
# # # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # # #     "salesforce_expertise": ["Expertise 1"],
# # # # # # # # #     "languages": ["Language 1"],
# # # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # # #     "database": ["DB 1"],
# # # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # # #     "soft_skills": ["Skill 1"],

# # # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # # #     "projects_descriptions": [
# # # # # # # # #       {
# # # # # # # # #         "project_name": "Project Name",
# # # # # # # # #         "Clouds": "Domain/Cloud",
# # # # # # # # #         "Role": "Role",
# # # # # # # # #         "Industry": "Industry",
# # # # # # # # #         "Duration": "Duration",
# # # # # # # # #         "description": ["Action verb task...", "Achieved result..."], 
# # # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # # #         "link": ""
# # # # # # # # #       }
# # # # # # # # #     ]
# # # # # # # # #   }
# # # # # # # # # }"""
# # # # # # # # #     return system_prompt, user_prompt

# # # # # # # # # # === THIS IS THE FUNCTION THE ERROR IS SAYING IS MISSING ===
# # # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # # #     genai = configure_gemini()
# # # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # # # #     response = None
# # # # # # # # #     success = False

# # # # # # # # #     for model_name in models_to_try:
# # # # # # # # #         if success: break
# # # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # # #         try:
# # # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # # #         except Exception: continue
        
# # # # # # # # #         for attempt in range(3):
# # # # # # # # #             try:
# # # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # # #                 else:
# # # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # # #                 success = True
# # # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # # #                 break
# # # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # # #                 time.sleep(10)
# # # # # # # # #             except exceptions.NotFound:
# # # # # # # # #                 break 
# # # # # # # # #             except Exception as e:
# # # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # # #                 break

# # # # # # # # #     if not success or not response:
# # # # # # # # #         print("❌ All models failed.")
# # # # # # # # #         return False

# # # # # # # # #     try:
# # # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # # #         return False

# # # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # # #     def format_list(key):
# # # # # # # # #         val = main_data.get(key, [])
# # # # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # # # #         return val if val else "NA"
    
# # # # # # # # #     template_data = {
# # # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # # #         "languages": format_list("languages"),
# # # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # # #         "database": format_list("database"),
# # # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # # #     }
    
# # # # # # # # #     try:
# # # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # # #             template = Template(f.read())
# # # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # # #             f.write(template.render(**template_data))
# # # # # # # # #         return True
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # # #         return False


# # # # # # # # import google.generativeai as genai
# # # # # # # # from google.api_core import exceptions
# # # # # # # # from dotenv import load_dotenv
# # # # # # # # import os
# # # # # # # # import json
# # # # # # # # import time
# # # # # # # # from jinja2 import Template
# # # # # # # # import re
# # # # # # # # import docx
# # # # # # # # import traceback

# # # # # # # # load_dotenv()

# # # # # # # # def configure_gemini():
# # # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # # #     if not api_key_gemini:
# # # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # # #     return genai

# # # # # # # # def clean_json_response(text):
# # # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # # #     return cleaned.strip()

# # # # # # # # def transform_projects_data(projects_list):
# # # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # # #     transformed_projects = []
# # # # # # # #     if not projects_list: return []
# # # # # # # #     for project in projects_list:
# # # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # # #         name = project.get("project_name", "Unknown Project")
# # # # # # # #         clouds = project.get("Clouds", "")
        
# # # # # # # #         if clouds and clouds not in ["NA", "None", ""]:
# # # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # # #         else:
# # # # # # # #             formatted_title = name

# # # # # # # #         # Ensure description is a LIST (for numbered list generation)
# # # # # # # #         desc = project.get("description", [])
# # # # # # # #         if isinstance(desc, str):
# # # # # # # #             # Clean up newlines if present
# # # # # # # #             if '\n' in desc:
# # # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # # #             else:
# # # # # # # #                 desc = [desc]
            
# # # # # # # #         transformed_projects.append({
# # # # # # # #             "title": formatted_title,
# # # # # # # #             "role": project.get("Role", ""),
# # # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # # #             "duration": project.get("Duration", ""),
# # # # # # # #             "link": project.get("link", ""),
# # # # # # # #             "description": desc,
# # # # # # # #             "tech_stack": tech_stack_str
# # # # # # # #         })
# # # # # # # #     return transformed_projects

# # # # # # # # def extract_text_from_docx(docx_path):
# # # # # # # #     """Extract text content from a DOCX file"""
# # # # # # # #     doc = docx.Document(docx_path)
# # # # # # # #     full_text = []
# # # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # # #     for table in doc.tables:
# # # # # # # #         for row in table.rows:
# # # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # # #     return '\n'.join(full_text)

# # # # # # # # def get_improved_extraction_prompt():
# # # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # # #     CRITICAL RULES:
# # # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Employment History' into the 'projects_descriptions' array.
# # # # # # # #     2. **List Format**: Break descriptions into a list of distinct strings so they can be numbered. Remove any existing bullet points.
# # # # # # # #     3. **Skill Mapping**: If the candidate is NOT a Salesforce Developer (e.g., QA, Data Scientist), map their **TECHNICAL** skills (Selenium, Python, MLOps, APIs) to 'salesforce_expertise'. 
# # # # # # # #     4. **Soft Skills Strictness**: 'soft_skills' MUST ONLY contain behavioral traits (e.g., Communication, Leadership, Problem Solving). **NEVER** put technical skills like Agile, Scrum, MLOps, or Data Processing here. Move them to 'salesforce_expertise' or 'crm_administration'.
# # # # # # # #     5. Output ONLY valid JSON."""

# # # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # # {
# # # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # # #     "name": "Full Name",
# # # # # # # #     "current_job_role": "Current Job Title",
# # # # # # # #     "experience": "Total Experience",
    
# # # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # # #     "certifications": ["Cert 1"],
# # # # # # # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # # # # # # #     "languages": ["Language 1"],
# # # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # # #     "database": ["DB 1"],
# # # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # # #     "soft_skills": ["Behavioral Skill 1"],

# # # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # # #     "development_tool": ["Tool 1"],
# # # # # # # #     "projects_descriptions": [
# # # # # # # #       {
# # # # # # # #         "project_name": "Project Name",
# # # # # # # #         "Clouds": "Domain/Cloud",
# # # # # # # #         "Role": "Role",
# # # # # # # #         "Industry": "Industry",
# # # # # # # #         "Duration": "Duration",
# # # # # # # #         "description": ["Action verb task...", "Achieved result..."], 
# # # # # # # #         "tech_stack": ["Tech"],
# # # # # # # #         "link": ""
# # # # # # # #       }
# # # # # # # #     ]
# # # # # # # #   }
# # # # # # # # }"""
# # # # # # # #     return system_prompt, user_prompt

# # # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # # #     genai = configure_gemini()
# # # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # # #     response = None
# # # # # # # #     success = False

# # # # # # # #     for model_name in models_to_try:
# # # # # # # #         if success: break
# # # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # # #         try:
# # # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # # #         except Exception: continue
        
# # # # # # # #         for attempt in range(3):
# # # # # # # #             try:
# # # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # # #                 else:
# # # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # # #                 success = True
# # # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # # #                 break
# # # # # # # #             except exceptions.ResourceExhausted:
# # # # # # # #                 time.sleep(10)
# # # # # # # #             except exceptions.NotFound:
# # # # # # # #                 break 
# # # # # # # #             except Exception as e:
# # # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # # #                 break

# # # # # # # #     if not success or not response:
# # # # # # # #         print("❌ All models failed.")
# # # # # # # #         return False

# # # # # # # #     try:
# # # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # # #         return False

# # # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # # #     def format_list(key):
# # # # # # # #         val = main_data.get(key, [])
# # # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # # #         return val if val else "NA"
    
# # # # # # # #     template_data = {
# # # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # # #         "certifications": format_list("certifications"),
# # # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # # #         "languages": format_list("languages"),
# # # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # # #         "database": format_list("database"),
# # # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # # #     }
    
# # # # # # # #     try:
# # # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # # #             template = Template(f.read())
# # # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # # #             f.write(template.render(**template_data))
# # # # # # # #         return True
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # # #         return False


# # # # # # # import google.generativeai as genai
# # # # # # # from google.api_core import exceptions
# # # # # # # from dotenv import load_dotenv
# # # # # # # import os
# # # # # # # import json
# # # # # # # import time
# # # # # # # from jinja2 import Template
# # # # # # # import re
# # # # # # # import docx
# # # # # # # import traceback
# # # # # # # from datetime import datetime

# # # # # # # load_dotenv()

# # # # # # # def configure_gemini():
# # # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # # #     if not api_key_gemini:
# # # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # # #     return genai

# # # # # # # def clean_json_response(text):
# # # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # # #     return cleaned.strip()

# # # # # # # def calculate_duration_string(duration_raw):
# # # # # # #     """
# # # # # # #     Parses a duration string (e.g., "Oct 2023 - Present") and calculates months.
# # # # # # #     """
# # # # # # #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# # # # # # #         return "N/A"

# # # # # # #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
    
# # # # # # #     if len(parts) != 2:
# # # # # # #         return duration_raw

# # # # # # #     start_str = parts[0].strip()
# # # # # # #     end_str = parts[1].strip()
    
# # # # # # #     def parse_date(d_str):
# # # # # # #         if d_str.lower() in ["present", "current", "now", "ongoing"]:
# # # # # # #             return datetime.now()
        
# # # # # # #         formats = [
# # # # # # #             "%b %Y", "%B %Y",       # Oct 2023, October 2023
# # # # # # #             "%b-%Y", "%B-%Y",       # Oct-2023
# # # # # # #             "%m/%Y", "%Y"           # 10/2023, 2023
# # # # # # #         ]
        
# # # # # # #         for fmt in formats:
# # # # # # #             try:
# # # # # # #                 return datetime.strptime(d_str, fmt)
# # # # # # #             except ValueError:
# # # # # # #                 continue
# # # # # # #         return None

# # # # # # #     start_date = parse_date(start_str)
# # # # # # #     end_date = parse_date(end_str)

# # # # # # #     if start_date and end_date:
# # # # # # #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
        
# # # # # # #         if diff_months < 1: diff_months = 1
        
# # # # # # #         if diff_months >= 12:
# # # # # # #             years = diff_months // 12
# # # # # # #             months = diff_months % 12
# # # # # # #             duration_txt = f"{years} Years" + (f" {months} Months" if months > 0 else "")
# # # # # # #         else:
# # # # # # #             duration_txt = f"{diff_months} Months"
            
# # # # # # #         return f"{start_str} - {end_str} ({duration_txt})"
    
# # # # # # #     return duration_raw

# # # # # # # def transform_projects_data(projects_list):
# # # # # # #     """Transform projects data from LLM format to template format"""
# # # # # # #     transformed_projects = []
# # # # # # #     if not projects_list: return []
# # # # # # #     for project in projects_list:
# # # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # # #         # === 1. GET RAW VALUES ===
# # # # # # #         raw_name = project.get("project_name", "")
# # # # # # #         role = project.get("Role", "")
# # # # # # #         clouds = project.get("Clouds", "")

# # # # # # #         # === 2. FIX "NONE" PROJECT NAME ===
# # # # # # #         # If project name is missing/None, fallback to Role or Generic Title
# # # # # # #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown", "unknown project"]:
# # # # # # #             if role and role.strip().lower() not in ["none", "na"]:
# # # # # # #                 name = f"{role} Project" # Use Role (e.g., "Developer Project")
# # # # # # #             else:
# # # # # # #                 name = "Professional Experience" # Last resort
# # # # # # #         else:
# # # # # # #             name = raw_name

# # # # # # #         # === 3. FORMAT TITLE ===
# # # # # # #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# # # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # # #         else:
# # # # # # #             formatted_title = name

# # # # # # #         # === 4. HANDLE DESCRIPTION ===
# # # # # # #         desc = project.get("description", [])
# # # # # # #         if isinstance(desc, str):
# # # # # # #             if '\n' in desc:
# # # # # # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # # #             else:
# # # # # # #                 desc = [desc]
        
# # # # # # #         # === 5. CALCULATE DURATION ===
# # # # # # #         raw_duration = project.get("Duration", "")
# # # # # # #         calculated_duration = calculate_duration_string(raw_duration)

# # # # # # #         transformed_projects.append({
# # # # # # #             "title": formatted_title,
# # # # # # #             "role": role,
# # # # # # #             "industry": project.get("Industry", "NA"),
# # # # # # #             "duration": calculated_duration,
# # # # # # #             "link": project.get("link", ""),
# # # # # # #             "description": desc,
# # # # # # #             "tech_stack": tech_stack_str
# # # # # # #         })
# # # # # # #     return transformed_projects

# # # # # # # def extract_text_from_docx(docx_path):
# # # # # # #     """Extract text content from a DOCX file"""
# # # # # # #     doc = docx.Document(docx_path)
# # # # # # #     full_text = []
# # # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # # #     for table in doc.tables:
# # # # # # #         for row in table.rows:
# # # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # # #     return '\n'.join(full_text)

# # # # # # # def get_improved_extraction_prompt():
# # # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # # #     CRITICAL RULES:
# # # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience', 'Employment History', or 'Projects' sections into the 'projects_descriptions' array.
# # # # # # #     2. **Project Names**: If a specific project name is not listed, **USE THE COMPANY NAME** as the Project Name. Do NOT return 'None' or 'N/A'.
# # # # # # #     3. **List Format**: Break descriptions into a list of distinct strings. Remove bullet points.
# # # # # # #     4. **Skill Mapping**: Map technical skills (Selenium, Python, etc.) to 'salesforce_expertise'. 
# # # # # # #     5. **Soft Skills**: Only behavioral traits (Communication, Leadership) in 'soft_skills'. NO technical skills here.
# # # # # # #     6. **Industry Inference**: Infer the industry (e.g., 'Banking') from the description.
# # # # # # #     7. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# # # # # # #     8. Output ONLY valid JSON."""

# # # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # # {
# # # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # # #     "name": "Full Name",
# # # # # # #     "current_job_role": "Current Job Title",
# # # # # # #     "experience": "Total Experience",
    
# # # # # # #     "crm_administration": ["Skill 1"],
# # # # # # #     "certifications": ["Cert 1"],
# # # # # # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # # # # # #     "languages": ["Language 1"],
# # # # # # #     "salesforce_components": ["Component 1"],
# # # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # # #     "database": ["DB 1"],
# # # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # # #     "soft_skills": ["Behavioral Skill 1"],

# # # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # # #     "development_tool": ["Tool 1"],
# # # # # # #     "projects_descriptions": [
# # # # # # #       {
# # # # # # #         "project_name": "Company Name OR Project Title (Do NOT use 'None')",
# # # # # # #         "Clouds": "Specific Cloud/Domain (e.g. Sales Cloud, AI)",
# # # # # # #         "Role": "Job Role",
# # # # # # #         "Industry": "Inferred Industry",
# # # # # # #         "Duration": "Date Range",
# # # # # # #         "description": ["Task 1...", "Result 2..."], 
# # # # # # #         "tech_stack": ["Tech"],
# # # # # # #         "link": ""
# # # # # # #       }
# # # # # # #     ]
# # # # # # #   }
# # # # # # # }"""
# # # # # # #     return system_prompt, user_prompt

# # # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # # #     genai = configure_gemini()
# # # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # # #     response = None
# # # # # # #     success = False

# # # # # # #     for model_name in models_to_try:
# # # # # # #         if success: break
# # # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # # #         try:
# # # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # # #         except Exception: continue
        
# # # # # # #         for attempt in range(3):
# # # # # # #             try:
# # # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # # #                 else:
# # # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # # #                 success = True
# # # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # # #                 break
# # # # # # #             except exceptions.ResourceExhausted:
# # # # # # #                 time.sleep(10)
# # # # # # #             except exceptions.NotFound:
# # # # # # #                 break 
# # # # # # #             except Exception as e:
# # # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # # #                 break

# # # # # # #     if not success or not response:
# # # # # # #         print("❌ All models failed.")
# # # # # # #         return False

# # # # # # #     try:
# # # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # # #     except Exception as e:
# # # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # # #         return False

# # # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # # #     if job_name not in llm_response_dict: return False
# # # # # # #     main_data = llm_response_dict[job_name]

# # # # # # #     def format_list(key):
# # # # # # #         val = main_data.get(key, [])
# # # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # # #         return val if val else "NA"
    
# # # # # # #     template_data = {
# # # # # # #         "name": main_data.get("name", "NA"),
# # # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # # #         "certifications": format_list("certifications"),
# # # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # # #         "languages": format_list("languages"),
# # # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # # #         "database": format_list("database"),
# # # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # # #         "development_tools": format_list("development_tool"),
# # # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # # #     }
    
# # # # # # #     try:
# # # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # # #             template = Template(f.read())
# # # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # # #             f.write(template.render(**template_data))
# # # # # # #         return True
# # # # # # #     except Exception as e:
# # # # # # #         print(f"❌ Template Error: {e}")
# # # # # # #         return False


# # # # # # import google.generativeai as genai
# # # # # # from google.api_core import exceptions
# # # # # # from dotenv import load_dotenv
# # # # # # import os
# # # # # # import json
# # # # # # import time
# # # # # # from jinja2 import Template
# # # # # # import re
# # # # # # import docx
# # # # # # import traceback
# # # # # # from datetime import datetime

# # # # # # load_dotenv()

# # # # # # def configure_gemini():
# # # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # # #     if not api_key_gemini:
# # # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # # #     genai.configure(api_key=api_key_gemini)
# # # # # #     return genai

# # # # # # def clean_json_response(text):
# # # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # # #     return cleaned.strip()

# # # # # # def calculate_duration_string(duration_raw):
# # # # # #     """Parses a duration string (e.g., 'Oct 2023 - Present') and calculates months."""
# # # # # #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# # # # # #         return "N/A"

# # # # # #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
# # # # # #     if len(parts) != 2: return duration_raw

# # # # # #     start_str = parts[0].strip()
# # # # # #     end_str = parts[1].strip()
    
# # # # # #     def parse_date(d_str):
# # # # # #         if d_str.lower() in ["present", "current", "now", "ongoing"]: return datetime.now()
# # # # # #         formats = ["%b %Y", "%B %Y", "%b-%Y", "%B-%Y", "%m/%Y", "%Y"]
# # # # # #         for fmt in formats:
# # # # # #             try: return datetime.strptime(d_str, fmt)
# # # # # #             except ValueError: continue
# # # # # #         return None

# # # # # #     start_date = parse_date(start_str)
# # # # # #     end_date = parse_date(end_str)

# # # # # #     if start_date and end_date:
# # # # # #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
# # # # # #         if diff_months < 1: diff_months = 1
        
# # # # # #         if diff_months >= 12:
# # # # # #             years = diff_months // 12
# # # # # #             months = diff_months % 12
# # # # # #             duration_txt = f"{years} Years" + (f" {months} Months" if months > 0 else "")
# # # # # #         else:
# # # # # #             duration_txt = f"{diff_months} Months"
            
# # # # # #         return f"{start_str} - {end_str} ({duration_txt})"
# # # # # #     return duration_raw

# # # # # # def transform_projects_data(projects_list):
# # # # # #     """Transform projects data from LLM format to template format"""
# # # # # #     transformed_projects = []
# # # # # #     if not projects_list: return []
# # # # # #     for project in projects_list:
# # # # # #         tech_stack = project.get("tech_stack", [])
# # # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # # #         # === FIX NONE PROJECT NAME ===
# # # # # #         raw_name = project.get("project_name", "")
# # # # # #         role = project.get("Role", "")
# # # # # #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown"]:
# # # # # #             name = f"{role} Project" if role else "Professional Experience"
# # # # # #         else:
# # # # # #             name = raw_name

# # # # # #         clouds = project.get("Clouds", "")
# # # # # #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# # # # # #             formatted_title = f"{name} ({clouds})"
# # # # # #         else:
# # # # # #             formatted_title = name

# # # # # #         desc = project.get("description", [])
# # # # # #         if isinstance(desc, str):
# # # # # #             if '\n' in desc: desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # # #             else: desc = [desc]
        
# # # # # #         transformed_projects.append({
# # # # # #             "title": formatted_title,
# # # # # #             "role": role,
# # # # # #             "industry": project.get("Industry", "NA"),
# # # # # #             "duration": calculate_duration_string(project.get("Duration", "")),
# # # # # #             "link": project.get("link", ""),
# # # # # #             "description": desc,
# # # # # #             "tech_stack": tech_stack_str
# # # # # #         })
# # # # # #     return transformed_projects

# # # # # # def extract_text_from_docx(docx_path):
# # # # # #     doc = docx.Document(docx_path)
# # # # # #     full_text = []
# # # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # # #     for table in doc.tables:
# # # # # #         for row in table.rows:
# # # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # # #     return '\n'.join(full_text)

# # # # # # def get_improved_extraction_prompt():
# # # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # # #     CRITICAL RULES:
# # # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Projects' sections into the 'projects_descriptions' array.
# # # # # #     2. **Project Names**: If a specific project name is not listed, **USE THE COMPANY NAME**. Do NOT return 'None'.
# # # # # #     3. **List Format**: Break descriptions into a list of distinct strings. Remove bullet points.
# # # # # #     4. **Skill Mapping**: Map technical skills (Selenium, Python) to 'salesforce_expertise'. 
# # # # # #     5. **Soft Skills**: Only behavioral traits (Communication, Leadership) in 'soft_skills'. **NO TECHNICAL SKILLS (like Data Processing, APIs, Cloud) ALLOWED HERE.**
# # # # # #     6. **Industry Inference**: Infer the industry (e.g., 'Banking') from the description.
# # # # # #     7. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# # # # # #     8. Output ONLY valid JSON."""

# # # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # # {
# # # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # # #     "name": "Full Name",
# # # # # #     "current_job_role": "Current Job Title",
# # # # # #     "experience": "Total Experience",
    
# # # # # #     "crm_administration": ["Skill 1"],
# # # # # #     "certifications": ["Cert 1"],
# # # # # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # # # # #     "languages": ["Language 1"],
# # # # # #     "salesforce_components": ["Component 1"],
# # # # # #     "ticketing_case_management": ["Tool 1"],
# # # # # #     "database": ["DB 1"],
# # # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # # #     "soft_skills": ["Behavioral Skill 1"],

# # # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # # #     "development_tool": ["Tool 1"],
# # # # # #     "projects_descriptions": [
# # # # # #       {
# # # # # #         "project_name": "Company Name OR Project Title (Do NOT use 'None')",
# # # # # #         "Clouds": "Specific Cloud/Domain (e.g. Sales Cloud, AI)",
# # # # # #         "Role": "Job Role",
# # # # # #         "Industry": "Inferred Industry",
# # # # # #         "Duration": "Date Range",
# # # # # #         "description": ["Task 1...", "Result 2..."], 
# # # # # #         "tech_stack": ["Tech"],
# # # # # #         "link": ""
# # # # # #       }
# # # # # #     ]
# # # # # #   }
# # # # # # }"""
# # # # # #     return system_prompt, user_prompt

# # # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # # #     genai = configure_gemini()
# # # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # # #     response = None
# # # # # #     success = False

# # # # # #     for model_name in models_to_try:
# # # # # #         if success: break
# # # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # # #         try:
# # # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # # #         except Exception: continue
        
# # # # # #         for attempt in range(3):
# # # # # #             try:
# # # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # # #                 else:
# # # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # # #                 success = True
# # # # # #                 print(f"✅ Success with {model_name}!")
# # # # # #                 break
# # # # # #             except exceptions.ResourceExhausted:
# # # # # #                 time.sleep(10)
# # # # # #             except exceptions.NotFound:
# # # # # #                 break 
# # # # # #             except Exception as e:
# # # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # # #                 break

# # # # # #     if not success or not response:
# # # # # #         print("❌ All models failed.")
# # # # # #         return False

# # # # # #     try:
# # # # # #         cleaned_response = clean_json_response(response.text)
# # # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # # #     except Exception as e:
# # # # # #         print(f"❌ Parsing Error: {e}")
# # # # # #         return False

# # # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # # #     if job_name not in llm_response_dict: return False
# # # # # #     main_data = llm_response_dict[job_name]

# # # # # #     def format_list(key):
# # # # # #         val = main_data.get(key, [])
# # # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # # #         return val if val else "NA"
    
# # # # # #     template_data = {
# # # # # #         "name": main_data.get("name", "NA"),
# # # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # # #         "experience": main_data.get("experience", "NA"),
# # # # # #         "crm_administration": format_list("crm_administration"),
# # # # # #         "certifications": format_list("certifications"),
# # # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # # #         "languages": format_list("languages"),
# # # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # # #         "database": format_list("database"),
# # # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # # #         "soft_skills": format_list("soft_skills"),
# # # # # #         "development_tools": format_list("development_tool"),
# # # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # # #     }
    
# # # # # #     try:
# # # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # # #             template = Template(f.read())
# # # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # # #             f.write(template.render(**template_data))
# # # # # #         return True
# # # # # #     except Exception as e:
# # # # # #         print(f"❌ Template Error: {e}")
# # # # # #         return False

# # # # # import google.generativeai as genai
# # # # # from google.api_core import exceptions
# # # # # from dotenv import load_dotenv
# # # # # import os
# # # # # import json
# # # # # import time
# # # # # from jinja2 import Template
# # # # # import re
# # # # # import docx
# # # # # import traceback
# # # # # from datetime import datetime

# # # # # load_dotenv()

# # # # # def configure_gemini():
# # # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # # #     if not api_key_gemini:
# # # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # # #     genai.configure(api_key=api_key_gemini)
# # # # #     return genai

# # # # # def clean_json_response(text):
# # # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # # #     return cleaned.strip()

# # # # # def calculate_duration_string(duration_raw):
# # # # #     """Parses a duration string and calculates months."""
# # # # #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# # # # #         return "N/A"

# # # # #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
# # # # #     if len(parts) != 2: return duration_raw

# # # # #     start_str = parts[0].strip()
# # # # #     end_str = parts[1].strip()
    
# # # # #     def parse_date(d_str):
# # # # #         if d_str.lower() in ["present", "current", "now", "ongoing"]: return datetime.now()
# # # # #         formats = ["%b %Y", "%B %Y", "%b-%Y", "%B-%Y", "%m/%Y", "%Y"]
# # # # #         for fmt in formats:
# # # # #             try: return datetime.strptime(d_str, fmt)
# # # # #             except ValueError: continue
# # # # #         return None

# # # # #     start_date = parse_date(start_str)
# # # # #     end_date = parse_date(end_str)

# # # # #     if start_date and end_date:
# # # # #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
# # # # #         if diff_months < 1: diff_months = 1
        
# # # # #         if diff_months >= 12:
# # # # #             years = diff_months // 12
# # # # #             months = diff_months % 12
# # # # #             duration_txt = f"{years} Years" + (f" {months} Months" if months > 0 else "")
# # # # #         else:
# # # # #             duration_txt = f"{diff_months} Months"
            
# # # # #         return f"{start_str} - {end_str} ({duration_txt})"
# # # # #     return duration_raw

# # # # # def transform_projects_data(projects_list):
# # # # #     """Transform projects data from LLM format to template format"""
# # # # #     transformed_projects = []
# # # # #     if not projects_list: return []
# # # # #     for project in projects_list:
# # # # #         tech_stack = project.get("tech_stack", [])
# # # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # # #         # === FIX NONE PROJECT NAME ===
# # # # #         raw_name = project.get("project_name", "")
# # # # #         role = project.get("Role", "")
# # # # #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown"]:
# # # # #             name = f"{role} Project" if role else "Professional Experience"
# # # # #         else:
# # # # #             name = raw_name

# # # # #         clouds = project.get("Clouds", "")
# # # # #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# # # # #             formatted_title = f"{name} ({clouds})"
# # # # #         else:
# # # # #             formatted_title = name

# # # # #         desc = project.get("description", [])
# # # # #         if isinstance(desc, str):
# # # # #             if '\n' in desc: desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # # #             else: desc = [desc]
        
# # # # #         transformed_projects.append({
# # # # #             "title": formatted_title,
# # # # #             "role": role,
# # # # #             "industry": project.get("Industry", "NA"),
# # # # #             "duration": calculate_duration_string(project.get("Duration", "")),
# # # # #             "link": project.get("link", ""),
# # # # #             "description": desc,
# # # # #             "tech_stack": tech_stack_str
# # # # #         })
# # # # #     return transformed_projects

# # # # # def extract_text_from_docx(docx_path):
# # # # #     doc = docx.Document(docx_path)
# # # # #     full_text = []
# # # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # # #     for table in doc.tables:
# # # # #         for row in table.rows:
# # # # #             for cell in row.cells: full_text.append(cell.text)
# # # # #     return '\n'.join(full_text)

# # # # # def get_improved_extraction_prompt():
# # # # #     system_prompt = """You are an expert resume parser AI. 
# # # # #     CRITICAL RULES:
# # # # #     1. **Experience = Projects**: You MUST extract the candidate's 'Work Experience' or 'Projects' sections into the 'projects_descriptions' array.
# # # # #     2. **Project Names**: If a specific project name is not listed, **USE THE COMPANY NAME**. Do NOT return 'None'.
# # # # #     3. **Exact Content (Verbatim)**: For 'description', extract the text EXACTLY as it appears in the resume. Do NOT summarize, rephrase, improve, or shorten it. Maintain the original meaning and detail.
# # # # #     4. **List Format**: Break descriptions into a list of distinct strings. Remove bullet points.
# # # # #     5. **Skill Mapping**: Map technical skills (Selenium, Python) to 'salesforce_expertise'. 
# # # # #     6. **Soft Skills**: Only behavioral traits (Communication, Leadership). **NO TECHNICAL SKILLS ALLOWED HERE.**
# # # # #     7. **Industry Inference**: Infer the industry (e.g., 'Banking') from the description.
# # # # #     8. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# # # # #     9. Output ONLY valid JSON."""

# # # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # # {
# # # # #   "Extract and Synthesize Candidate's Resume": {
# # # # #     "name": "Full Name",
# # # # #     "current_job_role": "Current Job Title",
# # # # #     "experience": "Total Experience",
    
# # # # #     "crm_administration": ["Skill 1"],
# # # # #     "certifications": ["Cert 1"],
# # # # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # # # #     "languages": ["Language 1"],
# # # # #     "salesforce_components": ["Component 1"],
# # # # #     "ticketing_case_management": ["Tool 1"],
# # # # #     "database": ["DB 1"],
# # # # #     "salesforce_clouds": ["Cloud 1"],
# # # # #     "soft_skills": ["Behavioral Skill 1"],

# # # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # # #     "development_tool": ["Tool 1"],
# # # # #     "projects_descriptions": [
# # # # #       {
# # # # #         "project_name": "Company Name OR Project Title (Do NOT use 'None')",
# # # # #         "Clouds": "Specific Cloud/Domain (e.g. Sales Cloud, AI)",
# # # # #         "Role": "Job Role",
# # # # #         "Industry": "Inferred Industry",
# # # # #         "Duration": "Date Range",
# # # # #         "description": ["Exact line 1 from resume...", "Exact line 2 from resume..."], 
# # # # #         "tech_stack": ["Tech"],
# # # # #         "link": ""
# # # # #       }
# # # # #     ]
# # # # #   }
# # # # # }"""
# # # # #     return system_prompt, user_prompt

# # # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # # #     genai = configure_gemini()
# # # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # # #     response = None
# # # # #     success = False

# # # # #     for model_name in models_to_try:
# # # # #         if success: break
# # # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # # #         try:
# # # # #             model = genai.GenerativeModel(model_name=model_name)
# # # # #         except Exception: continue
        
# # # # #         for attempt in range(3):
# # # # #             try:
# # # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # # #                 else:
# # # # #                     with open(pdf_file_path, 'rb') as f:
# # # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # # #                 success = True
# # # # #                 print(f"✅ Success with {model_name}!")
# # # # #                 break
# # # # #             except exceptions.ResourceExhausted:
# # # # #                 time.sleep(10)
# # # # #             except exceptions.NotFound:
# # # # #                 break 
# # # # #             except Exception as e:
# # # # #                 print(f"❌ Error with {model_name}: {e}")
# # # # #                 break

# # # # #     if not success or not response:
# # # # #         print("❌ All models failed.")
# # # # #         return False

# # # # #     try:
# # # # #         cleaned_response = clean_json_response(response.text)
# # # # #         parsed_output = json.loads(cleaned_response)
        
# # # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # # #     except Exception as e:
# # # # #         print(f"❌ Parsing Error: {e}")
# # # # #         return False

# # # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # # #     if job_name not in llm_response_dict: return False
# # # # #     main_data = llm_response_dict[job_name]

# # # # #     def format_list(key):
# # # # #         val = main_data.get(key, [])
# # # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # # #         return val if val else "NA"
    
# # # # #     template_data = {
# # # # #         "name": main_data.get("name", "NA"),
# # # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # # #         "experience": main_data.get("experience", "NA"),
# # # # #         "crm_administration": format_list("crm_administration"),
# # # # #         "certifications": format_list("certifications"),
# # # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # # #         "languages": format_list("languages"),
# # # # #         "salesforce_components": format_list("salesforce_components"),
# # # # #         "ticketing": format_list("ticketing_case_management"),
# # # # #         "database": format_list("database"),
# # # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # # #         "soft_skills": format_list("soft_skills"),
# # # # #         "development_tools": format_list("development_tool"),
# # # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # # #     }
    
# # # # #     try:
# # # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # # #             template = Template(f.read())
# # # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # # #             f.write(template.render(**template_data))
# # # # #         return True
# # # # #     except Exception as e:
# # # # #         print(f"❌ Template Error: {e}")
# # # # #         return False

# # # # import google.generativeai as genai
# # # # from google.api_core import exceptions
# # # # from dotenv import load_dotenv
# # # # import os
# # # # import json
# # # # import time
# # # # from jinja2 import Template
# # # # import re
# # # # import docx
# # # # import traceback
# # # # from datetime import datetime

# # # # load_dotenv()

# # # # def configure_gemini():
# # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # #     if not api_key_gemini:
# # # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # # #     genai.configure(api_key=api_key_gemini)
# # # #     return genai

# # # # def clean_json_response(text):
# # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # #     return cleaned.strip()

# # # # def calculate_duration_string(duration_raw):
# # # #     """Parses a duration string and calculates months."""
# # # #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# # # #         return "N/A"

# # # #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
# # # #     if len(parts) != 2: return duration_raw

# # # #     start_str = parts[0].strip()
# # # #     end_str = parts[1].strip()
    
# # # #     def parse_date(d_str):
# # # #         if d_str.lower() in ["present", "current", "now", "ongoing"]: return datetime.now()
# # # #         formats = ["%b %Y", "%B %Y", "%b-%Y", "%B-%Y", "%m/%Y", "%Y"]
# # # #         for fmt in formats:
# # # #             try: return datetime.strptime(d_str, fmt)
# # # #             except ValueError: continue
# # # #         return None

# # # #     start_date = parse_date(start_str)
# # # #     end_date = parse_date(end_str)

# # # #     if start_date and end_date:
# # # #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
# # # #         if diff_months < 1: diff_months = 1
        
# # # #         if diff_months >= 12:
# # # #             years = diff_months // 12
# # # #             months = diff_months % 12
# # # #             duration_txt = f"{years} Years" + (f" {months} Months" if months > 0 else "")
# # # #         else:
# # # #             duration_txt = f"{diff_months} Months"
            
# # # #         return f"{start_str} - {end_str} ({duration_txt})"
# # # #     return duration_raw

# # # # def transform_projects_data(projects_list):
# # # #     """Transform projects data from LLM format to template format"""
# # # #     transformed_projects = []
# # # #     if not projects_list: return []
# # # #     for project in projects_list:
# # # #         tech_stack = project.get("tech_stack", [])
# # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # #         # === FIX NONE PROJECT NAME ===
# # # #         raw_name = project.get("project_name", "")
# # # #         role = project.get("Role", "")
# # # #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown"]:
# # # #             name = f"{role} Project" if role else "Professional Experience"
# # # #         else:
# # # #             name = raw_name

# # # #         # === CLOUDS HANDLING ===
# # # #         clouds = project.get("Clouds", "")
# # # #         # Only add parentheses if we have a valid cloud/domain
# # # #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# # # #             formatted_title = f"{name} ({clouds})"
# # # #         else:
# # # #             formatted_title = name

# # # #         desc = project.get("description", [])
# # # #         if isinstance(desc, str):
# # # #             if '\n' in desc: desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # # #             else: desc = [desc]
        
# # # #         transformed_projects.append({
# # # #             "title": formatted_title,
# # # #             "role": role,
# # # #             "industry": project.get("Industry", "NA"),
# # # #             "duration": calculate_duration_string(project.get("Duration", "")),
# # # #             "link": project.get("link", ""),
# # # #             "description": desc,
# # # #             "tech_stack": tech_stack_str
# # # #         })
# # # #     return transformed_projects

# # # # def extract_text_from_docx(docx_path):
# # # #     doc = docx.Document(docx_path)
# # # #     full_text = []
# # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # #     for table in doc.tables:
# # # #         for row in table.rows:
# # # #             for cell in row.cells: full_text.append(cell.text)
# # # #     return '\n'.join(full_text)

# # # # def get_improved_extraction_prompt():
# # # #     system_prompt = """You are an expert resume parser AI. 
# # # #     CRITICAL RULES:
# # # #     1. **Experience = Projects**: Extract 'Work Experience' or 'Projects' into 'projects_descriptions'.
# # # #     2. **Project Names**: If a project name is missing, **USE THE COMPANY NAME**. Do NOT return 'None'.
# # # #     3. **Exact Content (Verbatim)**: Extract 'description' text EXACTLY as written in the resume. Do NOT summarize or rephrase.
# # # #     4. **Infer Clouds/Domain**: You **MUST** identify the Cloud or Technical Domain for EVERY project. If not explicitly stated, INFER it from the tech stack (e.g., 'Web App', 'Mobile', 'AI', 'FinTech', 'Sales Cloud'). **Do NOT return 'NA' or 'None'.**
# # # #     5. **Main Table Clouds**: For the 'salesforce_clouds' field in the top table, list ANY Salesforce Cloud mentioned in the projects (e.g. Sales, Service). If none, infer the domain (e.g. Web, QA).
# # # #     6. **Skill Mapping**: Technical skills go to 'salesforce_expertise'.
# # # #     7. **Soft Skills**: Only behavioral traits (Communication, Leadership). NO technical skills.
# # # #     8. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# # # #     9. Output ONLY valid JSON."""

# # # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # # {
# # # #   "Extract and Synthesize Candidate's Resume": {
# # # #     "name": "Full Name",
# # # #     "current_job_role": "Current Job Title",
# # # #     "experience": "Total Experience",
    
# # # #     "crm_administration": ["Skill 1"],
# # # #     "certifications": ["Cert 1"],
# # # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # # #     "languages": ["Language 1"],
# # # #     "salesforce_components": ["Component 1"],
# # # #     "ticketing_case_management": ["Tool 1"],
# # # #     "database": ["DB 1"],
# # # #     "salesforce_clouds": ["Sales Cloud", "Service Cloud", "OR 'NA'"],
# # # #     "soft_skills": ["Behavioral Skill 1"],

# # # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # # #     "development_tool": ["Tool 1"],
# # # #     "projects_descriptions": [
# # # #       {
# # # #         "project_name": "Company Name OR Project Title",
# # # #         "Clouds": "REQUIRED: Inferred Cloud or Domain (e.g. 'Web', 'Mobile', 'AI', 'Sales Cloud'). NO 'NA'.",
# # # #         "Role": "Job Role",
# # # #         "Industry": "Inferred Industry",
# # # #         "Duration": "Date Range",
# # # #         "description": ["Exact line 1...", "Exact line 2..."], 
# # # #         "tech_stack": ["Tech"],
# # # #         "link": ""
# # # #       }
# # # #     ]
# # # #   }
# # # # }"""
# # # #     return system_prompt, user_prompt

# # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # #     genai = configure_gemini()
# # # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # # #     response = None
# # # #     success = False

# # # #     for model_name in models_to_try:
# # # #         if success: break
# # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # #         try:
# # # #             model = genai.GenerativeModel(model_name=model_name)
# # # #         except Exception: continue
        
# # # #         for attempt in range(3):
# # # #             try:
# # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # #                 else:
# # # #                     with open(pdf_file_path, 'rb') as f:
# # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # #                 success = True
# # # #                 print(f"✅ Success with {model_name}!")
# # # #                 break
# # # #             except exceptions.ResourceExhausted:
# # # #                 time.sleep(10)
# # # #             except exceptions.NotFound:
# # # #                 break 
# # # #             except Exception as e:
# # # #                 print(f"❌ Error with {model_name}: {e}")
# # # #                 break

# # # #     if not success or not response:
# # # #         print("❌ All models failed.")
# # # #         return False

# # # #     try:
# # # #         cleaned_response = clean_json_response(response.text)
# # # #         parsed_output = json.loads(cleaned_response)
        
# # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # # #     except Exception as e:
# # # #         print(f"❌ Parsing Error: {e}")
# # # #         return False

# # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # #     if job_name not in llm_response_dict: return False
# # # #     main_data = llm_response_dict[job_name]

# # # #     def format_list(key):
# # # #         val = main_data.get(key, [])
# # # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # # #         return val if val else "NA"
    
# # # #     template_data = {
# # # #         "name": main_data.get("name", "NA"),
# # # #         "job_role": main_data.get("current_job_role", "NA"),
# # # #         "experience": main_data.get("experience", "NA"),
# # # #         "crm_administration": format_list("crm_administration"),
# # # #         "certifications": format_list("certifications"),
# # # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # # #         "languages": format_list("languages"),
# # # #         "salesforce_components": format_list("salesforce_components"),
# # # #         "ticketing": format_list("ticketing_case_management"),
# # # #         "database": format_list("database"),
# # # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # # #         "soft_skills": format_list("soft_skills"),
# # # #         "development_tools": format_list("development_tool"),
# # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # #     }
    
# # # #     try:
# # # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # # #             template = Template(f.read())
# # # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # # #             f.write(template.render(**template_data))
# # # #         return True
# # # #     except Exception as e:
# # # #         print(f"❌ Template Error: {e}")
# # # #         return False


# # # import google.generativeai as genai
# # # from google.api_core import exceptions
# # # from dotenv import load_dotenv
# # # import os
# # # import json
# # # import time
# # # from jinja2 import Template
# # # import re
# # # import docx
# # # import traceback
# # # from datetime import datetime

# # # load_dotenv()

# # # def configure_gemini():
# # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # #     if not api_key_gemini:
# # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# # #     genai.configure(api_key=api_key_gemini)
# # #     return genai

# # # def clean_json_response(text):
# # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # #     return cleaned.strip()

# # # def calculate_duration_string(duration_raw):
# # #     """Parses a duration string and calculates months."""
# # #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# # #         return "N/A"

# # #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
# # #     if len(parts) != 2: return duration_raw

# # #     start_str = parts[0].strip()
# # #     end_str = parts[1].strip()
    
# # #     def parse_date(d_str):
# # #         if d_str.lower() in ["present", "current", "now", "ongoing"]: return datetime.now()
# # #         formats = ["%b %Y", "%B %Y", "%b-%Y", "%B-%Y", "%m/%Y", "%Y"]
# # #         for fmt in formats:
# # #             try: return datetime.strptime(d_str, fmt)
# # #             except ValueError: continue
# # #         return None

# # #     start_date = parse_date(start_str)
# # #     end_date = parse_date(end_str)

# # #     if start_date and end_date:
# # #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
# # #         if diff_months < 1: diff_months = 1
        
# # #         if diff_months >= 12:
# # #             years = diff_months // 12
# # #             months = diff_months % 12
# # #             duration_txt = f"{years} Years" + (f" {months} Months" if months > 0 else "")
# # #         else:
# # #             duration_txt = f"{diff_months} Months"
            
# # #         return f"{start_str} - {end_str} ({duration_txt})"
# # #     return duration_raw

# # # def transform_projects_data(projects_list):
# # #     """Transform projects data from LLM format to template format"""
# # #     transformed_projects = []
# # #     if not projects_list: return []
# # #     for project in projects_list:
# # #         tech_stack = project.get("tech_stack", [])
# # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # #         # === FIX NONE PROJECT NAME ===
# # #         raw_name = project.get("project_name", "")
# # #         role = project.get("Role", "")
# # #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown"]:
# # #             name = f"{role} Project" if role else "Professional Experience"
# # #         else:
# # #             name = raw_name

# # #         # === CLOUDS HANDLING ===
# # #         clouds = project.get("Clouds", "")
# # #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# # #             formatted_title = f"{name} ({clouds})"
# # #         else:
# # #             formatted_title = name

# # #         desc = project.get("description", [])
# # #         if isinstance(desc, str):
# # #             if '\n' in desc: desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # #             else: desc = [desc]
        
# # #         transformed_projects.append({
# # #             "title": formatted_title,
# # #             "role": role,
# # #             "industry": project.get("Industry", "NA"),
# # #             "duration": calculate_duration_string(project.get("Duration", "")),
# # #             "link": project.get("link", ""),
# # #             "description": desc,
# # #             "tech_stack": tech_stack_str
# # #         })
# # #     return transformed_projects

# # # def extract_text_from_docx(docx_path):
# # #     doc = docx.Document(docx_path)
# # #     full_text = []
# # #     for para in doc.paragraphs: full_text.append(para.text)
# # #     for table in doc.tables:
# # #         for row in table.rows:
# # #             for cell in row.cells: full_text.append(cell.text)
# # #     return '\n'.join(full_text)

# # # def get_improved_extraction_prompt():
# # #     system_prompt = """You are an expert resume parser AI. 
# # #     CRITICAL RULES:
# # #     1. **Name Extraction**: Extract the candidate's ACTUAL NAME. Do NOT return 'Full Name', 'Candidate Name', or 'Name'. If not found, look for the text at the very top of the document.
# # #     2. **Experience = Projects**: Extract 'Work Experience' or 'Projects' into 'projects_descriptions'.
# # #     3. **Project Names**: If missing, **USE COMPANY NAME**. Do NOT return 'None'.
# # #     4. **Exact Content**: Extract 'description' text EXACTLY as written.
# # #     5. **Infer Clouds**: INFER Cloud/Domain (e.g. 'Web', 'Mobile', 'AI'). Do NOT return 'NA'.
# # #     6. **Main Table Clouds**: List Salesforce Clouds mentioned in projects (Sales, Service, etc.) or 'NA'.
# # #     7. **Skill Mapping**: Technical skills -> 'salesforce_expertise'. Soft skills -> Behavioral only.
# # #     8. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# # #     9. Output ONLY valid JSON."""

# # #     # CHANGED: Replaced "Full Name" example with "Actual Name" to prevent literal copying
# # #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # # {
# # #   "Extract and Synthesize Candidate's Resume": {
# # #     "name": "Actual Candidate Name",
# # #     "current_job_role": "Current Job Title",
# # #     "experience": "Total Experience",
    
# # #     "crm_administration": ["Skill 1"],
# # #     "certifications": ["Cert 1"],
# # #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# # #     "languages": ["Language 1"],
# # #     "salesforce_components": ["Component 1"],
# # #     "ticketing_case_management": ["Tool 1"],
# # #     "database": ["DB 1"],
# # #     "salesforce_clouds": ["Sales Cloud", "Service Cloud", "OR 'NA'"],
# # #     "soft_skills": ["Behavioral Skill 1"],

# # #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# # #     "development_tool": ["Tool 1"],
# # #     "projects_descriptions": [
# # #       {
# # #         "project_name": "Company Name OR Project Title",
# # #         "Clouds": "REQUIRED: Inferred Cloud or Domain (e.g. 'Web', 'Mobile'). NO 'NA'.",
# # #         "Role": "Job Role",
# # #         "Industry": "Inferred Industry",
# # #         "Duration": "Date Range",
# # #         "description": ["Exact line 1...", "Exact line 2..."], 
# # #         "tech_stack": ["Tech"],
# # #         "link": ""
# # #       }
# # #     ]
# # #   }
# # # }"""
# # #     return system_prompt, user_prompt

# # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # #     genai = configure_gemini()
# # #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # #     response = None
# # #     success = False

# # #     for model_name in models_to_try:
# # #         if success: break
# # #         print(f"🔄 Attempting with model: {model_name}...")
# # #         try:
# # #             model = genai.GenerativeModel(model_name=model_name)
# # #         except Exception: continue
        
# # #         for attempt in range(3):
# # #             try:
# # #                 if pdf_file_path.lower().endswith('.docx'):
# # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # #                 else:
# # #                     with open(pdf_file_path, 'rb') as f:
# # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # #                 success = True
# # #                 print(f"✅ Success with {model_name}!")
# # #                 break
# # #             except exceptions.ResourceExhausted:
# # #                 time.sleep(10)
# # #             except exceptions.NotFound:
# # #                 break 
# # #             except Exception as e:
# # #                 print(f"❌ Error with {model_name}: {e}")
# # #                 break

# # #     if not success or not response:
# # #         print("❌ All models failed.")
# # #         return False

# # #     try:
# # #         cleaned_response = clean_json_response(response.text)
# # #         parsed_output = json.loads(cleaned_response)
        
# # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# # #     except Exception as e:
# # #         print(f"❌ Parsing Error: {e}")
# # #         return False

# # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # #     job_name = "Extract and Synthesize Candidate's Resume"
# # #     if job_name not in llm_response_dict: return False
# # #     main_data = llm_response_dict[job_name]

# # #     def format_list(key):
# # #         val = main_data.get(key, [])
# # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # #         return val if val else "NA"
    
# # #     # === FIX NAME ===
# # #     raw_name = main_data.get("name", "NA")
# # #     if raw_name in ["Full Name", "Candidate Name", "Actual Candidate Name"]:
# # #         raw_name = "NA" # Or you could set it to "" to leave it blank

# # #     template_data = {
# # #         "name": raw_name,
# # #         "job_role": main_data.get("current_job_role", "NA"),
# # #         "experience": main_data.get("experience", "NA"),
# # #         "crm_administration": format_list("crm_administration"),
# # #         "certifications": format_list("certifications"),
# # #         "salesforce_expertise": format_list("salesforce_expertise"),
# # #         "languages": format_list("languages"),
# # #         "salesforce_components": format_list("salesforce_components"),
# # #         "ticketing": format_list("ticketing_case_management"),
# # #         "database": format_list("database"),
# # #         "salesforce_clouds": format_list("salesforce_clouds"),
# # #         "soft_skills": format_list("soft_skills"),
# # #         "development_tools": format_list("development_tool"),
# # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # #     }
    
# # #     try:
# # #         with open(template_file_path, 'r', encoding='utf-8') as f:
# # #             template = Template(f.read())
# # #         with open(output_file_path, 'w', encoding='utf-8') as f:
# # #             f.write(template.render(**template_data))
# # #         return True
# # #     except Exception as e:
# # #         print(f"❌ Template Error: {e}")
# # #         return False

# # import google.generativeai as genai
# # from google.api_core import exceptions
# # from dotenv import load_dotenv
# # import os
# # import json
# # import time
# # from jinja2 import Template
# # import re
# # import docx
# # import traceback
# # from datetime import datetime

# # load_dotenv()

# # def configure_gemini():
# #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# #     if not api_key_gemini:
# #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
# #     genai.configure(api_key=api_key_gemini)
# #     return genai

# # def clean_json_response(text):
# #     """Remove markdown code blocks and extra whitespace from JSON response"""
# #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# #     return cleaned.strip()

# # def calculate_duration_string(duration_raw):
# #     """
# #     Parses a duration string (e.g., "Oct 2023 - Present") and returns ONLY the total months.
# #     Output format: "29 Months"
# #     """
# #     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
# #         return "N/A"

# #     # Regex to find dates like "Oct 2023", "October 2023", "2023", etc.
# #     parts = re.split(r'\s*-\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)
    
# #     if len(parts) != 2:
# #         return duration_raw

# #     start_str = parts[0].strip()
# #     end_str = parts[1].strip()
    
# #     def parse_date(d_str):
# #         if d_str.lower() in ["present", "current", "now", "ongoing"]:
# #             return datetime.now()
        
# #         formats = [
# #             "%b %Y", "%B %Y",       # Oct 2023, October 2023
# #             "%b-%Y", "%B-%Y",       # Oct-2023
# #             "%m/%Y", "%Y",          # 10/2023, 2023
# #             "%b %d, %Y",            # Oct 1, 2023
# #             "%B %d, %Y"             # October 1, 2023
# #         ]
        
# #         for fmt in formats:
# #             try:
# #                 return datetime.strptime(d_str, fmt)
# #             except ValueError:
# #                 continue
# #         return None

# #     start_date = parse_date(start_str)
# #     end_date = parse_date(end_str)

# #     if start_date and end_date:
# #         # Calculate difference in months
# #         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
        
# #         if diff_months < 1: diff_months = 1
        
# #         # CHANGED: Return ONLY the total months count
# #         return f"{diff_months} Months"
    
# #     return duration_raw

# # def transform_projects_data(projects_list):
# #     """Transform projects data from LLM format to template format"""
# #     transformed_projects = []
# #     if not projects_list: return []
# #     for project in projects_list:
# #         tech_stack = project.get("tech_stack", [])
# #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# #         # === FIX NONE PROJECT NAME ===
# #         raw_name = project.get("project_name", "")
# #         role = project.get("Role", "")
# #         if not raw_name or raw_name.strip().lower() in ["none", "na", "n/a", "unknown"]:
# #             name = f"{role} Project" if role else "Professional Experience"
# #         else:
# #             name = raw_name

# #         # === CLOUDS HANDLING ===
# #         clouds = project.get("Clouds", "")
# #         if clouds and clouds not in ["NA", "None", "", "N/A"]:
# #             formatted_title = f"{name} ({clouds})"
# #         else:
# #             formatted_title = name

# #         desc = project.get("description", [])
# #         if isinstance(desc, str):
# #             if '\n' in desc: desc = [line.strip() for line in desc.split('\n') if line.strip()]
# #             else: desc = [desc]
        
# #         transformed_projects.append({
# #             "title": formatted_title,
# #             "role": role,
# #             "industry": project.get("Industry", "NA"),
# #             "duration": calculate_duration_string(project.get("Duration", "")),
# #             "link": project.get("link", ""),
# #             "description": desc,
# #             "tech_stack": tech_stack_str
# #         })
# #     return transformed_projects

# # def extract_text_from_docx(docx_path):
# #     doc = docx.Document(docx_path)
# #     full_text = []
# #     for para in doc.paragraphs: full_text.append(para.text)
# #     for table in doc.tables:
# #         for row in table.rows:
# #             for cell in row.cells: full_text.append(cell.text)
# #     return '\n'.join(full_text)

# # def get_improved_extraction_prompt():
# #     system_prompt = """You are an expert resume parser AI. 
# #     CRITICAL RULES:
# #     1. **Name Extraction**: Extract the candidate's ACTUAL NAME. Do NOT return 'Full Name', 'Candidate Name', or 'Name'.
# #     2. **Experience = Projects**: Extract 'Work Experience' or 'Projects' into 'projects_descriptions'.
# #     3. **Project Names**: If missing, **USE COMPANY NAME**. Do NOT return 'None'.
# #     4. **Exact Content**: Extract 'description' text EXACTLY as written.
# #     5. **Infer Clouds**: INFER Cloud/Domain (e.g. 'Web', 'Mobile', 'AI'). Do NOT return 'NA'.
# #     6. **Main Table Clouds**: List Salesforce Clouds mentioned in projects (Sales, Service, etc.) or 'NA'.
# #     7. **Skill Mapping**: Technical skills -> 'salesforce_expertise'. Soft skills -> Behavioral only.
# #     8. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
# #     9. Output ONLY valid JSON."""

# #     user_prompt = """Extract the resume content into this EXACT JSON structure.

# # {
# #   "Extract and Synthesize Candidate's Resume": {
# #     "name": "Actual Candidate Name",
# #     "current_job_role": "Current Job Title",
# #     "experience": "Total Experience",
    
# #     "crm_administration": ["Skill 1"],
# #     "certifications": ["Cert 1"],
# #     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
# #     "languages": ["Language 1"],
# #     "salesforce_components": ["Component 1"],
# #     "ticketing_case_management": ["Tool 1"],
# #     "database": ["DB 1"],
# #     "salesforce_clouds": ["Sales Cloud", "Service Cloud", "OR 'NA'"],
# #     "soft_skills": ["Behavioral Skill 1"],

# #     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
# #     "development_tool": ["Tool 1"],
# #     "projects_descriptions": [
# #       {
# #         "project_name": "Company Name OR Project Title",
# #         "Clouds": "REQUIRED: Inferred Cloud or Domain (e.g. 'Web', 'Mobile'). NO 'NA'.",
# #         "Role": "Job Role",
# #         "Industry": "Inferred Industry",
# #         "Duration": "Date Range",
# #         "description": ["Exact line 1...", "Exact line 2..."], 
# #         "tech_stack": ["Tech"],
# #         "link": ""
# #       }
# #     ]
# #   }
# # }"""
# #     return system_prompt, user_prompt

# # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# #     genai = configure_gemini()
# #     print("\n🚀 Starting Salesforce Resume Extraction...")
    
# #     system_prompt, user_prompt = get_improved_extraction_prompt()
# #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# #     response = None
# #     success = False

# #     for model_name in models_to_try:
# #         if success: break
# #         print(f"🔄 Attempting with model: {model_name}...")
# #         try:
# #             model = genai.GenerativeModel(model_name=model_name)
# #         except Exception: continue
        
# #         for attempt in range(3):
# #             try:
# #                 if pdf_file_path.lower().endswith('.docx'):
# #                     resume_text = extract_text_from_docx(pdf_file_path)
# #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# #                 else:
# #                     with open(pdf_file_path, 'rb') as f:
# #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# #                 success = True
# #                 print(f"✅ Success with {model_name}!")
# #                 break
# #             except exceptions.ResourceExhausted:
# #                 time.sleep(10)
# #             except exceptions.NotFound:
# #                 break 
# #             except Exception as e:
# #                 print(f"❌ Error with {model_name}: {e}")
# #                 break

# #     if not success or not response:
# #         print("❌ All models failed.")
# #         return False

# #     try:
# #         cleaned_response = clean_json_response(response.text)
# #         parsed_output = json.loads(cleaned_response)
        
# #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
# #     except Exception as e:
# #         print(f"❌ Parsing Error: {e}")
# #         return False

# # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# #     job_name = "Extract and Synthesize Candidate's Resume"
# #     if job_name not in llm_response_dict: return False
# #     main_data = llm_response_dict[job_name]

# #     def format_list(key):
# #         val = main_data.get(key, [])
# #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# #         return val if val else "NA"
    
# #     # === FIX NAME ===
# #     raw_name = main_data.get("name", "NA")
# #     if raw_name in ["Full Name", "Candidate Name", "Actual Candidate Name"]:
# #         raw_name = "NA" 

# #     template_data = {
# #         "name": raw_name,
# #         "job_role": main_data.get("current_job_role", "NA"),
# #         "experience": main_data.get("experience", "NA"),
# #         "crm_administration": format_list("crm_administration"),
# #         "certifications": format_list("certifications"),
# #         "salesforce_expertise": format_list("salesforce_expertise"),
# #         "languages": format_list("languages"),
# #         "salesforce_components": format_list("salesforce_components"),
# #         "ticketing": format_list("ticketing_case_management"),
# #         "database": format_list("database"),
# #         "salesforce_clouds": format_list("salesforce_clouds"),
# #         "soft_skills": format_list("soft_skills"),
# #         "development_tools": format_list("development_tool"),
# #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# #     }
    
# #     try:
# #         with open(template_file_path, 'r', encoding='utf-8') as f:
# #             template = Template(f.read())
# #         with open(output_file_path, 'w', encoding='utf-8') as f:
# #             f.write(template.render(**template_data))
# #         return True
# #     except Exception as e:
# #         print(f"❌ Template Error: {e}")
# #         return False

# import google.generativeai as genai
# from google.api_core import exceptions
# from dotenv import load_dotenv
# import os
# import json
# import time
# from jinja2 import Template
# import re
# import docx
# import traceback
# from datetime import datetime

# load_dotenv()

# def configure_gemini():
#     api_key_gemini = os.getenv("GEMINI_API_KEY")
#     if not api_key_gemini:
#         raise ValueError("GEMINI_API_KEY not found in environment variables.")
#     genai.configure(api_key=api_key_gemini)
#     return genai

# def clean_json_response(text):
#     """Remove markdown code blocks and extra whitespace from JSON response"""
#     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
#     return cleaned.strip()

# def calculate_duration_string(duration_raw):
#     """
#     Parses a duration string (e.g., "Oct 2023 - Present") and returns ONLY the total months.
#     Output format: "29 Months"
#     """
#     if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
#         return "N/A"

#     # Regex to find dates like "Oct 2023", "October 2023", "2023", etc.
#     parts = re.split(r'\s*[-–]\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)

    
#     if len(parts) != 2:
#         return duration_raw

#     start_str = parts[0].strip()
#     end_str = parts[1].strip()
    
#     def parse_date(d_str):
#         if d_str.lower() in ["present", "current", "now", "ongoing"]:
#             return datetime.now()
        
#         formats = [
#             "%b %Y", "%B %Y",       # Oct 2023, October 2023
#             "%b-%Y", "%B-%Y",       # Oct-2023
#             "%m/%Y", "%Y",          # 10/2023, 2023
#             "%b %d, %Y",            # Oct 1, 2023
#             "%B %d, %Y"             # October 1, 2023
#         ]
        
#         for fmt in formats:
#             try:
#                 return datetime.strptime(d_str, fmt)
#             except ValueError:
#                 continue
#         return None

#     start_date = parse_date(start_str)
#     end_date = parse_date(end_str)

#     if start_date and end_date:
#         # Calculate difference in months
#         diff_months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month) + 1 
        
#         if diff_months < 1: diff_months = 1
        
#         # CHANGED: Return ONLY the total months count
#         return f"{diff_months} Months"
    
#     return duration_raw

# def transform_projects_data(projects_list):
#     transformed_projects = []
#     if not projects_list:
#         return []

#     for idx, project in enumerate(projects_list, 1):

#         tech_stack = project.get("tech_stack", [])
#         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)

#         name = project.get("project_name", "Unknown Project")
#         clouds = project.get("Clouds", "")
#         sf_clouds = project.get("Industry", "")

#         extra = ""

#         if clouds and clouds not in ["NA", "None", ""]:
#             extra = clouds

#         if sf_clouds and sf_clouds not in ["NA", "None", ""]:
#             if extra:
#                 extra = f"{extra}, {sf_clouds}"
#             else:
#                 extra = sf_clouds

#         if extra:
#             formatted_title = f"Project {idx}: {name} ({extra})"
#         else:
#             formatted_title = f"Project {idx}: {name}"


#         desc = project.get("description", [])
#         if isinstance(desc, str):
#             desc = [desc]

#         transformed_projects.append({
#             "title": formatted_title,
#             "role": project.get("Role", ""),
#             "industry": project.get("Industry", "NA"),

#             # ⭐ CHANGE IS HERE
#             "duration": calculate_duration_string(project.get("Duration", "")),

#             "link": project.get("link", ""),
#             "description": desc,
#             "tech_stack": tech_stack_str
#         })

#     return transformed_projects


# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     full_text = []
#     for para in doc.paragraphs: full_text.append(para.text)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells: full_text.append(cell.text)
#     return '\n'.join(full_text)

# def get_improved_extraction_prompt():
#     system_prompt = """You are an expert resume parser AI. 
#     CRITICAL RULES:
#     1. **Name Extraction**: Extract the candidate's ACTUAL NAME. Do NOT return 'Full Name', 'Candidate Name', or 'Name'.
#     2. **Experience = Projects**: Extract 'Work Experience' or 'Projects' into 'projects_descriptions'.
#     3. **Project Names**: If missing, **USE COMPANY NAME**. Do NOT return 'None'.
#     4. **Exact Content**: Extract 'description' text EXACTLY as written.
#     5. **Infer Clouds**: INFER Cloud/Domain (e.g. 'Web', 'Mobile', 'AI'). Do NOT return 'NA'.
#     6. **Main Table Clouds**: List Salesforce Clouds mentioned in projects (Sales, Service, etc.) or 'NA'.
#     7. **Skill Mapping**: Technical skills -> 'salesforce_expertise'. Soft skills -> Behavioral only.
#     8. **Duration**: Extract dates (e.g., "Oct 2022 - Present").
#     9. Output ONLY valid JSON."""

#     user_prompt = """Extract the resume content into this EXACT JSON structure.

# {
#   "Extract and Synthesize Candidate's Resume": {
#     "name": "Actual Candidate Name",
#     "current_job_role": "Current Job Title",
#     "experience": "Total Experience",
    
#     "crm_administration": ["Skill 1"],
#     "certifications": ["Cert 1"],
#     "salesforce_expertise": ["Technical Skill 1", "Technical Skill 2"],
#     "languages": ["Language 1"],
#     "salesforce_components": ["Component 1"],
#     "ticketing_case_management": ["Tool 1"],
#     "database": ["DB 1"],
#     "salesforce_clouds": ["Sales Cloud", "Service Cloud", "OR 'NA'"],
#     "soft_skills": ["Behavioral Skill 1"],

#     "expertise": [{"category_name": "Category", "skills": [{"skill_name": "Skill", "details": ["Tech"]}]}],
#     "development_tool": ["Tool 1"],
#     "projects_descriptions": [
#       {
#         "project_name": "Company Name OR Project Title",
#         "Clouds": "REQUIRED: Inferred Cloud or Domain (e.g. 'Web', 'Mobile'). NO 'NA'.",
#         "Role": "Job Role",
#         "Industry": "Inferred Industry",
#         "Duration": "Date Range",
#         "description": ["Exact line 1...", "Exact line 2..."], 
#         "tech_stack": ["Tech"],
#         "link": ""
#       }
#     ]
#   }
# }"""
#     return system_prompt, user_prompt

# def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
#     genai = configure_gemini()
#     print("\n🚀 Starting Salesforce Resume Extraction...")
    
#     system_prompt, user_prompt = get_improved_extraction_prompt()
#     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
#     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
#     response = None
#     success = False

#     for model_name in models_to_try:
#         if success: break
#         print(f"🔄 Attempting with model: {model_name}...")
#         try:
#             model = genai.GenerativeModel(model_name=model_name)
#         except Exception: continue
        
#         for attempt in range(3):
#             try:
#                 if pdf_file_path.lower().endswith('.docx'):
#                     resume_text = extract_text_from_docx(pdf_file_path)
#                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
#                 else:
#                     with open(pdf_file_path, 'rb') as f:
#                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
#                 success = True
#                 print(f"✅ Success with {model_name}!")
#                 break
#             except exceptions.ResourceExhausted:
#                 time.sleep(10)
#             except exceptions.NotFound:
#                 break 
#             except Exception as e:
#                 print(f"❌ Error with {model_name}: {e}")
#                 break

#     if not success or not response:
#         print("❌ All models failed.")
#         return False

#     try:
#         cleaned_response = clean_json_response(response.text)
#         parsed_output = json.loads(cleaned_response)
        
#         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
#             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
#         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
#     except Exception as e:
#         print(f"❌ Parsing Error: {e}")
#         return False

# def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
#     job_name = "Extract and Synthesize Candidate's Resume"
#     if job_name not in llm_response_dict: return False
#     main_data = llm_response_dict[job_name]

#     def format_list(key):
#         val = main_data.get(key, [])
#         if isinstance(val, list): return ", ".join(val) if val else "NA"
#         return val if val else "NA"
    
#     # === FIX NAME ===
#     raw_name = main_data.get("name", "NA")
#     if raw_name in ["Full Name", "Candidate Name", "Actual Candidate Name"]:
#         raw_name = "NA" 

#     template_data = {
#         "name": raw_name,
#         "job_role": main_data.get("current_job_role", "NA"),
#         "experience": main_data.get("experience", "NA"),
#         "crm_administration": format_list("crm_administration"),
#         "certifications": format_list("certifications"),
#         "salesforce_expertise": format_list("salesforce_expertise"),
#         "languages": format_list("languages"),
#         "salesforce_components": format_list("salesforce_components"),
#         "ticketing": format_list("ticketing_case_management"),
#         "database": format_list("database"),
#         "salesforce_clouds": format_list("salesforce_clouds"),
#         "soft_skills": format_list("soft_skills"),
#         "development_tools": format_list("development_tool"),
#         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
#     }
    
#     try:
#         with open(template_file_path, 'r', encoding='utf-8') as f:
#             template = Template(f.read())
#         with open(output_file_path, 'w', encoding='utf-8') as f:
#             f.write(template.render(**template_data))
#         return True
#     except Exception as e:
#         print(f"❌ Template Error: {e}")
#         return False

import google.generativeai as genai
from google.api_core import exceptions
from dotenv import load_dotenv
import os
import json
import time
from jinja2 import Template
import re
import docx
from datetime import datetime

load_dotenv()

def configure_gemini():
    api_key_gemini = os.getenv("GEMINI_API_KEY")
    if not api_key_gemini:
        raise ValueError("GEMINI_API_KEY not found in environment variables.")
    genai.configure(api_key=api_key_gemini)
    return genai

def clean_json_response(text):
    cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
    return cleaned.strip()

# =========================
# DURATION → MONTHS
# =========================
def calculate_duration_string(duration_raw):
    if not duration_raw or duration_raw.lower() in ["na", "none", "", "n/a"]:
        return "N/A"

    parts = re.split(r'\s*[-–]\s*|\s+to\s+', duration_raw, flags=re.IGNORECASE)

    if len(parts) != 2:
        return duration_raw

    start_str = parts[0].strip()
    end_str = parts[1].strip()

    def parse_date(d_str):
        if d_str.lower() in ["present", "current", "now", "ongoing"]:
            return datetime.now()

        formats = [
            "%b %Y","%B %Y","%b-%Y","%B-%Y",
            "%m/%Y","%Y","%b %d, %Y","%B %d, %Y"
        ]

        for fmt in formats:
            try:
                return datetime.strptime(d_str, fmt)
            except:
                continue
        return None

    start_date = parse_date(start_str)
    end_date = parse_date(end_str)

    if start_date and end_date:
        diff_months = (end_date.year - start_date.year)*12 + (end_date.month - start_date.month) + 1
        if diff_months < 1:
            diff_months = 1
        return f"{diff_months} Months"

    return duration_raw

# =========================
# PROJECT TRANSFORM
# (LONG DESCRIPTIONS)
# =========================
def transform_projects_data(projects_list):
    transformed_projects = []
    if not projects_list:
        return []

    for idx, project in enumerate(projects_list, 1):

        tech_stack = project.get("tech_stack", [])
        tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)

        name = project.get("project_name", "Unknown Project")
        clouds = project.get("Clouds", "")
        industry = project.get("Industry", "")

        extra = ""
        if clouds and clouds not in ["NA","None",""]:
            extra = clouds
        

        formatted_title = f"Project {idx}: ({extra})" if extra else f"Project {idx}:"

        desc = project.get("description", [])
        if isinstance(desc, str):
            desc = [desc]

        transformed_projects.append({
            "title": formatted_title,
            "role": project.get("Role",""),
            "industry": project.get("Industry","NA"),
            "duration": calculate_duration_string(project.get("Duration","")),
            "link": project.get("link",""),
            "description": desc,
            "tech_stack": tech_stack_str
        })

    return transformed_projects

# =========================
# DOCX TEXT EXTRACT
# =========================
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)

# =========================
# PROMPT
# =========================
def get_improved_extraction_prompt():
    system_prompt = """You are an expert resume parser AI.

CRITICAL EXTRACTION RULES:

1. Extract ALL bullet points for every project.
2. Do NOT summarize.
3. Do NOT shorten.
4. If project description has paragraph → split into multiple points.
5. Provide all points as it is.
6. Preserve full sentences.
7. Return long descriptions exactly as written.
8. Duration must remain date range (parser converts to months).
9. Output ONLY JSON.
10. If certification is shortformed then return full form (e.g. 'Salesforce PD-I' → 'Salesforce Platform Developer ').
"""

    user_prompt = """Extract resume into this JSON:

{
 "Extract and Synthesize Candidate's Resume":{
   "name":"",
   "current_job_role":"",
   "experience":"in years and months format (e.g. '3 years 2 months')",
   "crm_administration":[],
   "certifications":[],
   "salesforce_expertise":[],
   "languages":[],
   "salesforce_components":[],
   "ticketing_case_management":[],
   "database":[],
   "salesforce_clouds":[],
   "soft_skills":[],
   "development_tool":[],
   "projects_descriptions":[
     {
       "project_name":"",
       "Clouds":"",
       "Role":"",
       "Industry":"it should not be na, infer industry from description instead of technology prefer IT",
       "Duration":"Only in months and also add months word (e.g. '29 Months')",
       "description":[
         "Full long point 1",
         "Full long point 2",
         "Full long point 3",
         "Full long point 4"
       ],
       "tech_stack":[],
       "link":""
     }
   ]
 }
}"""
    return system_prompt, user_prompt


# =========================
# MAIN EXTRACTION
# =========================
def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
    genai = configure_gemini()
    system_prompt, user_prompt = get_improved_extraction_prompt()
    full_prompt = f"{system_prompt}\n\n{user_prompt}"

    models = [
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
    "gemini-2.5-flash",
    "gemini-2.5-pro"
    ]

    response = None
    for model_name in models:
        try:
            print("Trying model:", model_name)
            model = genai.GenerativeModel(model_name=model_name)

            if pdf_file_path.lower().endswith(".docx"):
                text = extract_text_from_docx(pdf_file_path)
                response = model.generate_content(f"{full_prompt}\n\n{text}")
            else:
                with open(pdf_file_path,'rb') as f:
                    response = model.generate_content([
                        full_prompt,
                        {"mime_type":"application/pdf","data":f.read()}
                    ])

            print("SUCCESS MODEL:", model_name)
            break

        except Exception as e:
            # Keep terminal output readable – log only a short error summary.
            err_msg = str(e)
            # Strip very long provider payloads (quota JSON, links, etc.)
            if " [" in err_msg:
                err_msg = err_msg.split(" [", 1)[0]
            if len(err_msg) > 200:
                err_msg = err_msg[:200] + "..."
            print(f"MODEL FAILED: {model_name} - {err_msg}")
            continue


    if not response:
        print("❌ Failed")
        return False

    cleaned = clean_json_response(response.text)
    parsed = json.loads(cleaned)

    if "Extract and Synthesize Candidate's Resume" not in parsed:
        parsed = {"Extract and Synthesize Candidate's Resume": parsed}

    return populate_template_with_llm_response(parsed, template_file_path, output_html_path)

# =========================
# TEMPLATE RENDER
# =========================
def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
    main = llm_response_dict["Extract and Synthesize Candidate's Resume"]

    def fmt(key):
        v = main.get(key,[])
        if isinstance(v,list):
            return ", ".join(v) if v else "NA"
        return v if v else "NA"

    template_data = {
        "name": main.get("name","NA"),
        "job_role": main.get("current_job_role","NA"),
        "experience": main.get("experience","NA"),
        "crm_administration": fmt("crm_administration"),
        "certifications": fmt("certifications"),
        "salesforce_expertise": fmt("salesforce_expertise"),
        "languages": fmt("languages"),
        "salesforce_components": fmt("salesforce_components"),
        "ticketing": fmt("ticketing_case_management"),
        "database": fmt("database"),
        "salesforce_clouds": fmt("salesforce_clouds"),
        "soft_skills": fmt("soft_skills"),
        "development_tools": fmt("development_tool"),
        "projects": transform_projects_data(main.get("projects_descriptions",[]))
    }

    with open(template_file_path,'r',encoding='utf-8') as f:
        template = Template(f.read())

    with open(output_file_path,'w',encoding='utf-8') as f:
        f.write(template.render(**template_data))

    return True
