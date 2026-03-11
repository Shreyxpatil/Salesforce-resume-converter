# # # # import google.generativeai as genai
# # # # from google.api_core import exceptions
# # # # from dotenv import load_dotenv
# # # # import os
# # # # import json
# # # # import time
# # # # from jinja2 import Template
# # # # import re
# # # # import docx
# # # # import traceback

# # # # load_dotenv()

# # # # def configure_gemini():
# # # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # # #     if not api_key_gemini:
# # # #         raise ValueError("GEMINI_API_KEY not found in environment variables. Please check your .env file")
    
# # # #     genai.configure(api_key=api_key_gemini)
# # # #     return genai

# # # # def clean_json_response(text):
# # # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # # #     return cleaned.strip()

# # # # def transform_expertise_data(expertise_list):
# # # #     """Transform expertise data from LLM format to template format"""
# # # #     transformed_expertise = []
    
# # # #     for category in expertise_list:
# # # #         transformed_skills = []
# # # #         for skill in category.get("skills", []):
# # # #             details = skill.get("details", [])
# # # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # # #             transformed_skills.append({
# # # #                 "area": skill.get("skill_name", ""),
# # # #                 "technologies": technologies
# # # #             })
# # # #         transformed_expertise.append({
# # # #             "category_name": category.get("category_name", ""),
# # # #             "skills": transformed_skills
# # # #         })
# # # #     return transformed_expertise

# # # # def transform_projects_data(projects_list):
# # # #     """Transform projects data from LLM format to template format"""
# # # #     transformed_projects = []
    
# # # #     for project in projects_list:
# # # #         tech_stack = project.get("tech_stack", [])
# # # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # # #         # Ensure description is a list for bullet points
# # # #         desc = project.get("description", [])
# # # #         if isinstance(desc, str):
# # # #             desc = [desc]
            
# # # #         transformed_projects.append({
# # # #             "title": project.get("project_name", ""),
# # # #             "role": project.get("Role", ""),
# # # #             "duration": project.get("Duration", ""),
# # # #             "link": project.get("link", None),
# # # #             "description": desc,
# # # #             "tech_stack": tech_stack_str
# # # #         })
# # # #     return transformed_projects

# # # # def extract_text_from_docx(docx_path):
# # # #     """Extract text content from a DOCX file"""
# # # #     doc = docx.Document(docx_path)
# # # #     full_text = []
# # # #     for para in doc.paragraphs: full_text.append(para.text)
# # # #     for table in doc.tables:
# # # #         for row in table.rows:
# # # #             for cell in row.cells: full_text.append(cell.text)
# # # #     return '\n'.join(full_text)

# # # # def get_improved_extraction_prompt():
# # # #     system_prompt = """You are an expert resume parser AI. Extract ALL information EXACTLY as written.
# # # #     CRITICAL: For project descriptions, return a list of bullet points strings."""
    
# # # #     user_prompt = """Extract and structure the resume content according to this EXACT JSON schema:

# # # # {
# # # #   "Extract and Synthesize Candidate's Resume": {
# # # #     "name": "Full name",
# # # #     "current_job_role": "Current job title",
# # # #     "experience": "Total experience",
# # # #     "expertise": [
# # # #       {
# # # #         "category_name": "Main category",
# # # #         "skills": [
# # # #           {
# # # #             "skill_name": "Specific skill",
# # # #             "details": ["technology1", "technology2"]
# # # #           }
# # # #         ]
# # # #       }
# # # #     ],
# # # #     "development_tool": ["tool1", "tool2"],
# # # #     "database": ["database1"],
# # # #     "projects_descriptions": [
# # # #       {
# # # #         "project_name": "Project title",
# # # #         "Role": "Role",
# # # #         "Duration": "Duration",
# # # #         "description": [
# # # #           "Bullet point 1",
# # # #           "Bullet point 2"
# # # #         ],
# # # #         "tech_stack": ["tech1"],
# # # #         "link": "Project URL"
# # # #       }
# # # #     ]
# # # #   }
# # # # }"""
# # # #     return system_prompt, user_prompt

# # # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # # #     genai = configure_gemini()
# # # #     print("\n🚀 Starting General Resume Extraction...")
    
# # # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # # #     # === UPDATED: VALID MODELS ONLY (From your check_models.py output) ===
# # # #     models_to_try = [
# # # #         "gemini-2.5-flash-lite", 
# # # #         "gemini-2.0-flash-lite",
# # # #         "gemini-2.5-flash",
# # # #         "gemini-2.0-flash",
# # # #         "gemini-2.5-pro",
# # # #     ]
    
# # # #     response = None
# # # #     success = False

# # # #     for model_name in models_to_try:
# # # #         if success: break
# # # #         print(f"🔄 Attempting with model: {model_name}...")
# # # #         try:
# # # #             model = genai.GenerativeModel(model_name=model_name)
# # # #         except Exception:
# # # #             continue
        
# # # #         for attempt in range(3):
# # # #             try:
# # # #                 if pdf_file_path.lower().endswith('.docx'):
# # # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # # #                 else:
# # # #                     with open(pdf_file_path, 'rb') as f:
# # # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # # #                 success = True
# # # #                 print(f"✅ Success with {model_name}!")
# # # #                 break
# # # #             except exceptions.ResourceExhausted:
# # # #                 print(f"⚠️ Quota exceeded for {model_name}. Waiting 10s...")
# # # #                 time.sleep(10)
# # # #             except exceptions.NotFound:
# # # #                 print(f"❌ Model {model_name} not found. Switching...")
# # # #                 break 
# # # #             except Exception as e:
# # # #                 print(f"❌ Error with {model_name}: {e}")
# # # #                 break

# # # #     if not success or not response:
# # # #         print("❌ All models failed.")
# # # #         return False

# # # #     try:
# # # #         cleaned_response = clean_json_response(response.text)
# # # #         parsed_output = json.loads(cleaned_response)
        
# # # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
        
# # # #     except Exception as e:
# # # #         print(f"❌ General Parser Error: {e}")
# # # #         return False

# # # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # # #     job_name = "Extract and Synthesize Candidate's Resume"
# # # #     if job_name not in llm_response_dict: return False
    
# # # #     main_data = llm_response_dict[job_name]
    
# # # #     template_data = {
# # # #         "name": main_data.get("name", ""),
# # # #         "job_role": main_data.get("current_job_role", ""),
# # # #         "experience": main_data.get("experience", ""),
# # # #         "expertise": transform_expertise_data(main_data.get("expertise", [])),
# # # #         "development_tools": ", ".join(main_data.get("development_tool", [])),
# # # #         "databases": ", ".join(main_data.get("database", [])),
# # # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # # #     }
    
# # # #     try:
# # # #         with open(template_file_path, 'r', encoding='utf-8') as file:
# # # #             template = Template(file.read())
# # # #         with open(output_file_path, 'w', encoding='utf-8') as file:
# # # #             file.write(template.render(**template_data))
# # # #         return True
# # # #     except Exception as e:
# # # #         print(f"❌ Error saving HTML file: {e}")
# # # #         return False

# # # #new

# # # import google.generativeai as genai
# # # from google.api_core import exceptions
# # # from dotenv import load_dotenv
# # # import os
# # # import json
# # # import time
# # # from jinja2 import Template
# # # import re
# # # import docx

# # # load_dotenv()

# # # def configure_gemini():
# # #     api_key_gemini = os.getenv("GEMINI_API_KEY")
# # #     # Fallback to backend config if available
# # #     if not api_key_gemini:
# # #         try:
# # #             from backend.config.config import settings
# # #             api_key_gemini = settings.GEMINI_API_KEY
# # #         except ImportError:
# # #             pass

# # #     if not api_key_gemini:
# # #         raise ValueError("GEMINI_API_KEY not found in environment variables.")
    
# # #     genai.configure(api_key=api_key_gemini)
# # #     return genai

# # # def clean_json_response(text):
# # #     """Remove markdown code blocks and extra whitespace from JSON response"""
# # #     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
# # #     return cleaned.strip()

# # # def transform_expertise_data(expertise_list):
# # #     transformed_expertise = []
# # #     if not expertise_list: return []
    
# # #     for category in expertise_list:
# # #         transformed_skills = []
# # #         for skill in category.get("skills", []):
# # #             details = skill.get("details", [])
# # #             technologies = ", ".join(details) if isinstance(details, list) else str(details)
# # #             transformed_skills.append({
# # #                 "area": skill.get("skill_name", ""),
# # #                 "technologies": technologies
# # #             })
# # #         transformed_expertise.append({
# # #             "category_name": category.get("category_name", ""),
# # #             "skills": transformed_skills
# # #         })
# # #     return transformed_expertise

# # # def transform_projects_data(projects_list):
# # #     transformed_projects = []
# # #     if not projects_list: return []
    
# # #     for project in projects_list:
# # #         tech_stack = project.get("tech_stack", [])
# # #         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
# # #         # === KEY CHANGE: Ensure description is always a LIST ===
# # #         desc = project.get("description", [])
# # #         if isinstance(desc, str):
# # #             # If AI returns a string despite instructions, split it by newlines or wrap it
# # #             if '\n' in desc:
# # #                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
# # #             else:
# # #                 desc = [desc]
        
# # #         transformed_projects.append({
# # #             "title": project.get("project_name", "Project"),
# # #             "role": project.get("Role", ""),
# # #             "duration": project.get("Duration", ""),
# # #             "link": project.get("link", None),
# # #             "description": desc, # Now strictly a list
# # #             "tech_stack": tech_stack_str
# # #         })
# # #     return transformed_projects

# # # def extract_text_from_docx(docx_path):
# # #     doc = docx.Document(docx_path)
# # #     full_text = []
# # #     for para in doc.paragraphs: full_text.append(para.text)
# # #     for table in doc.tables:
# # #         for row in table.rows:
# # #             for cell in row.cells: full_text.append(cell.text)
# # #     return '\n'.join(full_text)

# # # def get_improved_extraction_prompt():
# # #     # === YOUR EXACT PROMPT ===
# # #     system_prompt = """You are an expert resume parser AI with exceptional attention to detail. Your task is to extract ALL information from the provided resume and structure it into a clean, valid JSON object.

# # # CRITICAL RULES:
# # # 1. Extract information EXACTLY as written - do not summarize, rephrase, or modify
# # # 2. For project descriptions: Return them as a LIST of strings (bullet points). Include EVERY detail, responsibility, achievement, and technical specification.
# # # 3. Preserve ALL technical terms, frameworks, tools, and metrics exactly as stated
# # # 4. If information is missing or unclear, use empty string "" or empty array []
# # # 5. Output ONLY valid JSON - no explanations, no markdown, no extra text

# # # IMPORTANT: Your entire response must be a single, parseable JSON object that strictly follows the provided schema."""

# # #     user_prompt = """Extract and structure the resume content according to this EXACT JSON schema:

# # # {
# # #   "Extract and Synthesize Candidate's Resume": {
# # #     "name": "Full name of the candidate",
# # #     "current_job_role": "Current or most recent job title",
# # #     "experience": "Total years/months of experience (e.g., '5 Years', '3+ Years')",
# # #     "expertise": [
# # #       {
# # #         "category_name": "Main category (e.g., 'AI & Machine Learning', 'Web Development')",
# # #         "skills": [
# # #           {
# # #             "skill_name": "Specific skill area (e.g., 'Frameworks', 'Languages')",
# # #             "details": ["technology1", "technology2", "technology3"]
# # #           }
# # #         ]
# # #       }
# # #     ],
# # #     "development_tool": ["tool1", "tool2", "tool3"],
# # #     "database": ["database1", "database2"],
# # #     "projects_descriptions": [
# # #       {
# # #         "project_name": "Exact project title",
# # #         "Role": "Exact role for that project, if not specified, use empty string",
# # #         "Duration": "Exact duration for that project, if not specified, use empty string",
# # #         "description": [
# # #           "Bullet point 1 - detail",
# # #           "Bullet point 2 - detail",
# # #           "Bullet point 3 - detail"
# # #         ],
# # #         "tech_stack": ["tech1", "tech2", "tech3"],
# # #         "impact": ["impact statement 1", "impact statement 2"],
# # #         "link": "Project URL if provided, otherwise empty string"
# # #       }
# # #     ]
# # #   }
# # # }

# # # CRITICAL: The JSON must have the outer key "Extract and Synthesize Candidate's Resume" containing all the data.

# # # EXTRACTION GUIDELINES:

# # # 1. NAME & BASICS:
# # #    - Extract full name exactly as written
# # #    - Current job role should be the most recent/current position title
# # #    - Experience format: "X Years" or "X+ Years" or "X Months"

# # # 2. EXPERTISE SECTION:
# # #    - Group skills into logical categories (AI/ML, Web Dev, Cloud, etc.)
# # #    - Each category should have clear skill areas with specific technologies
# # #    - Example structure:
# # #      * Category: "Programming Languages"
# # #        - Skill: "Backend Languages" → ["Python", "Java", "Go"]
# # #        - Skill: "Frontend Languages" → ["JavaScript", "TypeScript"]

# # # 3. DEVELOPMENT TOOLS:
# # #    - List all tools: Git, Docker, Kubernetes, VS Code, Jenkins, etc.
# # #    - Extract as a flat array

# # # 4. DATABASES:
# # #    - List all databases: PostgreSQL, MongoDB, Redis, etc.
# # #    - Extract as a flat array

# # # 5. PROJECT DESCRIPTIONS (MOST IMPORTANT):
# # #    - Return the description as a LIST OF STRINGS (Array).
# # #    - Break the text into distinct bullet points.
# # #    - Capture ALL bullet points, features, and responsibilities.
# # #    - Preserve ALL technical details and metrics.
# # #    - Include impact statements (e.g., "Improved performance by 40%").
# # #    - Extract tech stack as array.
# # #    - Include project URL if mentioned.
# # #    - Extract "Role" and "Duration" if they are mentioned in the project description.

# # # REMEMBER: Wrap everything inside the "Extract and Synthesize Candidate's Resume" key.

# # # Now extract all information from the provided resume following these guidelines exactly."""
# # #     return system_prompt, user_prompt

# # # def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
# # #     genai = configure_gemini()
# # #     print("\n🚀 Starting General Resume Extraction...")
    
# # #     system_prompt, user_prompt = get_improved_extraction_prompt()
# # #     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
# # #     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.5-pro"]
    
# # #     response = None
# # #     success = False

# # #     for model_name in models_to_try:
# # #         if success: break
# # #         print(f"🔄 Attempting with model: {model_name}...")
# # #         try:
# # #             model = genai.GenerativeModel(model_name=model_name)
# # #         except Exception: continue
        
# # #         for attempt in range(3):
# # #             try:
# # #                 if pdf_file_path.lower().endswith('.docx'):
# # #                     resume_text = extract_text_from_docx(pdf_file_path)
# # #                     response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
# # #                 else:
# # #                     with open(pdf_file_path, 'rb') as f:
# # #                         response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
# # #                 success = True
# # #                 print(f"✅ Success with {model_name}!")
# # #                 break
# # #             except exceptions.ResourceExhausted:
# # #                 time.sleep(10)
# # #             except exceptions.NotFound:
# # #                 break 
# # #             except Exception as e:
# # #                 print(f"❌ Error with {model_name}: {e}")
# # #                 break

# # #     if not success or not response:
# # #         print("❌ All models failed.")
# # #         return False

# # #     try:
# # #         cleaned_response = clean_json_response(response.text)
# # #         parsed_output = json.loads(cleaned_response)
        
# # #         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
# # #             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
# # #         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
        
# # #     except Exception as e:
# # #         print(f"❌ General Parser Error: {e}")
# # #         return False

# # # def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
# # #     job_name = "Extract and Synthesize Candidate's Resume"
# # #     if job_name not in llm_response_dict: return False
    
# # #     main_data = llm_response_dict[job_name]
    
# # #     raw_name = main_data.get("name", "NA")
# # #     if raw_name in ["Full Name", "Actual Candidate Name", "Candidate Name"]:
# # #         raw_name = "NA"

# # #     def format_list(val):
# # #         if isinstance(val, list): return ", ".join(val) if val else "NA"
# # #         return val if val else "NA"

# # #     template_data = {
# # #         "name": raw_name,
# # #         "job_role": main_data.get("current_job_role", "NA"),
# # #         "experience": main_data.get("experience", "NA"),
# # #         "expertise": transform_expertise_data(main_data.get("expertise", [])),
# # #         "development_tools": format_list(main_data.get("development_tool", [])),
# # #         "databases": format_list(main_data.get("database", [])),
# # #         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
# # #     }
    
# # #     try:
# # #         with open(template_file_path, 'r', encoding='utf-8') as file:
# # #             template = Template(file.read())
# # #         with open(output_file_path, 'w', encoding='utf-8') as file:
# # #             file.write(template.render(**template_data))
# # #         return True
# # #     except Exception as e:
# # #         print(f"❌ Error saving HTML file: {e}")
# # #         return False

# import google.generativeai as genai
# from google.api_core import exceptions
# from dotenv import load_dotenv
# import os
# import json
# import time
# from jinja2 import Template
# import re
# import docx

# load_dotenv()

# def configure_gemini():
#     api_key_gemini = os.getenv("GEMINI_API_KEY")
#     if not api_key_gemini:
#         # Fallback check
#         try:
#             from backend.config.config import settings
#             api_key_gemini = settings.GEMINI_API_KEY
#         except ImportError:
#             pass
            
#     if not api_key_gemini:
#         raise ValueError("GEMINI_API_KEY not found.")
    
#     genai.configure(api_key=api_key_gemini)
#     return genai

# def clean_json_response(text):
#     cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
#     return cleaned.strip()

# def transform_expertise_data(expertise_list):
#     transformed_expertise = []
#     if not expertise_list: return []
    
#     for category in expertise_list:
#         transformed_skills = []
#         for skill in category.get("skills", []):
#             details = skill.get("details", [])
#             technologies = ", ".join(details) if isinstance(details, list) else str(details)
#             transformed_skills.append({
#                 "area": skill.get("skill_name", ""),
#                 "technologies": technologies
#             })
#         transformed_expertise.append({
#             "category_name": category.get("category_name", ""),
#             "skills": transformed_skills
#         })
#     return transformed_expertise

# def transform_projects_data(projects_list):
#     transformed_projects = []
#     if not projects_list: return []
    
#     for project in projects_list:
#         tech_stack = project.get("tech_stack", [])
#         tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
#         # Ensure description is a list
#         desc = project.get("description", [])
#         if isinstance(desc, str):
#             if '\n' in desc:
#                 desc = [line.strip() for line in desc.split('\n') if line.strip()]
#             else:
#                 desc = [desc]
        
#         # === FIX: INFER ROLE IF EMPTY ===
#         role = project.get("Role", "")
#         if not role or role.lower() in ["na", "none", ""]:
#             # Fallback: Use the main Job Role if available, or generic "Developer"
#             role = "Salesforce Developer" # Default fallback
            
#         # === FIX: DURATION FORMATTING ===
#         duration = project.get("Duration", "")
#         if not duration or duration.lower() in ["na", "none", ""]:
#             duration = "Present" # Default if totally missing

#         transformed_projects.append({
#             "title": project.get("project_name", "Project"),
#             "role": role,
#             "duration": duration,
#             "link": project.get("link", None),
#             "description": desc,
#             "tech_stack": tech_stack_str
#         })
#     return transformed_projects

# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     full_text = []
#     for para in doc.paragraphs: full_text.append(para.text)
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells: full_text.append(cell.text)
#     return '\n'.join(full_text)

# def get_improved_extraction_prompt():
#     system_prompt = """You are an expert resume parser AI.
#     CRITICAL RULES:
#     1. **Role & Duration**: You MUST extract a Role and Duration for EVERY project.
#        - **Duration Format**: Convert ALL dates to 'Mon YYYY - Mon YYYY' (e.g., 'Jan 2024 - Oct 2024').
#        - If Role is not stated, INFER it from the project context (e.g., "Developer", "Tester").
#     2. **Expertise Grouping**: Structure skills exactly as the schema requests (Category -> Sub-Category -> Details).
#     3. **Descriptions**: Extract as bullet points.
#     4. **Output**: JSON only."""

#     user_prompt = """Extract the resume into this EXACT JSON structure:

# {
#   "Extract and Synthesize Candidate's Resume": {
#     "name": "Candidate Name",
#     "current_job_role": "Current Job Title",
#     "experience": "Total Experience",
#     "expertise": [
#       {
#         "category_name": "Main Category (e.g. Salesforce Development)",
#         "skills": [
#           {
#             "skill_name": "Sub-Category (e.g. Core Salesforce)",
#             "details": ["Apex", "Triggers", "LWC"]
#           }
#         ]
#       }
#     ],
#     "development_tool": ["Jira", "Git"],
#     "database": ["MySQL"],
#     "projects_descriptions": [
#       {
#         "project_name": "Project Title",
#         "Role": "Job Role (REQUIRED)",
#         "Duration": "Jan 2023 - Dec 2023 (REQUIRED FORMAT)",
#         "description": ["Bullet 1", "Bullet 2"],
#         "tech_stack": ["Java", "AWS"],
#         "link": ""
#       }
#     ]
#   }
# }"""
#     return system_prompt, user_prompt

# def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
#     genai = configure_gemini()
#     print("\n🚀 Starting General Resume Extraction...")
    
#     system_prompt, user_prompt = get_improved_extraction_prompt()
#     full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
#     models_to_try = ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite", "gemini-2.5-flash"]
    
#     response = None
#     success = False

#     for model_name in models_to_try:
#         if success: break
#         print(f"🔄 Attempting with model: {model_name}...")
#         try:
#             model = genai.GenerativeModel(model_name=model_name)
#             if pdf_file_path.lower().endswith('.docx'):
#                 resume_text = extract_text_from_docx(pdf_file_path)
#                 response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
#             else:
#                 with open(pdf_file_path, 'rb') as f:
#                     response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
#             success = True
#         except Exception as e:
#             print(f"❌ Error with {model_name}: {e}")
#             continue

#     if not success or not response:
#         print("❌ All models failed.")
#         return False

#     try:
#         cleaned_response = clean_json_response(response.text)
#         parsed_output = json.loads(cleaned_response)
        
#         if "Extract and Synthesize Candidate's Resume" not in parsed_output:
#             parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
#         return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
        
#     except Exception as e:
#         print(f"❌ General Parser Error: {e}")
#         return False

# def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
#     job_name = "Extract and Synthesize Candidate's Resume"
#     if job_name not in llm_response_dict: return False
    
#     main_data = llm_response_dict[job_name]
    
#     raw_name = main_data.get("name", "NA")
#     if raw_name in ["Full Name", "Candidate Name"]: raw_name = "NA"

#     def format_list(val):
#         if isinstance(val, list): return ", ".join(val) if val else "NA"
#         return val if val else "NA"

#     template_data = {
#         "name": raw_name,
#         "job_role": main_data.get("current_job_role", "NA"),
#         "experience": main_data.get("experience", "NA"),
#         "expertise": transform_expertise_data(main_data.get("expertise", [])),
#         "development_tools": format_list(main_data.get("development_tool", [])),
#         "databases": format_list(main_data.get("database", [])),
#         "projects": transform_projects_data(main_data.get("projects_descriptions", []))
#     }
    
#     try:
#         with open(template_file_path, 'r', encoding='utf-8') as file:
#             template = Template(file.read())
#         with open(output_file_path, 'w', encoding='utf-8') as file:
#             file.write(template.render(**template_data))
#         return True
#     except Exception as e:
#         print(f"❌ Error saving HTML file: {e}")
#         return False


import warnings

# Keep terminal output clean (library deprecation warning).
warnings.filterwarnings(
    "ignore",
    category=FutureWarning,
    message=r"All support for the `google\.generativeai` package has ended\..*",
)

import google.generativeai as genai
from google.api_core import exceptions
from dotenv import load_dotenv
import os
import json
import time
from jinja2 import Template
import re
import docx

load_dotenv()

def configure_gemini():
    api_key_gemini = os.getenv("GEMINI_API_KEY")
    if not api_key_gemini:
        try:
            from backend.config.config import settings
            api_key_gemini = settings.GEMINI_API_KEY
        except ImportError:
            pass

    if not api_key_gemini:
        raise ValueError("GEMINI_API_KEY not found in environment variables.")
    
    genai.configure(api_key=api_key_gemini)
    return genai

def clean_json_response(text):
    """Remove markdown code blocks and extra whitespace from JSON response"""
    cleaned = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.MULTILINE)
    return cleaned.strip()

def transform_expertise_data(expertise_list):
    transformed_expertise = []
    if not expertise_list: return []
    
    for category in expertise_list:
        transformed_skills = []
        for skill in category.get("skills", []):
            details = skill.get("details", [])
            technologies = ", ".join(details) if isinstance(details, list) else str(details)
            transformed_skills.append({
                "area": skill.get("skill_name", ""),
                "technologies": technologies
            })
        transformed_expertise.append({
            "category_name": category.get("category_name", ""),
            "skills": transformed_skills
        })
    return transformed_expertise

def transform_projects_data(projects_list):
    transformed_projects = []
    if not projects_list: return []
    
    for project in projects_list:
        tech_stack = project.get("tech_stack", [])
        tech_stack_str = ", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack)
        
        # Ensure description is a list
        desc = project.get("description", [])
        if isinstance(desc, str):
            if '\n' in desc:
                desc = [line.strip() for line in desc.split('\n') if line.strip()]
            else:
                desc = [desc]
        
        # Infer Role if missing
        role = project.get("Role", "")
        if not role or role.lower() in ["na", "none", ""]:
            role = "Salesforce Developer" 
            
        # Format Duration if missing
        duration = project.get("Duration", "")
        if not duration or duration.lower() in ["na", "none", ""]:
            duration = "Present"

        transformed_projects.append({
            "title": project.get("project_name", "Project"),
            "role": role,
            "duration": duration,
            "link": project.get("link", None),
            "description": desc,
            "tech_stack": tech_stack_str
        })
    return transformed_projects

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs: full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells: full_text.append(cell.text)
    return '\n'.join(full_text)

def get_improved_extraction_prompt():
    system_prompt = """You are an expert resume parser AI.
    CRITICAL RULES:
    1. **Role & Duration**: You MUST extract a Role and Duration for EVERY project.
       - **Duration Format**: Convert ALL dates to 'Mon YYYY - Mon YYYY' (e.g., 'Jan 2024 - Oct 2024').
       - If Role is not stated, INFER it from the project context.
    2. **Expertise Grouping**: Structure skills exactly as the schema requests (Category -> Sub-Category -> Details).
    3. **Descriptions**: Extract the text EXACTLY as it appears in the resume ("As Is"). Do not summarize or rephrase. Return as a list of strings (split by bullet points or new lines).
    4. **Output**: JSON only."""

    user_prompt = """Extract the resume into this EXACT JSON structure:

{
  "Extract and Synthesize Candidate's Resume": {
    "name": "Candidate Name",
    "current_job_role": "Current Job Title",
    "experience": "Total Experience",
    "expertise": [
      {
        "category_name": "Main Category (e.g. Salesforce Development)",
        "skills": [
          {
            "skill_name": "Sub-Category (e.g. Core Salesforce)",
            "details": ["Apex", "Triggers", "LWC"]
          }
        ]
      }
    ],
    "development_tool": ["Jira", "Git"],
    "database": ["MySQL"],
    "projects_descriptions": [
      {
        "project_name": "Project Title",
        "Role": "Job Role (REQUIRED)",
        "Duration": "Jan 2023 - Dec 2023 (REQUIRED FORMAT)",
        "description": [
          "Exact text from resume line 1",
          "Exact text from resume line 2"
        ],
        "tech_stack": ["Java", "AWS"],
        "link": ""
      }
    ]
  }
}"""
    return system_prompt, user_prompt

def extract_resume_data(pdf_file_path, template_file_path, output_html_path):
    genai = configure_gemini()
    print("\n🚀 Starting General Resume Extraction...")
    
    system_prompt, user_prompt = get_improved_extraction_prompt()
    full_prompt = f"{system_prompt}\n\n{user_prompt}"
    
    # === UPDATED MODEL LIST (Based on your available models) ===
    models_to_try = [
        "gemini-2.5-flash",       # Primary
        "gemini-2.0-flash",       # Secondary
        "gemini-2.5-pro",         # Fallback (Powerful)
        "gemini-2.0-flash-lite"   # Last resort (Low quota)
    ]
    
    response = None
    success = False

    for model_name in models_to_try:
        if success: break
        print(f"🔄 Attempting with model: {model_name}...")
        try:
            model = genai.GenerativeModel(model_name=model_name)
            
            if pdf_file_path.lower().endswith('.docx'):
                resume_text = extract_text_from_docx(pdf_file_path)
                response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
            else:
                with open(pdf_file_path, 'rb') as f:
                    response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
            
            success = True
            print(f"✅ Success with {model_name}!")
        
        except exceptions.ResourceExhausted:
            print(f"⚠️ Quota exceeded for {model_name}. Waiting 30 seconds to retry...")
            time.sleep(30)
            try:
                print(f"🔄 Retrying {model_name}...")
                if pdf_file_path.lower().endswith('.docx'):
                    response = model.generate_content(f"{full_prompt}\n\nRESUME CONTENT:\n{resume_text}")
                else:
                    with open(pdf_file_path, 'rb') as f:
                        response = model.generate_content([full_prompt, {"mime_type": "application/pdf", "data": f.read()}])
                success = True
                print(f"✅ Success on retry with {model_name}!")
            except Exception as e:
                 print(f"❌ Retry failed: {e}")
                 
        except exceptions.NotFound:
            print(f"❌ Model {model_name} not found. Skipping...")
            continue
        except Exception as e:
            print(f"❌ Error with {model_name}: {e}")
            continue

    if not success or not response:
        print("❌ All models failed.")
        return False

    try:
        cleaned_response = clean_json_response(response.text)
        parsed_output = json.loads(cleaned_response)
        
        if "Extract and Synthesize Candidate's Resume" not in parsed_output:
            parsed_output = {"Extract and Synthesize Candidate's Resume": parsed_output}
            
        return populate_template_with_llm_response(parsed_output, template_file_path, output_html_path)
        
    except Exception as e:
        print(f"❌ General Parser Error: {e}")
        return False

def populate_template_with_llm_response(llm_response_dict, template_file_path, output_file_path):
    job_name = "Extract and Synthesize Candidate's Resume"
    if job_name not in llm_response_dict: return False
    
    main_data = llm_response_dict[job_name]
    
    raw_name = main_data.get("name", "NA")
    if raw_name in ["Full Name", "Candidate Name"]: raw_name = "NA"

    def format_list(val):
        if isinstance(val, list): return ", ".join(val) if val else "NA"
        return val if val else "NA"

    template_data = {
        "name": raw_name,
        "job_role": main_data.get("current_job_role", "NA"),
        "experience": main_data.get("experience", "NA"),
        "expertise": transform_expertise_data(main_data.get("expertise", [])),
        "development_tools": format_list(main_data.get("development_tool", [])),
        "databases": format_list(main_data.get("database", [])),
        "projects": transform_projects_data(main_data.get("projects_descriptions", []))
    }
    
    try:
        with open(template_file_path, 'r', encoding='utf-8') as file:
            template = Template(file.read())
        with open(output_file_path, 'w', encoding='utf-8') as file:
            file.write(template.render(**template_data))
        return True
    except Exception as e:
        print(f"❌ Error saving HTML file: {e}")
        return False