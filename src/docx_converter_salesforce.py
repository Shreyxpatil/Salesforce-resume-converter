# # # # # # # # # # # from docx import Document
# # # # # # # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # # # # # # from docx.oxml.ns import qn
# # # # # # # # # # # from docx.oxml import OxmlElement
# # # # # # # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # # # # # # import os

# # # # # # # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # # # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # # # # # # #         html_content = file.read()
    
# # # # # # # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # # # # # # #     doc = Document()
    
# # # # # # # # # # #     # Margins
# # # # # # # # # # #     for section in doc.sections:
# # # # # # # # # # #         section.top_margin = Inches(0.5)
# # # # # # # # # # #         section.bottom_margin = Inches(0.5)
# # # # # # # # # # #         section.left_margin = Inches(0.5)
# # # # # # # # # # #         section.right_margin = Inches(0.5)
    
# # # # # # # # # # #     # Logo
# # # # # # # # # # #     try:
# # # # # # # # # # #         logo_filename = "download.png"
# # # # # # # # # # #         if os.path.exists(logo_filename):
# # # # # # # # # # #             logo_para = doc.add_paragraph()
# # # # # # # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # # # # # # #             run = logo_para.add_run()
# # # # # # # # # # #             run.add_picture(logo_filename, width=Inches(3.0))
# # # # # # # # # # #     except Exception:
# # # # # # # # # # #         pass
    
# # # # # # # # # # #     # --- 1. TOP TABLE (Blue Header) ---
# # # # # # # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # # # # # # #     if main_table:
# # # # # # # # # # #         doc_table = doc.add_table(rows=0, cols=2)
# # # # # # # # # # #         doc_table.style = 'Table Grid'
# # # # # # # # # # #         doc_table.columns[0].width = Inches(1.875)
# # # # # # # # # # #         doc_table.columns[1].width = Inches(5.625)
# # # # # # # # # # #         set_table_borders(doc_table)
        
# # # # # # # # # # #         rows = main_table.find_all('tr')
# # # # # # # # # # #         for row in rows:
# # # # # # # # # # #             cells = row.find_all('td')
# # # # # # # # # # #             if len(cells) == 2:
# # # # # # # # # # #                 doc_row = doc_table.add_row().cells
# # # # # # # # # # #                 # is_header=True forces BLUE background
# # # # # # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=True, center_align=True)
# # # # # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # # # #                 process_html_content(doc_row[1], cells[1])
        
# # # # # # # # # # #     # --- 2. PROJECTS TABLE (White Header for Salesforce) ---
# # # # # # # # # # #     project_title_div = soup.find('div', {'class': 'project-title-bar'})
# # # # # # # # # # #     projects_table_html = soup.find('table', {'class': 'projects-table'})
    
# # # # # # # # # # #     if projects_table_html and project_title_div:
# # # # # # # # # # #         proj_table = doc.add_table(rows=0, cols=2)
# # # # # # # # # # #         proj_table.style = 'Table Grid'
# # # # # # # # # # #         proj_table.columns[0].width = Inches(1.875)
# # # # # # # # # # #         proj_table.columns[1].width = Inches(5.625)
# # # # # # # # # # #         set_table_borders(proj_table)

# # # # # # # # # # #         # Title Bar
# # # # # # # # # # #         title_row_cells = proj_table.add_row().cells
# # # # # # # # # # #         merged_cell = title_row_cells[0].merge(title_row_cells[1])
# # # # # # # # # # #         set_cell_content(merged_cell, project_title_div.get_text(strip=True), is_header=True, font_size=14)
        
# # # # # # # # # # #         rows = projects_table_html.find_all('tr')
# # # # # # # # # # #         for row in rows:
# # # # # # # # # # #             cells = row.find_all('td')
# # # # # # # # # # #             if len(cells) == 2:
# # # # # # # # # # #                 doc_row = proj_table.add_row().cells
                
# # # # # # # # # # #                 # CHANGED: is_header=False (No Blue), but italic=True and center_align=True
# # # # # # # # # # #                 set_cell_content(
# # # # # # # # # # #                     doc_row[0], 
# # # # # # # # # # #                     cells[0].get_text(strip=True), 
# # # # # # # # # # #                     is_header=False, 
# # # # # # # # # # #                     center_align=True, 
# # # # # # # # # # #                     italic=True,
# # # # # # # # # # #                     bold=False # Ensure text is not white
# # # # # # # # # # #                 )
                
# # # # # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # # # # # # #     doc.save(docx_file_path)

# # # # # # # # # # # def set_table_borders(table):
# # # # # # # # # # #     tbl = table._tbl
# # # # # # # # # # #     tblPr = tbl.tblPr
# # # # # # # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # # # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # # # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # # # # # # #         border.set(qn('w:val'), 'single')
# # # # # # # # # # #         border.set(qn('w:sz'), '4')
# # # # # # # # # # #         border.set(qn('w:space'), '0')
# # # # # # # # # # #         border.set(qn('w:color'), '000000')
# # # # # # # # # # #         tblBorders.append(border)
# # # # # # # # # # #     tblPr.append(tblBorders)

# # # # # # # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # # # # # # # #     cell.text = ""
# # # # # # # # # # #     paragraph = cell.paragraphs[0]
# # # # # # # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # # # # # # #     paragraph.paragraph_format.space_after = Pt(0)
    
# # # # # # # # # # #     if center_align: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# # # # # # # # # # #     else: paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # # # # # # #     if is_header:
# # # # # # # # # # #         # BLUE Background logic
# # # # # # # # # # #         set_cell_background(cell, "0d3d62")
# # # # # # # # # # #         run = paragraph.add_run(text)
# # # # # # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
# # # # # # # # # # #         run.font.bold = True
# # # # # # # # # # #         run.font.size = Pt(font_size)
# # # # # # # # # # #     else:
# # # # # # # # # # #         # WHITE Background logic
# # # # # # # # # # #         if text:
# # # # # # # # # # #             run = paragraph.add_run(text)
# # # # # # # # # # #             run.font.size = Pt(font_size)
# # # # # # # # # # #             run.font.italic = italic
# # # # # # # # # # #             run.font.bold = bold
# # # # # # # # # # #             run.font.color.rgb = RGBColor(0, 0, 0) # Black Text
    
# # # # # # # # # # #     if center_align:
# # # # # # # # # # #         cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # # # # # # def set_cell_background(cell, color_hex):
# # # # # # # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # # # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # # # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # # # # # # def process_html_content(cell, html_element):
# # # # # # # # # # #     text = html_element.get_text(strip=True)
# # # # # # # # # # #     if text:
# # # # # # # # # # #         para = cell.add_paragraph()
# # # # # # # # # # #         run = para.add_run(text)
# # # # # # # # # # #         run.font.size = Pt(12)
# # # # # # # # # # #         run.font.name = 'Arial'

# # # # # # # # # # # def process_project_content(cell, html_element):
# # # # # # # # # # #     cell.text = ""
    
# # # # # # # # # # #     for element in html_element.children:
# # # # # # # # # # #         if isinstance(element, Tag):
# # # # # # # # # # #             if element.name == 'div':
# # # # # # # # # # #                 para = cell.add_paragraph()
# # # # # # # # # # #                 para.paragraph_format.space_after = Pt(2)
                
# # # # # # # # # # #                 strong = element.find('strong')
# # # # # # # # # # #                 if strong:
# # # # # # # # # # #                     label_text = strong.get_text(strip=True)
# # # # # # # # # # #                     run = para.add_run(label_text + " ")
# # # # # # # # # # #                     run.bold = True
# # # # # # # # # # #                     run.font.size = Pt(12)
# # # # # # # # # # #                     run.font.name = 'Arial'
                    
# # # # # # # # # # #                     value_text = element.get_text().replace(label_text, '', 1).strip()
# # # # # # # # # # #                     if value_text:
# # # # # # # # # # #                         run2 = para.add_run(value_text)
# # # # # # # # # # #                         run2.font.size = Pt(12)
# # # # # # # # # # #                         run2.font.name = 'Arial'
# # # # # # # # # # #                 else:
# # # # # # # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # # # # # # #                     run.font.size = Pt(12)
# # # # # # # # # # #                     run.font.name = 'Arial'
# # # # # # # # # # #                     if "Project Description" in run.text:
# # # # # # # # # # #                          run.bold = True

# # # # # # # # # # #             elif element.name == 'ol':
# # # # # # # # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # # # # # # # #                     para = cell.add_paragraph()
# # # # # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
                    
# # # # # # # # # # #                     run_num = para.add_run(f"{i}. ")
# # # # # # # # # # #                     run_num.font.size = Pt(12)
# # # # # # # # # # #                     run_num.font.name = 'Arial'
                    
# # # # # # # # # # #                     run_text = para.add_run(li.get_text(strip=True))
# # # # # # # # # # #                     run_text.font.size = Pt(12)
# # # # # # # # # # #                     run_text.font.name = 'Arial'

# # # # # # # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # # # # # # #     try:
# # # # # # # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # # # # # # #     except Exception as e:
# # # # # # # # # # #         print(f"Conversion failed: {e}")

# # # # # # # # # # from docx import Document
# # # # # # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # # # # # from docx.oxml.ns import qn
# # # # # # # # # # from docx.oxml import OxmlElement
# # # # # # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # # # # # import os

# # # # # # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # # # # # #         html_content = file.read()
    
# # # # # # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # # # # # #     doc = Document()
    
# # # # # # # # # #     # Set Margins
# # # # # # # # # #     for section in doc.sections:
# # # # # # # # # #         section.top_margin = Inches(0.5)
# # # # # # # # # #         section.bottom_margin = Inches(0.5)
# # # # # # # # # #         section.left_margin = Inches(0.5)
# # # # # # # # # #         section.right_margin = Inches(0.5)
    
# # # # # # # # # #     # === 1. ADD LOGO ===
# # # # # # # # # #     # Using the specific path you requested
# # # # # # # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # # # # # # #     if os.path.exists(logo_path):
# # # # # # # # # #         try:
# # # # # # # # # #             logo_para = doc.add_paragraph()
# # # # # # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # # # # # #             run = logo_para.add_run()
# # # # # # # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # # # # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # # # # # # #         except Exception as e:
# # # # # # # # # #             print(f"⚠️ Could not add logo: {e}")
# # # # # # # # # #     else:
# # # # # # # # # #         # Fallback to current directory if specific path fails
# # # # # # # # # #         if os.path.exists("download.png"):
# # # # # # # # # #             try:
# # # # # # # # # #                 logo_para = doc.add_paragraph()
# # # # # # # # # #                 run = logo_para.add_run()
# # # # # # # # # #                 run.add_picture("download.png", width=Inches(2.5))
# # # # # # # # # #             except: pass

# # # # # # # # # #     # === 2. MAIN INFO TABLE (Blue Headers) ===
# # # # # # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # # # # # #     if main_table:
# # # # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # # # #         t.style = 'Table Grid'
# # # # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # # # #         set_table_borders(t)
        
# # # # # # # # # #         for row in main_table.find_all('tr'):
# # # # # # # # # #             cells = row.find_all('td')
# # # # # # # # # #             if len(cells) == 2:
# # # # # # # # # #                 doc_row = t.add_row().cells
# # # # # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=True, center_align=True)
# # # # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # # #                 process_html_content(doc_row[1], cells[1])

# # # # # # # # # #     # === 3. PROJECTS TABLE (White Sidebar) ===
# # # # # # # # # #     proj_div = soup.find('div', {'class': 'project-title-bar'})
# # # # # # # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
    
# # # # # # # # # #     if proj_table and proj_div:
# # # # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # # # #         t.style = 'Table Grid'
# # # # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # # # #         set_table_borders(t)
        
# # # # # # # # # #         # Title Bar Row
# # # # # # # # # #         row = t.add_row().cells
# # # # # # # # # #         merged = row[0].merge(row[1])
# # # # # # # # # #         set_cell_content(merged, proj_div.get_text(strip=True), is_header=True, font_size=14)
        
# # # # # # # # # #         # Project Rows
# # # # # # # # # #         for row in proj_table.find_all('tr'):
# # # # # # # # # #             cells = row.find_all('td')
# # # # # # # # # #             if len(cells) == 2:
# # # # # # # # # #                 doc_row = t.add_row().cells
                
# # # # # # # # # #                 # Left Column: White background, Bold/Italic text, Centered
# # # # # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # # # # # # #                 # Right Column: Project Details
# # # # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # # # # # #     doc.save(docx_file_path)

# # # # # # # # # # def set_table_borders(table):
# # # # # # # # # #     tbl = table._tbl
# # # # # # # # # #     tblPr = tbl.tblPr
# # # # # # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # # # # # #         border.set(qn('w:val'), 'single')
# # # # # # # # # #         border.set(qn('w:sz'), '4')
# # # # # # # # # #         border.set(qn('w:space'), '0')
# # # # # # # # # #         border.set(qn('w:color'), '000000')
# # # # # # # # # #         tblBorders.append(border)
# # # # # # # # # #     tblPr.append(tblBorders)

# # # # # # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # # # # # # #     cell.text = ""
# # # # # # # # # #     paragraph = cell.paragraphs[0]
# # # # # # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # # # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # # # # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # # # # # #     if is_header:
# # # # # # # # # #         set_cell_background(cell, "0d3d62")
# # # # # # # # # #         run = paragraph.add_run(text)
# # # # # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # # # # #         run.font.bold = True
# # # # # # # # # #         run.font.size = Pt(font_size)
# # # # # # # # # #     else:
# # # # # # # # # #         if text:
# # # # # # # # # #             run = paragraph.add_run(text)
# # # # # # # # # #             run.font.size = Pt(font_size)
# # # # # # # # # #             run.font.italic = italic
# # # # # # # # # #             run.font.bold = bold
    
# # # # # # # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # # # # # def set_cell_background(cell, color_hex):
# # # # # # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # # # # # def process_html_content(cell, html_element):
# # # # # # # # # #     """Handles standard content in the main table"""
# # # # # # # # # #     text = html_element.get_text(strip=True)
# # # # # # # # # #     if text:
# # # # # # # # # #         para = cell.add_paragraph()
# # # # # # # # # #         para.add_run(text).font.size = Pt(12)

# # # # # # # # # # def process_project_content(cell, html_element):
# # # # # # # # # #     """Handles content in the Project Right Column, including Bullet Points"""
# # # # # # # # # #     cell.text = ""
# # # # # # # # # #     for element in html_element.children:
# # # # # # # # # #         if isinstance(element, Tag):
# # # # # # # # # #             # 1. Handle DIVs (Role, Duration, Tech Stack labels)
# # # # # # # # # #             if element.name == 'div':
# # # # # # # # # #                 para = cell.add_paragraph()
# # # # # # # # # #                 para.paragraph_format.space_after = Pt(2)
                
# # # # # # # # # #                 strong = element.find('strong')
# # # # # # # # # #                 if strong:
# # # # # # # # # #                     label = strong.get_text(strip=True)
# # # # # # # # # #                     para.add_run(label + " ").bold = True
# # # # # # # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # # # # # # #                     if val: para.add_run(val).font.size = Pt(12)
# # # # # # # # # #                 else:
# # # # # # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # # # # # #                     run.font.size = Pt(12)
# # # # # # # # # #                     if "Project Description" in run.text: 
# # # # # # # # # #                         run.bold = True
# # # # # # # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # # # # # # #             # 2. Handle UL (Bullet Points) - This is the key logic
# # # # # # # # # #             elif element.name == 'ul':
# # # # # # # # # #                 for li in element.find_all('li'):
# # # # # # # # # #                     para = cell.add_paragraph()
# # # # # # # # # #                     # Indent to make it look like a list
# # # # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # # # # #                     # Add bullet char manually to ensure compatibility
# # # # # # # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # # # # # # #                     run.font.size = Pt(12)

# # # # # # # # # #             # 3. Fallback for Paragraphs
# # # # # # # # # #             elif element.name == 'p':
# # # # # # # # # #                 para = cell.add_paragraph()
# # # # # # # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # # # # # # #                 para.add_run(element.get_text(strip=True)).font.size = Pt(12)

# # # # # # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # # # # # #     try:
# # # # # # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # # # # # #     except Exception as e:
# # # # # # # # # #         print(f"Conversion failed: {e}")



# # # # # # # # # from docx import Document
# # # # # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # # # # from docx.oxml.ns import qn
# # # # # # # # # from docx.oxml import OxmlElement
# # # # # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # # # # import os

# # # # # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # # # # #         html_content = file.read()
    
# # # # # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # # # # #     doc = Document()
    
# # # # # # # # #     # Set Margins
# # # # # # # # #     for section in doc.sections:
# # # # # # # # #         section.top_margin = Inches(0.5)
# # # # # # # # #         section.bottom_margin = Inches(0.5)
# # # # # # # # #         section.left_margin = Inches(0.5)
# # # # # # # # #         section.right_margin = Inches(0.5)
    
# # # # # # # # #     # === 1. ADD LOGO ===
# # # # # # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # # # # # #     if os.path.exists(logo_path):
# # # # # # # # #         try:
# # # # # # # # #             logo_para = doc.add_paragraph()
# # # # # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # # # # #             run = logo_para.add_run()
# # # # # # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # # # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # # # # # #         except Exception: pass
# # # # # # # # #     else:
# # # # # # # # #         if os.path.exists("download.png"):
# # # # # # # # #             try:
# # # # # # # # #                 logo_para = doc.add_paragraph()
# # # # # # # # #                 run = logo_para.add_run()
# # # # # # # # #                 run.add_picture("download.png", width=Inches(2.5))
# # # # # # # # #             except: pass

# # # # # # # # #     # === 2. MAIN INFO TABLE (Blue Headers) ===
# # # # # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # # # # #     if main_table:
# # # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # # #         t.style = 'Table Grid'
# # # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # # #         set_table_borders(t)
        
# # # # # # # # #         for row in main_table.find_all('tr'):
# # # # # # # # #             cells = row.find_all('td')
# # # # # # # # #             if len(cells) == 2:
# # # # # # # # #                 doc_row = t.add_row().cells
# # # # # # # # #                 left_text = cells[0].get_text(strip=True)
                
# # # # # # # # #                 # LEFT CELL
# # # # # # # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=True)
                
# # # # # # # # #                 # RIGHT CELL (CHECK FOR NAME STYLE)
# # # # # # # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # # # # # # #                 if is_name_cell:
# # # # # # # # #                     # BLUE BG, WHITE TEXT
# # # # # # # # #                     set_cell_content(doc_row[1], "", is_header=True)
# # # # # # # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # # # # # # #                 else:
# # # # # # # # #                     # STANDARD
# # # # # # # # #                     set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # #                     process_html_content(doc_row[1], cells[1])

# # # # # # # # #     # === 3. PROJECTS TABLE ===
# # # # # # # # #     # Look for the new class name for text-only header
# # # # # # # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # # # # # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # # # # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
    
# # # # # # # # #     if proj_table and proj_div:
# # # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # # #         t.style = 'Table Grid'
# # # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # # #         set_table_borders(t)
        
# # # # # # # # #         # TITLE ROW (White BG, Blue Text)
# # # # # # # # #         row = t.add_row().cells
# # # # # # # # #         merged = row[0].merge(row[1])
# # # # # # # # #         set_cell_content(merged, proj_div.get_text(strip=True), is_header=False, font_size=14, bold=True, text_color="0d3d62")
        
# # # # # # # # #         # Project Rows
# # # # # # # # #         for row in proj_table.find_all('tr'):
# # # # # # # # #             cells = row.find_all('td')
# # # # # # # # #             if len(cells) == 2:
# # # # # # # # #                 doc_row = t.add_row().cells
                
# # # # # # # # #                 # Left Column: White background, Bold/Italic
# # # # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # # # # # #                 # Right Column: Project Details
# # # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # # # # #     doc.save(docx_file_path)

# # # # # # # # # def set_table_borders(table):
# # # # # # # # #     tbl = table._tbl
# # # # # # # # #     tblPr = tbl.tblPr
# # # # # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # # # # #         border.set(qn('w:val'), 'single')
# # # # # # # # #         border.set(qn('w:sz'), '4')
# # # # # # # # #         border.set(qn('w:space'), '0')
# # # # # # # # #         border.set(qn('w:color'), '000000')
# # # # # # # # #         tblBorders.append(border)
# # # # # # # # #     tblPr.append(tblBorders)

# # # # # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False, text_color=None):
# # # # # # # # #     cell.text = ""
# # # # # # # # #     paragraph = cell.paragraphs[0]
# # # # # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # # # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # # # # #     if is_header:
# # # # # # # # #         set_cell_background(cell, "0d3d62")
# # # # # # # # #         run = paragraph.add_run(text)
# # # # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # # # #         run.font.bold = True
# # # # # # # # #         run.font.size = Pt(font_size)
# # # # # # # # #     else:
# # # # # # # # #         if text:
# # # # # # # # #             run = paragraph.add_run(text)
# # # # # # # # #             run.font.size = Pt(font_size)
# # # # # # # # #             run.font.italic = italic
# # # # # # # # #             run.font.bold = bold
            
# # # # # # # # #             if text_color:
# # # # # # # # #                 r = int(text_color[0:2], 16)
# # # # # # # # #                 g = int(text_color[2:4], 16)
# # # # # # # # #                 b = int(text_color[4:6], 16)
# # # # # # # # #                 run.font.color.rgb = RGBColor(r, g, b)
    
# # # # # # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # # # # def set_cell_background(cell, color_hex):
# # # # # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # # # # def process_html_content(cell, html_element, text_color=None):
# # # # # # # # #     text = html_element.get_text(strip=True)
# # # # # # # # #     if text:
# # # # # # # # #         para = cell.add_paragraph()
# # # # # # # # #         run = para.add_run(text)
# # # # # # # # #         run.font.size = Pt(12)
# # # # # # # # #         if text_color == "white":
# # # # # # # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # # # #             run.font.bold = True

# # # # # # # # # def process_project_content(cell, html_element):
# # # # # # # # #     cell.text = ""
# # # # # # # # #     for element in html_element.children:
# # # # # # # # #         if isinstance(element, Tag):
# # # # # # # # #             # 1. DIVs (Role, Duration)
# # # # # # # # #             if element.name == 'div':
# # # # # # # # #                 para = cell.add_paragraph()
# # # # # # # # #                 para.paragraph_format.space_after = Pt(2)
# # # # # # # # #                 strong = element.find('strong')
# # # # # # # # #                 if strong:
# # # # # # # # #                     label = strong.get_text(strip=True)
# # # # # # # # #                     para.add_run(label + " ").bold = True
# # # # # # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # # # # # #                     if val: para.add_run(val).font.size = Pt(12)
# # # # # # # # #                 else:
# # # # # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # # # # #                     run.font.size = Pt(12)
# # # # # # # # #                     if "Project Description" in run.text: 
# # # # # # # # #                         run.bold = True
# # # # # # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # # # # # #             # 2. OL (Numbered List) - NEW LOGIC
# # # # # # # # #             elif element.name == 'ol':
# # # # # # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # # # # # #                     para = cell.add_paragraph()
# # # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # # # #                     # Manually adding numbers 1. 2. 3.
# # # # # # # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # # # # # # #                     run.font.size = Pt(12)

# # # # # # # # #             # 3. UL (Bullet Points) - Fallback
# # # # # # # # #             elif element.name == 'ul':
# # # # # # # # #                 for li in element.find_all('li'):
# # # # # # # # #                     para = cell.add_paragraph()
# # # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # # # # # #                     run.font.size = Pt(12)

# # # # # # # # #             # 4. Paragraphs
# # # # # # # # #             elif element.name == 'p':
# # # # # # # # #                 para = cell.add_paragraph()
# # # # # # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # # # # # #                 para.add_run(element.get_text(strip=True)).font.size = Pt(12)

# # # # # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # # # # #     try:
# # # # # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # # # # #     except Exception as e:
# # # # # # # # #         print(f"Conversion failed: {e}")


# # # # # # # # from docx import Document
# # # # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # # # from docx.oxml.ns import qn
# # # # # # # # from docx.oxml import OxmlElement
# # # # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # # # import os

# # # # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # # # #         html_content = file.read()
    
# # # # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # # # #     doc = Document()
    
# # # # # # # #     for section in doc.sections:
# # # # # # # #         section.top_margin = Inches(0.5)
# # # # # # # #         section.bottom_margin = Inches(0.5)
# # # # # # # #         section.left_margin = Inches(0.5)
# # # # # # # #         section.right_margin = Inches(0.5)
    
# # # # # # # #     # === 1. ADD LOGO ===
# # # # # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # # # # #     if os.path.exists(logo_path):
# # # # # # # #         try:
# # # # # # # #             logo_para = doc.add_paragraph()
# # # # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # # # #             run = logo_para.add_run()
# # # # # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # # # # #         except Exception: pass
# # # # # # # #     else:
# # # # # # # #         if os.path.exists("download.png"):
# # # # # # # #             try:
# # # # # # # #                 logo_para = doc.add_paragraph()
# # # # # # # #                 run = logo_para.add_run()
# # # # # # # #                 run.add_picture("download.png", width=Inches(2.5))
# # # # # # # #             except: pass

# # # # # # # #     # === 2. MAIN INFO TABLE ===
# # # # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # # # #     if main_table:
# # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # #         t.style = 'Table Grid'
# # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # #         set_table_borders(t)
        
# # # # # # # #         for row in main_table.find_all('tr'):
# # # # # # # #             cells = row.find_all('td')
# # # # # # # #             if len(cells) == 2:
# # # # # # # #                 doc_row = t.add_row().cells
# # # # # # # #                 left_text = cells[0].get_text(strip=True)
                
# # # # # # # #                 # LEFT CELL
# # # # # # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=True)
                
# # # # # # # #                 # RIGHT CELL (CHECK FOR NAME STYLE)
# # # # # # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # # # # # #                 if is_name_cell:
# # # # # # # #                     # Cell BG: Blue, Text: White
# # # # # # # #                     set_cell_content(doc_row[1], "", is_header=True) 
# # # # # # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # # # # # #                 else:
# # # # # # # #                     set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # #                     process_html_content(doc_row[1], cells[1])

# # # # # # # #     # === 3. PROJECTS HEADER & TABLE ===
# # # # # # # #     # Look for the special text-only class
# # # # # # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # # # # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # # # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
    
# # # # # # # #     if proj_table and proj_div:
# # # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # # #         t.style = 'Table Grid'
# # # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # # #         set_table_borders(t)
        
# # # # # # # #         # TITLE ROW: White Cell, but Text has Blue Background (Highlight)
# # # # # # # #         row = t.add_row().cells
# # # # # # # #         merged = row[0].merge(row[1])
        
# # # # # # # #         # We manually build this cell to apply Run Shading
# # # # # # # #         merged.text = ""
# # # # # # # #         paragraph = merged.paragraphs[0]
# # # # # # # #         paragraph.paragraph_format.space_before = Pt(0)
# # # # # # # #         paragraph.paragraph_format.space_after = Pt(0)
# # # # # # # #         paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
# # # # # # # #         run = paragraph.add_run(proj_div.get_text(strip=True))
# # # # # # # #         run.font.bold = True
# # # # # # # #         run.font.size = Pt(14)
# # # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# # # # # # # #         # Apply Blue Background to the TEXT RUN only
# # # # # # # #         set_run_background(run, "0D3D62")
        
# # # # # # # #         # Project Rows
# # # # # # # #         for row in proj_table.find_all('tr'):
# # # # # # # #             cells = row.find_all('td')
# # # # # # # #             if len(cells) == 2:
# # # # # # # #                 doc_row = t.add_row().cells
                
# # # # # # # #                 # Left Column: White background, Bold/Italic
# # # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # # # # #                 # Right Column: Project Details
# # # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # # # #     doc.save(docx_file_path)

# # # # # # # # # === HELPER FUNCTIONS ===

# # # # # # # # def set_run_background(run, color_hex):
# # # # # # # #     """Applies shading (background color) to a specific text run."""
# # # # # # # #     rPr = run._r.get_or_add_rPr()
# # # # # # # #     shd = OxmlElement('w:shd')
# # # # # # # #     shd.set(qn('w:val'), 'clear')
# # # # # # # #     shd.set(qn('w:color'), 'auto')
# # # # # # # #     shd.set(qn('w:fill'), color_hex)
# # # # # # # #     rPr.append(shd)

# # # # # # # # def set_table_borders(table):
# # # # # # # #     tbl = table._tbl
# # # # # # # #     tblPr = tbl.tblPr
# # # # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # # # #         border.set(qn('w:val'), 'single')
# # # # # # # #         border.set(qn('w:sz'), '4')
# # # # # # # #         border.set(qn('w:space'), '0')
# # # # # # # #         border.set(qn('w:color'), '000000')
# # # # # # # #         tblBorders.append(border)
# # # # # # # #     tblPr.append(tblBorders)

# # # # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # # # # #     cell.text = ""
# # # # # # # #     paragraph = cell.paragraphs[0]
# # # # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # # # #     if is_header:
# # # # # # # #         set_cell_background(cell, "0d3d62")
# # # # # # # #         run = paragraph.add_run(text)
# # # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # # #         run.font.bold = True
# # # # # # # #         run.font.size = Pt(font_size)
# # # # # # # #     else:
# # # # # # # #         if text:
# # # # # # # #             run = paragraph.add_run(text)
# # # # # # # #             run.font.size = Pt(font_size)
# # # # # # # #             run.font.italic = italic
# # # # # # # #             run.font.bold = bold
    
# # # # # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # # # def set_cell_background(cell, color_hex):
# # # # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # # # def process_html_content(cell, html_element, text_color=None):
# # # # # # # #     text = html_element.get_text(strip=True)
# # # # # # # #     if text:
# # # # # # # #         para = cell.add_paragraph()
# # # # # # # #         run = para.add_run(text)
# # # # # # # #         run.font.size = Pt(12)
# # # # # # # #         if text_color == "white":
# # # # # # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # # #             run.font.bold = True

# # # # # # # # def process_project_content(cell, html_element):
# # # # # # # #     cell.text = ""
# # # # # # # #     for element in html_element.children:
# # # # # # # #         if isinstance(element, Tag):
# # # # # # # #             # 1. DIVs (Role, Duration)
# # # # # # # #             if element.name == 'div':
# # # # # # # #                 para = cell.add_paragraph()
# # # # # # # #                 para.paragraph_format.space_after = Pt(2)
# # # # # # # #                 strong = element.find('strong')
# # # # # # # #                 if strong:
# # # # # # # #                     label = strong.get_text(strip=True)
# # # # # # # #                     para.add_run(label + " ").bold = True
# # # # # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # # # # #                     if val: para.add_run(val).font.size = Pt(12)
# # # # # # # #                 else:
# # # # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # # # #                     run.font.size = Pt(12)
# # # # # # # #                     if "Project Description" in run.text: 
# # # # # # # #                         run.bold = True
# # # # # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # # # # #             # 2. OL (Numbered List) - Handles 1. 2. 3.
# # # # # # # #             elif element.name == 'ol':
# # # # # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # # # # #                     para = cell.add_paragraph()
# # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # # #                     # Insert number manually
# # # # # # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # # # # # #                     run.font.size = Pt(12)

# # # # # # # #             # 3. UL (Fallback for Bullets)
# # # # # # # #             elif element.name == 'ul':
# # # # # # # #                 for li in element.find_all('li'):
# # # # # # # #                     para = cell.add_paragraph()
# # # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # # # # #                     run.font.size = Pt(12)

# # # # # # # #             # 4. Paragraphs
# # # # # # # #             elif element.name == 'p':
# # # # # # # #                 para = cell.add_paragraph()
# # # # # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # # # # #                 para.add_run(element.get_text(strip=True)).font.size = Pt(12)

# # # # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # # # #     try:
# # # # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # # # #     except Exception as e:
# # # # # # # #         print(f"Conversion failed: {e}")

# # # # # # # #changed

# # # # # # # from docx import Document
# # # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # # from docx.oxml.ns import qn
# # # # # # # from docx.oxml import OxmlElement
# # # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # # import os

# # # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # # #         html_content = file.read()
    
# # # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # # #     doc = Document()
    
# # # # # # #     for section in doc.sections:
# # # # # # #         section.top_margin = Inches(0.5)
# # # # # # #         section.bottom_margin = Inches(0.5)
# # # # # # #         section.left_margin = Inches(0.5)
# # # # # # #         section.right_margin = Inches(0.5)
    
# # # # # # #     # === 1. ADD LOGO ===
# # # # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # # # #     if os.path.exists(logo_path):
# # # # # # #         try:
# # # # # # #             logo_para = doc.add_paragraph()
# # # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # # #             run = logo_para.add_run()
# # # # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # # # #         except Exception: pass
# # # # # # #     elif os.path.exists("download.png"):
# # # # # # #         try:
# # # # # # #             logo_para = doc.add_paragraph()
# # # # # # #             run = logo_para.add_run()
# # # # # # #             run.add_picture("download.png", width=Inches(2.5))
# # # # # # #         except: pass

# # # # # # #     # === 2. MAIN INFO TABLE ===
# # # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # # #     if main_table:
# # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # #         t.style = 'Table Grid'
# # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # #         set_table_borders(t)
        
# # # # # # #         for row in main_table.find_all('tr'):
# # # # # # #             cells = row.find_all('td')
# # # # # # #             if len(cells) == 2:
# # # # # # #                 doc_row = t.add_row().cells
# # # # # # #                 left_text = cells[0].get_text(strip=True)
                
# # # # # # #                 # LEFT CELL (Standard)
# # # # # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=True)
                
# # # # # # #                 # RIGHT CELL (Check for Name Cell Styling)
# # # # # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # # # # #                 if is_name_cell:
# # # # # # #                     # Special: Blue BG, White Text
# # # # # # #                     set_cell_content(doc_row[1], "", is_header=True)
# # # # # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # # # # #                 else:
# # # # # # #                     # Standard: White BG, Black Text
# # # # # # #                     set_cell_content(doc_row[1], "", is_header=False)
# # # # # # #                     process_html_content(doc_row[1], cells[1])

# # # # # # #     # === 3. SEPARATE PROJECT HEADER ===
# # # # # # #     # This creates the visual "gap" by making the header a paragraph, not a table row
# # # # # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # # # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # # # # #     if proj_div:
# # # # # # #         p = doc.add_paragraph()
# # # # # # #         p.paragraph_format.space_before = Pt(12) # This creates the GAP
# # # # # # #         p.paragraph_format.space_after = Pt(4)
        
# # # # # # #         run = p.add_run(proj_div.get_text(strip=True))
# # # # # # #         run.bold = True
# # # # # # #         run.font.size = Pt(14)
# # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# # # # # # #         # Apply Blue Background ONLY to the text run (creating the label effect)
# # # # # # #         set_run_background(run, "0D3D62")

# # # # # # #     # === 4. PROJECTS TABLE ===
# # # # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
# # # # # # #     if proj_table:
# # # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # # #         t.style = 'Table Grid'
# # # # # # #         t.columns[0].width = Inches(1.875)
# # # # # # #         t.columns[1].width = Inches(5.625)
# # # # # # #         set_table_borders(t)
        
# # # # # # #         for row in proj_table.find_all('tr'):
# # # # # # #             cells = row.find_all('td')
# # # # # # #             if len(cells) == 2:
# # # # # # #                 doc_row = t.add_row().cells
                
# # # # # # #                 # Left (Project X)
# # # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # # # #                 # Right (Details)
# # # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # # #     doc.save(docx_file_path)

# # # # # # # # === HELPER FUNCTIONS ===

# # # # # # # def set_run_background(run, color_hex):
# # # # # # #     """Applies shading (background color) to a specific text run."""
# # # # # # #     rPr = run._r.get_or_add_rPr()
# # # # # # #     shd = OxmlElement('w:shd')
# # # # # # #     shd.set(qn('w:val'), 'clear')
# # # # # # #     shd.set(qn('w:color'), 'auto')
# # # # # # #     shd.set(qn('w:fill'), color_hex)
# # # # # # #     rPr.append(shd)

# # # # # # # def set_table_borders(table):
# # # # # # #     tbl = table._tbl
# # # # # # #     tblPr = tbl.tblPr
# # # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # # #         border.set(qn('w:val'), 'single')
# # # # # # #         border.set(qn('w:sz'), '4')
# # # # # # #         border.set(qn('w:space'), '0')
# # # # # # #         border.set(qn('w:color'), '000000')
# # # # # # #         tblBorders.append(border)
# # # # # # #     tblPr.append(tblBorders)

# # # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # # # #     cell.text = ""
# # # # # # #     paragraph = cell.paragraphs[0]
# # # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # # #     if is_header:
# # # # # # #         set_cell_background(cell, "0d3d62")
# # # # # # #         run = paragraph.add_run(text)
# # # # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # #         run.font.bold = True
# # # # # # #         run.font.size = Pt(font_size)
# # # # # # #     else:
# # # # # # #         if text:
# # # # # # #             run = paragraph.add_run(text)
# # # # # # #             run.font.size = Pt(font_size)
# # # # # # #             run.font.italic = italic
# # # # # # #             run.font.bold = bold
    
# # # # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # # def set_cell_background(cell, color_hex):
# # # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # # def process_html_content(cell, html_element, text_color=None):
# # # # # # #     text = html_element.get_text(strip=True)
# # # # # # #     if text:
# # # # # # #         para = cell.add_paragraph()
# # # # # # #         run = para.add_run(text)
# # # # # # #         run.font.size = Pt(12)
# # # # # # #         if text_color == "white":
# # # # # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # # #             run.font.bold = True

# # # # # # # def process_project_content(cell, html_element):
# # # # # # #     cell.text = ""
# # # # # # #     for element in html_element.children:
# # # # # # #         if isinstance(element, Tag):
# # # # # # #             # 1. DIVs (Role, Duration)
# # # # # # #             if element.name == 'div':
# # # # # # #                 para = cell.add_paragraph()
# # # # # # #                 para.paragraph_format.space_after = Pt(2)
# # # # # # #                 strong = element.find('strong')
# # # # # # #                 if strong:
# # # # # # #                     label = strong.get_text(strip=True)
# # # # # # #                     para.add_run(label + " ").bold = True
# # # # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # # # #                     if val: para.add_run(val).font.size = Pt(12)
# # # # # # #                 else:
# # # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # # #                     run.font.size = Pt(12)
# # # # # # #                     if "Project Description" in run.text: 
# # # # # # #                         run.bold = True
# # # # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # # # #             # 2. OL (Numbered List) - Handles 1. 2. 3.
# # # # # # #             elif element.name == 'ol':
# # # # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # # # #                     para = cell.add_paragraph()
# # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # #                     # Manually insert number
# # # # # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # # # # #                     run.font.size = Pt(12)

# # # # # # #             # 3. UL (Fallback for Bullets)
# # # # # # #             elif element.name == 'ul':
# # # # # # #                 for li in element.find_all('li'):
# # # # # # #                     para = cell.add_paragraph()
# # # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # # # #                     run.font.size = Pt(12)

# # # # # # #             # 4. Paragraphs
# # # # # # #             elif element.name == 'p':
# # # # # # #                 para = cell.add_paragraph()
# # # # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # # # #                 para.add_run(element.get_text(strip=True)).font.size = Pt(12)

# # # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # # #     try:
# # # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # # #     except Exception as e:
# # # # # # #         print(f"Conversion failed: {e}")

# # # # # # #color


# # # # # # from docx import Document
# # # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # # from docx.oxml.ns import qn
# # # # # # from docx.oxml import OxmlElement
# # # # # # from bs4 import BeautifulSoup, Tag
# # # # # # import os

# # # # # # # === CONFIGURATION ===
# # # # # # # Updated to Dark Navy Blue (Deepak Singh style)
# # # # # # THEME_COLOR = "1C4587"

# # # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # # #         html_content = file.read()
    
# # # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # # #     doc = Document()
    
# # # # # #     for section in doc.sections:
# # # # # #         section.top_margin = Inches(0.5)
# # # # # #         section.bottom_margin = Inches(0.5)
# # # # # #         section.left_margin = Inches(0.5)
# # # # # #         section.right_margin = Inches(0.5)
    
# # # # # #     # === 1. ADD LOGO ===
# # # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # # #     if os.path.exists(logo_path):
# # # # # #         try:
# # # # # #             logo_para = doc.add_paragraph()
# # # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # # #             run = logo_para.add_run()
# # # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # # #         except Exception: pass
# # # # # #     elif os.path.exists("download.png"):
# # # # # #         try:
# # # # # #             logo_para = doc.add_paragraph()
# # # # # #             run = logo_para.add_run()
# # # # # #             run.add_picture("download.png", width=Inches(2.5))
# # # # # #         except: pass

# # # # # #     # === 2. MAIN INFO TABLE ===
# # # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # # #     if main_table:
# # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # #         t.style = 'Table Grid'
# # # # # #         t.columns[0].width = Inches(1.875)
# # # # # #         t.columns[1].width = Inches(5.625)
# # # # # #         set_table_borders(t)
        
# # # # # #         for row in main_table.find_all('tr'):
# # # # # #             cells = row.find_all('td')
# # # # # #             if len(cells) == 2:
# # # # # #                 doc_row = t.add_row().cells
# # # # # #                 left_text = cells[0].get_text(strip=True)
                
# # # # # #                 # LEFT CELL (Header)
# # # # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=True)
                
# # # # # #                 # RIGHT CELL (Check for Name Cell Styling)
# # # # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # # # #                 if is_name_cell:
# # # # # #                     # Blue BG, White Text for Name
# # # # # #                     set_cell_content(doc_row[1], "", is_header=True)
# # # # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # # # #                 else:
# # # # # #                     # Standard White BG
# # # # # #                     set_cell_content(doc_row[1], "", is_header=False)
# # # # # #                     process_html_content(doc_row[1], cells[1])

# # # # # #     # === 3. SEPARATE PROJECT HEADER ===
# # # # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # # # #     if proj_div:
# # # # # #         p = doc.add_paragraph()
# # # # # #         p.paragraph_format.space_before = Pt(12) 
# # # # # #         p.paragraph_format.space_after = Pt(4)
        
# # # # # #         run = p.add_run(proj_div.get_text(strip=True))
# # # # # #         run.bold = True
# # # # # #         run.font.size = Pt(14)
# # # # # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# # # # # #         # Apply Dark Blue Background to the text run
# # # # # #         set_run_background(run, THEME_COLOR)

# # # # # #     # === 4. PROJECTS TABLE ===
# # # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
# # # # # #     if proj_table:
# # # # # #         t = doc.add_table(rows=0, cols=2)
# # # # # #         t.style = 'Table Grid'
# # # # # #         t.columns[0].width = Inches(1.875)
# # # # # #         t.columns[1].width = Inches(5.625)
# # # # # #         set_table_borders(t)
        
# # # # # #         for row in proj_table.find_all('tr'):
# # # # # #             cells = row.find_all('td')
# # # # # #             if len(cells) == 2:
# # # # # #                 doc_row = t.add_row().cells
                
# # # # # #                 # Left Column: Project Name
# # # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # # #                 # Right Column: Details
# # # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # # #     doc.save(docx_file_path)

# # # # # # # === HELPER FUNCTIONS ===

# # # # # # def set_run_background(run, color_hex):
# # # # # #     """Applies shading (background color) to a specific text run."""
# # # # # #     rPr = run._r.get_or_add_rPr()
# # # # # #     shd = OxmlElement('w:shd')
# # # # # #     shd.set(qn('w:val'), 'clear')
# # # # # #     shd.set(qn('w:color'), 'auto')
# # # # # #     shd.set(qn('w:fill'), color_hex)
# # # # # #     rPr.append(shd)

# # # # # # def set_table_borders(table):
# # # # # #     tbl = table._tbl
# # # # # #     tblPr = tbl.tblPr
# # # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # # #         border = OxmlElement(f'w:{border_name}')
# # # # # #         border.set(qn('w:val'), 'single')
# # # # # #         border.set(qn('w:sz'), '4')
# # # # # #         border.set(qn('w:space'), '0')
# # # # # #         border.set(qn('w:color'), '000000')
# # # # # #         tblBorders.append(border)
# # # # # #     tblPr.append(tblBorders)

# # # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # # #     cell.text = ""
# # # # # #     paragraph = cell.paragraphs[0]
# # # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # # #     if is_header:
# # # # # #         set_cell_background(cell, THEME_COLOR) # Uses darker blue
# # # # # #         run = paragraph.add_run(text)
# # # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # #         run.font.bold = True
# # # # # #         run.font.size = Pt(font_size)
# # # # # #     else:
# # # # # #         if text:
# # # # # #             run = paragraph.add_run(text)
# # # # # #             run.font.size = Pt(font_size)
# # # # # #             run.font.italic = italic
# # # # # #             run.font.bold = bold
    
# # # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # # def set_cell_background(cell, color_hex):
# # # # # #     shading_elm = OxmlElement('w:shd')
# # # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # # def process_html_content(cell, html_element, text_color=None):
# # # # # #     text = html_element.get_text(strip=True)
# # # # # #     if text:
# # # # # #         para = cell.add_paragraph()
# # # # # #         run = para.add_run(text)
# # # # # #         run.font.size = Pt(12)
# # # # # #         if text_color == "white":
# # # # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # # # #             run.font.bold = True

# # # # # # def process_project_content(cell, html_element):
# # # # # #     cell.text = ""
# # # # # #     for element in html_element.children:
# # # # # #         if isinstance(element, Tag):
# # # # # #             # 1. DIVs (Role, Duration)
# # # # # #             if element.name == 'div':
# # # # # #                 para = cell.add_paragraph()
# # # # # #                 para.paragraph_format.space_after = Pt(2)
# # # # # #                 strong = element.find('strong')
# # # # # #                 if strong:
# # # # # #                     label = strong.get_text(strip=True)
# # # # # #                     para.add_run(label + " ").bold = True
# # # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # # #                     if val: para.add_run(val).font.size = Pt(12)
# # # # # #                 else:
# # # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # # #                     run.font.size = Pt(12)
# # # # # #                     if "Project Description" in run.text: 
# # # # # #                         run.bold = True
# # # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # # #             # 2. OL (Numbered List)
# # # # # #             elif element.name == 'ol':
# # # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # # #                     para = cell.add_paragraph()
# # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # # # #                     run.font.size = Pt(12)

# # # # # #             # 3. UL (Fallback)
# # # # # #             elif element.name == 'ul':
# # # # # #                 for li in element.find_all('li'):
# # # # # #                     para = cell.add_paragraph()
# # # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # # #                     run.font.size = Pt(12)

# # # # # #             # 4. Paragraphs
# # # # # #             elif element.name == 'p':
# # # # # #                 para = cell.add_paragraph()
# # # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # # #                 para.add_run(element.get_text(strip=True)).font.size = Pt(12)

# # # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # # #     try:
# # # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # # #     except Exception as e:
# # # # # #         print(f"Conversion failed: {e}")


# # # # # from docx import Document
# # # # # from docx.shared import Inches, Pt, RGBColor
# # # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # # from docx.oxml.ns import qn
# # # # # from docx.oxml import OxmlElement
# # # # # from bs4 import BeautifulSoup, Tag
# # # # # import os

# # # # # # === CONFIGURATION ===
# # # # # # Dark Navy Blue (Deepak Singh style)
# # # # # THEME_COLOR = "0d3d62"

# # # # # def html_to_docx(html_file_path, docx_file_path):
# # # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # # #         html_content = file.read()
    
# # # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # # #     doc = Document()
    
# # # # #     # Set Margins
# # # # #     for section in doc.sections:
# # # # #         section.top_margin = Inches(0.5)
# # # # #         section.bottom_margin = Inches(0.5)
# # # # #         section.left_margin = Inches(0.5)
# # # # #         section.right_margin = Inches(0.5)
    
# # # # #     # === 1. ADD LOGO ===
# # # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # # #     if os.path.exists(logo_path):
# # # # #         try:
# # # # #             logo_para = doc.add_paragraph()
# # # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # # #             run = logo_para.add_run()
# # # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # # #         except Exception: pass
# # # # #     elif os.path.exists("download.png"):
# # # # #         try:
# # # # #             logo_para = doc.add_paragraph()
# # # # #             run = logo_para.add_run()
# # # # #             run.add_picture("download.png", width=Inches(2.5))
# # # # #         except: pass

# # # # #     # === 2. MAIN INFO TABLE ===
# # # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # # #     if main_table:
# # # # #         t = doc.add_table(rows=0, cols=2)
# # # # #         t.style = 'Table Grid'
# # # # #         t.columns[0].width = Inches(1.875)
# # # # #         t.columns[1].width = Inches(5.625)
# # # # #         set_table_borders(t)
        
# # # # #         for row in main_table.find_all('tr'):
# # # # #             cells = row.find_all('td')
# # # # #             if len(cells) == 2:
# # # # #                 doc_row = t.add_row().cells
# # # # #                 left_text = cells[0].get_text(strip=True)
                
# # # # #                 # LEFT CELL (Header: Blue BG, White Text)
# # # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=True)
                
# # # # #                 # RIGHT CELL
# # # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # # #                 if is_name_cell:
# # # # #                     # Special: Blue BG, White Text
# # # # #                     set_cell_content(doc_row[1], "", is_header=True) 
# # # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # # #                 else:
# # # # #                     # Standard: White BG, Black Text
# # # # #                     set_cell_content(doc_row[1], "", is_header=False)
# # # # #                     process_html_content(doc_row[1], cells[1])

# # # # #     # === 3. SEPARATE PROJECT HEADER ===
# # # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # # #     if proj_div:
# # # # #         p = doc.add_paragraph()
# # # # #         p.paragraph_format.space_before = Pt(12)
# # # # #         p.paragraph_format.space_after = Pt(4)
        
# # # # #         run = p.add_run(proj_div.get_text(strip=True))
# # # # #         run.font.name = 'Arial'  # <--- FONT
# # # # #         run.bold = True
# # # # #         run.font.size = Pt(14)
# # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
        
# # # # #         # Apply Blue Background ONLY to the text run
# # # # #         set_run_background(run, THEME_COLOR)

# # # # #     # === 4. PROJECTS TABLE ===
# # # # #     proj_table = soup.find('table', {'class': 'projects-table'})
# # # # #     if proj_table:
# # # # #         t = doc.add_table(rows=0, cols=2)
# # # # #         t.style = 'Table Grid'
# # # # #         t.columns[0].width = Inches(1.875)
# # # # #         t.columns[1].width = Inches(5.625)
# # # # #         set_table_borders(t)
        
# # # # #         for row in proj_table.find_all('tr'):
# # # # #             cells = row.find_all('td')
# # # # #             if len(cells) == 2:
# # # # #                 doc_row = t.add_row().cells
                
# # # # #                 # Left Column: Project Name
# # # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # # #                 # Right Column: Details
# # # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # # #                 process_project_content(doc_row[1], cells[1])
    
# # # # #     doc.save(docx_file_path)

# # # # # # === HELPER FUNCTIONS ===

# # # # # def set_run_background(run, color_hex):
# # # # #     """Applies shading (background color) to a specific text run."""
# # # # #     rPr = run._r.get_or_add_rPr()
# # # # #     shd = OxmlElement('w:shd')
# # # # #     shd.set(qn('w:val'), 'clear')
# # # # #     shd.set(qn('w:color'), 'auto')
# # # # #     shd.set(qn('w:fill'), color_hex)
# # # # #     rPr.append(shd)

# # # # # def set_table_borders(table):
# # # # #     tbl = table._tbl
# # # # #     tblPr = tbl.tblPr
# # # # #     tblBorders = OxmlElement('w:tblBorders')
# # # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # # #         border = OxmlElement(f'w:{border_name}')
# # # # #         border.set(qn('w:val'), 'single')
# # # # #         border.set(qn('w:sz'), '4')
# # # # #         border.set(qn('w:space'), '0')
# # # # #         border.set(qn('w:color'), '000000')
# # # # #         tblBorders.append(border)
# # # # #     tblPr.append(tblBorders)

# # # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # # #     cell.text = ""
# # # # #     paragraph = cell.paragraphs[0]
# # # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # # #     paragraph.paragraph_format.space_after = Pt(0)
# # # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # # #     if is_header:
# # # # #         set_cell_background(cell, THEME_COLOR)
# # # # #         run = paragraph.add_run(text)
# # # # #         run.font.name = 'Arial' # <--- FONT
# # # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # # #         run.font.bold = True
# # # # #         run.font.size = Pt(font_size)
# # # # #     else:
# # # # #         if text:
# # # # #             run = paragraph.add_run(text)
# # # # #             run.font.name = 'Arial' # <--- FONT
# # # # #             run.font.size = Pt(font_size)
# # # # #             run.font.italic = italic
# # # # #             run.font.bold = bold
    
# # # # #     if center_align: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# # # # # def set_cell_background(cell, color_hex):
# # # # #     shading_elm = OxmlElement('w:shd')
# # # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # # def process_html_content(cell, html_element, text_color=None):
# # # # #     text = html_element.get_text(strip=True)
# # # # #     if text:
# # # # #         para = cell.add_paragraph()
# # # # #         run = para.add_run(text)
# # # # #         run.font.name = 'Arial' # <--- FONT
# # # # #         run.font.size = Pt(12)
# # # # #         if text_color == "white":
# # # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # # #             run.font.bold = True

# # # # # def process_project_content(cell, html_element):
# # # # #     cell.text = ""
# # # # #     for element in html_element.children:
# # # # #         if isinstance(element, Tag):
# # # # #             # 1. DIVs (Role, Duration, Tech Stack labels)
# # # # #             if element.name == 'div':
# # # # #                 para = cell.add_paragraph()
# # # # #                 para.paragraph_format.space_after = Pt(2)
# # # # #                 strong = element.find('strong')
# # # # #                 if strong:
# # # # #                     label = strong.get_text(strip=True)
# # # # #                     run = para.add_run(label + " ")
# # # # #                     run.font.name = 'Arial' # <--- FONT
# # # # #                     run.bold = True
                    
# # # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # # #                     if val: 
# # # # #                         r2 = para.add_run(val)
# # # # #                         r2.font.name = 'Arial' # <--- FONT
# # # # #                         r2.font.size = Pt(12)
# # # # #                 else:
# # # # #                     run = para.add_run(element.get_text(strip=True))
# # # # #                     run.font.name = 'Arial' # <--- FONT
# # # # #                     run.font.size = Pt(12)
# # # # #                     if "Project Description" in run.text: 
# # # # #                         run.bold = True
# # # # #                         para.paragraph_format.space_before = Pt(6)

# # # # #             # 2. OL (Numbered List)
# # # # #             elif element.name == 'ol':
# # # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # # #                     para = cell.add_paragraph()
# # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # # #                     run.font.name = 'Arial' # <--- FONT
# # # # #                     run.font.size = Pt(12)

# # # # #             # 3. UL (Fallback for Bullets)
# # # # #             elif element.name == 'ul':
# # # # #                 for li in element.find_all('li'):
# # # # #                     para = cell.add_paragraph()
# # # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # # #                     run.font.name = 'Arial' # <--- FONT
# # # # #                     run.font.size = Pt(12)

# # # # #             # 4. Paragraphs
# # # # #             elif element.name == 'p':
# # # # #                 para = cell.add_paragraph()
# # # # #                 para.paragraph_format.space_before = Pt(3)
# # # # #                 run = para.add_run(element.get_text(strip=True))
# # # # #                 run.font.name = 'Arial' # <--- FONT
# # # # #                 run.font.size = Pt(12)

# # # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # # #     try:
# # # # #         html_to_docx(html_file_path, docx_file_path)
# # # # #     except Exception as e:
# # # # #         print(f"Conversion failed: {e}")


# # # # from docx import Document
# # # # from docx.shared import Inches, Pt, RGBColor
# # # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # # from docx.oxml.ns import qn
# # # # from docx.oxml import OxmlElement
# # # # from bs4 import BeautifulSoup, Tag
# # # # import os

# # # # # === CONFIGURATION ===
# # # # # Dark Navy Blue (Deepak Singh style)
# # # # THEME_COLOR = "#1c4587"

# # # # def html_to_docx(html_file_path, docx_file_path):
# # # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # # #         html_content = file.read()
    
# # # #     soup = BeautifulSoup(html_content, 'html.parser')
# # # #     doc = Document()
    
# # # #     # Set Margins
# # # #     for section in doc.sections:
# # # #         section.top_margin = Inches(0.5)
# # # #         section.bottom_margin = Inches(0.5)
# # # #         section.left_margin = Inches(0.5)
# # # #         section.right_margin = Inches(0.5)
    
# # # #     # === 1. ADD LOGO ===
# # # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # # #     if os.path.exists(logo_path):
# # # #         try:
# # # #             logo_para = doc.add_paragraph()
# # # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # # #             run = logo_para.add_run()
# # # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # # #             logo_para.paragraph_format.space_after = Pt(12)
# # # #         except Exception: pass
# # # #     elif os.path.exists("download.png"):
# # # #         try:
# # # #             logo_para = doc.add_paragraph()
# # # #             run = logo_para.add_run()
# # # #             run.add_picture("download.png", width=Inches(2.5))
# # # #         except: pass

# # # #     # === 2. MAIN INFO TABLE ===
# # # #     main_table = soup.find('table', {'class': 'resume-table'})
# # # #     if main_table:
# # # #         t = doc.add_table(rows=0, cols=2)
# # # #         t.style = 'Table Grid'
# # # #         t.columns[0].width = Inches(1.875)
# # # #         t.columns[1].width = Inches(5.625)
# # # #         set_table_borders(t)
        
# # # #         for row in main_table.find_all('tr'):
# # # #             cells = row.find_all('td')
# # # #             if len(cells) == 2:
# # # #                 doc_row = t.add_row().cells
# # # #                 left_text = cells[0].get_text(strip=True)
                
# # # #                 # LEFT CELL: Blue BG, White Text, Left Aligned, Vertically Centered
# # # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=False)
                
# # # #                 # RIGHT CELL (Check for Name Cell Styling)
# # # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # # #                 if is_name_cell:
# # # #                     # Name Value: Blue BG, White Text, Centered (Deepak Singh style)
# # # #                     set_cell_content(doc_row[1], "", is_header=True, center_align=True) 
# # # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # # #                 else:
# # # #                     # Standard: White BG, Black Text, Left Aligned
# # # #                     set_cell_content(doc_row[1], "", is_header=False, center_align=False)
# # # #                     process_html_content(doc_row[1], cells[1])

# # # #     # === 3. SEPARATE PROJECT HEADER ===
# # # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # # #     if proj_div:
# # # #         p = doc.add_paragraph()
# # # #         p.paragraph_format.space_before = Pt(12)
# # # #         p.paragraph_format.space_after = Pt(4)
        
# # # #         run = p.add_run(proj_div.get_text(strip=True))
# # # #         run.font.name = 'Arial'
# # # #         run.bold = True
# # # #         run.font.size = Pt(14)
# # # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# # # #         # Apply Blue Background ONLY to the text run
# # # #         set_run_background(run, THEME_COLOR)

# # # #     # === 4. PROJECTS TABLE ===
# # # #     proj_table = soup.find('table', {'class': 'projects-table'})
# # # #     if proj_table:
# # # #         t = doc.add_table(rows=0, cols=2)
# # # #         t.style = 'Table Grid'
# # # #         t.columns[0].width = Inches(1.875)
# # # #         t.columns[1].width = Inches(5.625)
# # # #         set_table_borders(t)
        
# # # #         for row in proj_table.find_all('tr'):
# # # #             cells = row.find_all('td')
# # # #             if len(cells) == 2:
# # # #                 doc_row = t.add_row().cells
                
# # # #                 # Left Column: Project Name (White BG, Centered Vertically & Horizontally)
# # # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # # #                 # Right Column: Details
# # # #                 set_cell_content(doc_row[1], "", is_header=False)
# # # #                 process_project_content(doc_row[1], cells[1])
    
# # # #     doc.save(docx_file_path)

# # # # # === HELPER FUNCTIONS ===

# # # # def set_run_background(run, color_hex):
# # # #     """Applies shading (background color) to a specific text run."""
# # # #     rPr = run._r.get_or_add_rPr()
# # # #     shd = OxmlElement('w:shd')
# # # #     shd.set(qn('w:val'), 'clear')
# # # #     shd.set(qn('w:color'), 'auto')
# # # #     shd.set(qn('w:fill'), color_hex)
# # # #     rPr.append(shd)

# # # # def set_table_borders(table):
# # # #     tbl = table._tbl
# # # #     tblPr = tbl.tblPr
# # # #     tblBorders = OxmlElement('w:tblBorders')
# # # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # # #         border = OxmlElement(f'w:{border_name}')
# # # #         border.set(qn('w:val'), 'single')
# # # #         border.set(qn('w:sz'), '4')
# # # #         border.set(qn('w:space'), '0')
# # # #         border.set(qn('w:color'), '000000')
# # # #         tblBorders.append(border)
# # # #     tblPr.append(tblBorders)

# # # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # # #     cell.text = ""
# # # #     paragraph = cell.paragraphs[0]
# # # #     paragraph.paragraph_format.space_before = Pt(0)
# # # #     paragraph.paragraph_format.space_after = Pt(0)
    
# # # #     # Horizontal Alignment
# # # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # # #     # Vertical Alignment (Always Center)
# # # #     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
# # # #     if is_header:
# # # #         set_cell_background(cell, THEME_COLOR)
# # # #         run = paragraph.add_run(text)
# # # #         run.font.name = 'Arial'
# # # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # # #         run.font.bold = True
# # # #         run.font.size = Pt(font_size)
# # # #     else:
# # # #         if text:
# # # #             run = paragraph.add_run(text)
# # # #             run.font.name = 'Arial'
# # # #             run.font.size = Pt(font_size)
# # # #             run.font.italic = italic
# # # #             run.font.bold = bold

# # # # def set_cell_background(cell, color_hex):
# # # #     shading_elm = OxmlElement('w:shd')
# # # #     shading_elm.set(qn('w:fill'), color_hex)
# # # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # # def process_html_content(cell, html_element, text_color=None):
# # # #     text = html_element.get_text(strip=True)
# # # #     if text:
# # # #         para = cell.add_paragraph()
# # # #         run = para.add_run(text)
# # # #         run.font.name = 'Arial'
# # # #         run.font.size = Pt(12)
# # # #         if text_color == "white":
# # # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # # #             run.font.bold = True

# # # # def process_project_content(cell, html_element):
# # # #     cell.text = ""
# # # #     for element in html_element.children:
# # # #         if isinstance(element, Tag):
# # # #             # 1. DIVs (Role, Duration)
# # # #             if element.name == 'div':
# # # #                 para = cell.add_paragraph()
# # # #                 para.paragraph_format.space_after = Pt(2)
# # # #                 strong = element.find('strong')
# # # #                 if strong:
# # # #                     label = strong.get_text(strip=True)
# # # #                     run = para.add_run(label + " ")
# # # #                     run.font.name = 'Arial'
# # # #                     run.bold = True
                    
# # # #                     val = element.get_text().replace(label, '', 1).strip()
# # # #                     if val: 
# # # #                         r2 = para.add_run(val)
# # # #                         r2.font.name = 'Arial'
# # # #                         r2.font.size = Pt(12)
# # # #                 else:
# # # #                     run = para.add_run(element.get_text(strip=True))
# # # #                     run.font.name = 'Arial'
# # # #                     run.font.size = Pt(12)
# # # #                     if "Project Description" in run.text: 
# # # #                         run.bold = True
# # # #                         para.paragraph_format.space_before = Pt(6)

# # # #             # 2. OL (Numbered List)
# # # #             elif element.name == 'ol':
# # # #                 for i, li in enumerate(element.find_all('li'), 1):
# # # #                     para = cell.add_paragraph()
# # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # # #                     run.font.name = 'Arial'
# # # #                     run.font.size = Pt(12)

# # # #             # 3. UL (Fallback for Bullets)
# # # #             elif element.name == 'ul':
# # # #                 for li in element.find_all('li'):
# # # #                     para = cell.add_paragraph()
# # # #                     para.paragraph_format.left_indent = Inches(0.25)
# # # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # # #                     run = para.add_run("• " + li.get_text(strip=True))
# # # #                     run.font.name = 'Arial'
# # # #                     run.font.size = Pt(12)

# # # #             # 4. Paragraphs
# # # #             elif element.name == 'p':
# # # #                 para = cell.add_paragraph()
# # # #                 para.paragraph_format.space_before = Pt(3)
# # # #                 run = para.add_run(element.get_text(strip=True))
# # # #                 run.font.name = 'Arial'
# # # #                 run.font.size = Pt(12)

# # # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # # #     try:
# # # #         html_to_docx(html_file_path, docx_file_path)
# # # #     except Exception as e:
# # # #         print(f"Conversion failed: {e}")


# # # from docx import Document
# # # from docx.shared import Inches, Pt, RGBColor
# # # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # # from docx.oxml.ns import qn
# # # from docx.oxml import OxmlElement
# # # from bs4 import BeautifulSoup, Tag
# # # import os

# # # # === CONFIGURATION ===
# # # # Dark Navy Blue (Deepak Singh style)
# # # THEME_COLOR = "0d3d62"

# # # def html_to_docx(html_file_path, docx_file_path):
# # #     with open(html_file_path, 'r', encoding='utf-8') as file:
# # #         html_content = file.read()
    
# # #     soup = BeautifulSoup(html_content, 'html.parser')
# # #     doc = Document()
    
# # #     # Set Margins
# # #     for section in doc.sections:
# # #         section.top_margin = Inches(0.5)
# # #         section.bottom_margin = Inches(0.5)
# # #         section.left_margin = Inches(0.5)
# # #         section.right_margin = Inches(0.5)
    
# # #     # === 1. ADD LOGO ===
# # #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# # #     if os.path.exists(logo_path):
# # #         try:
# # #             logo_para = doc.add_paragraph()
# # #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# # #             run = logo_para.add_run()
# # #             run.add_picture(logo_path, width=Inches(2.5)) 
# # #             logo_para.paragraph_format.space_after = Pt(12)
# # #         except Exception: pass
# # #     elif os.path.exists("download.png"):
# # #         try:
# # #             logo_para = doc.add_paragraph()
# # #             run = logo_para.add_run()
# # #             run.add_picture("download.png", width=Inches(2.5))
# # #         except: pass

# # #     # === 2. MAIN INFO TABLE ===
# # #     main_table = soup.find('table', {'class': 'resume-table'})
# # #     if main_table:
# # #         t = doc.add_table(rows=0, cols=2)
# # #         t.style = 'Table Grid'
# # #         t.columns[0].width = Inches(1.875)
# # #         t.columns[1].width = Inches(5.625)
# # #         set_table_borders(t)
        
# # #         for row in main_table.find_all('tr'):
# # #             cells = row.find_all('td')
# # #             if len(cells) == 2:
# # #                 doc_row = t.add_row().cells
# # #                 left_text = cells[0].get_text(strip=True)
                
# # #                 # LEFT CELL: Blue BG, White Text, Left Aligned, Vertically Centered
# # #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=False)
                
# # #                 # RIGHT CELL: Always Centered now (for Skills, Experience, Languages, etc.)
# # #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# # #                 if is_name_cell:
# # #                     # Name Value: Blue BG, White Text, Centered
# # #                     set_cell_content(doc_row[1], "", is_header=True, center_align=True) 
# # #                     process_html_content(doc_row[1], cells[1], text_color="white")
# # #                 else:
# # #                     # Standard Values: White BG, Black Text, CENTERED (Fixes "all not in centre")
# # #                     set_cell_content(doc_row[1], "", is_header=False, center_align=True)
# # #                     process_html_content(doc_row[1], cells[1])

# # #     # === 3. SEPARATE PROJECT HEADER ===
# # #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# # #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# # #     if proj_div:
# # #         p = doc.add_paragraph()
# # #         p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # <--- CHANGED: Left Align for Project Header
# # #         p.paragraph_format.space_before = Pt(12)
# # #         p.paragraph_format.space_after = Pt(4)
        
# # #         run = p.add_run(proj_div.get_text(strip=True))
# # #         run.font.name = 'Arial'
# # #         run.bold = True
# # #         run.font.size = Pt(14)
# # #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# # #         # Apply Blue Background ONLY to the text run
# # #         set_run_background(run, THEME_COLOR)

# # #     # === 4. PROJECTS TABLE ===
# # #     proj_table = soup.find('table', {'class': 'projects-table'})
# # #     if proj_table:
# # #         t = doc.add_table(rows=0, cols=2)
# # #         t.style = 'Table Grid'
# # #         t.columns[0].width = Inches(1.875)
# # #         t.columns[1].width = Inches(5.625)
# # #         set_table_borders(t)
        
# # #         for row in proj_table.find_all('tr'):
# # #             cells = row.find_all('td')
# # #             if len(cells) == 2:
# # #                 doc_row = t.add_row().cells
                
# # #                 # Left Column: Project Name (White BG, Centered Vertically & Horizontally)
# # #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# # #                 # Right Column: Details
# # #                 set_cell_content(doc_row[1], "", is_header=False)
# # #                 process_project_content(doc_row[1], cells[1])
    
# # #     doc.save(docx_file_path)

# # # # === HELPER FUNCTIONS ===

# # # def set_run_background(run, color_hex):
# # #     """Applies shading (background color) to a specific text run."""
# # #     rPr = run._r.get_or_add_rPr()
# # #     shd = OxmlElement('w:shd')
# # #     shd.set(qn('w:val'), 'clear')
# # #     shd.set(qn('w:color'), 'auto')
# # #     shd.set(qn('w:fill'), color_hex)
# # #     rPr.append(shd)

# # # def set_table_borders(table):
# # #     tbl = table._tbl
# # #     tblPr = tbl.tblPr
# # #     tblBorders = OxmlElement('w:tblBorders')
# # #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# # #         border = OxmlElement(f'w:{border_name}')
# # #         border.set(qn('w:val'), 'single')
# # #         border.set(qn('w:sz'), '4')
# # #         border.set(qn('w:space'), '0')
# # #         border.set(qn('w:color'), '000000')
# # #         tblBorders.append(border)
# # #     tblPr.append(tblBorders)

# # # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# # #     cell.text = ""
# # #     paragraph = cell.paragraphs[0]
# # #     paragraph.paragraph_format.space_before = Pt(0)
# # #     paragraph.paragraph_format.space_after = Pt(0)
    
# # #     # Horizontal Alignment (Controls Left/Center of text)
# # #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# # #     # Vertical Alignment (Always Center)
# # #     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
# # #     if is_header:
# # #         set_cell_background(cell, THEME_COLOR)
# # #         run = paragraph.add_run(text)
# # #         run.font.name = 'Arial'
# # #         run.font.color.rgb = RGBColor(255, 255, 255)
# # #         run.font.bold = True
# # #         run.font.size = Pt(font_size)
# # #     else:
# # #         if text:
# # #             run = paragraph.add_run(text)
# # #             run.font.name = 'Arial'
# # #             run.font.size = Pt(font_size)
# # #             run.font.italic = italic
# # #             run.font.bold = bold

# # # def set_cell_background(cell, color_hex):
# # #     shading_elm = OxmlElement('w:shd')
# # #     shading_elm.set(qn('w:fill'), color_hex)
# # #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # # def process_html_content(cell, html_element, text_color=None):
# # #     text = html_element.get_text(strip=True)
# # #     if text:
# # #         # Use EXISTING paragraph to keep alignment
# # #         para = cell.paragraphs[0]
        
# # #         run = para.add_run(text)
# # #         run.font.name = 'Arial'
# # #         run.font.size = Pt(12)
# # #         if text_color == "white":
# # #             run.font.color.rgb = RGBColor(255, 255, 255)
# # #             run.font.bold = True

# # # def process_project_content(cell, html_element):
# # #     cell.text = "" 
# # #     first_element = True
    
# # #     for element in html_element.children:
# # #         if isinstance(element, Tag):
# # #             if element.name == 'div':
# # #                 if first_element:
# # #                     para = cell.paragraphs[0]
# # #                     first_element = False
# # #                 else:
# # #                     para = cell.add_paragraph()
                
# # #                 para.paragraph_format.space_after = Pt(2)
# # #                 strong = element.find('strong')
# # #                 if strong:
# # #                     label = strong.get_text(strip=True)
# # #                     run = para.add_run(label + " ")
# # #                     run.font.name = 'Arial'
# # #                     run.bold = True
                    
# # #                     val = element.get_text().replace(label, '', 1).strip()
# # #                     if val: 
# # #                         r2 = para.add_run(val)
# # #                         r2.font.name = 'Arial'
# # #                         r2.font.size = Pt(12)
# # #                 else:
# # #                     run = para.add_run(element.get_text(strip=True))
# # #                     run.font.name = 'Arial'
# # #                     run.font.size = Pt(12)
# # #                     if "Project Description" in run.text: 
# # #                         run.bold = True
# # #                         para.paragraph_format.space_before = Pt(6)

# # #             elif element.name == 'ol':
# # #                 for i, li in enumerate(element.find_all('li'), 1):
# # #                     if first_element:
# # #                         para = cell.paragraphs[0]
# # #                         first_element = False
# # #                     else:
# # #                         para = cell.add_paragraph()
                        
# # #                     para.paragraph_format.left_indent = Inches(0.25)
# # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# # #                     run.font.name = 'Arial'
# # #                     run.font.size = Pt(12)

# # #             elif element.name == 'ul':
# # #                 for li in element.find_all('li'):
# # #                     if first_element:
# # #                         para = cell.paragraphs[0]
# # #                         first_element = False
# # #                     else:
# # #                         para = cell.add_paragraph()

# # #                     para.paragraph_format.left_indent = Inches(0.25)
# # #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# # #                     run = para.add_run("• " + li.get_text(strip=True))
# # #                     run.font.name = 'Arial'
# # #                     run.font.size = Pt(12)

# # #             elif element.name == 'p':
# # #                 if first_element:
# # #                     para = cell.paragraphs[0]
# # #                     first_element = False
# # #                 else:
# # #                     para = cell.add_paragraph()
                    
# # #                 para.paragraph_format.space_before = Pt(3)
# # #                 run = para.add_run(element.get_text(strip=True))
# # #                 run.font.name = 'Arial'
# # #                 run.font.size = Pt(12)

# # # def convert_salesforce_resume(html_file_path, docx_file_path):
# # #     try:
# # #         html_to_docx(html_file_path, docx_file_path)
# # #     except Exception as e:
# # #         print(f"Conversion failed: {e}")
# # from docx import Document
# # from docx.shared import Inches, Pt, RGBColor
# # from docx.enum.text import WD_ALIGN_PARAGRAPH
# # from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# # from docx.oxml.ns import qn
# # from docx.oxml import OxmlElement
# # from bs4 import BeautifulSoup, Tag
# # import os

# # # === CONFIGURATION ===
# # # Dark Navy Blue (Deepak Singh style)
# # THEME_COLOR = "#1c4587"

# # def html_to_docx(html_file_path, docx_file_path):
# #     with open(html_file_path, 'r', encoding='utf-8') as file:
# #         html_content = file.read()
    
# #     soup = BeautifulSoup(html_content, 'html.parser')
# #     doc = Document()
    
# #     # Set Margins
# #     for section in doc.sections:
# #         section.top_margin = Inches(0.5)
# #         section.bottom_margin = Inches(0.5)
# #         section.left_margin = Inches(0.5)
# #         section.right_margin = Inches(0.5)
    
# #     # === 1. ADD LOGO ===
# #     logo_path = "/home/ca/Projects/resume_converter/download.png"
# #     if os.path.exists(logo_path):
# #         try:
# #             logo_para = doc.add_paragraph()
# #             logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
# #             run = logo_para.add_run()
# #             run.add_picture(logo_path, width=Inches(2.5)) 
# #             logo_para.paragraph_format.space_after = Pt(12)
# #         except Exception: pass
# #     elif os.path.exists("download.png"):
# #         try:
# #             logo_para = doc.add_paragraph()
# #             run = logo_para.add_run()
# #             run.add_picture("download.png", width=Inches(2.5))
# #         except: pass

# #     # === 2. MAIN INFO TABLE ===
# #     main_table = soup.find('table', {'class': 'resume-table'})
# #     if main_table:
# #         t = doc.add_table(rows=0, cols=2)
# #         t.style = 'Table Grid'
# #         t.columns[0].width = Inches(1.875)
# #         t.columns[1].width = Inches(5.625)
# #         set_table_borders(t)
        
# #         for row in main_table.find_all('tr'):
# #             cells = row.find_all('td')
# #             if len(cells) == 2:
# #                 doc_row = t.add_row().cells
# #                 left_text = cells[0].get_text(strip=True)
                
# #                 # LEFT CELL: Blue BG, White Text, Left Aligned, Vertically Centered
# #                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=False)
                
# #                 # RIGHT CELL Logic:
# #                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
# #                 if is_name_cell:
# #                     # Name Value: Blue BG, White Text, Centered
# #                     set_cell_content(doc_row[1], "", is_header=True, center_align=True) 
# #                     process_html_content(doc_row[1], cells[1], text_color="white")
# #                 else:
# #                     # Standard Values: White BG, Black Text, ALWAYS CENTERED (Matches Reference)
# #                     set_cell_content(doc_row[1], "", is_header=False, center_align=True)
# #                     process_html_content(doc_row[1], cells[1])

# #     # === 3. SEPARATE PROJECT HEADER ===
# #     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
# #     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
# #     if proj_div:
# #         p = doc.add_paragraph()
# #         p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left Align for Project Header
# #         p.paragraph_format.space_before = Pt(12)
# #         p.paragraph_format.space_after = Pt(4)
        
# #         run = p.add_run(proj_div.get_text(strip=True))
# #         run.font.name = 'Arial'
# #         run.bold = True
# #         run.font.size = Pt(14)
# #         run.font.color.rgb = RGBColor(255, 255, 255) # White Text
        
# #         # Apply Blue Background ONLY to the text run
# #         set_run_background(run, THEME_COLOR)

# #     # === 4. PROJECTS TABLE ===
# #     proj_table = soup.find('table', {'class': 'projects-table'})
# #     if proj_table:
# #         t = doc.add_table(rows=0, cols=2)
# #         t.style = 'Table Grid'
# #         t.columns[0].width = Inches(1.875)
# #         t.columns[1].width = Inches(5.625)
# #         set_table_borders(t)
        
# #         for row in proj_table.find_all('tr'):
# #             cells = row.find_all('td')
# #             if len(cells) == 2:
# #                 doc_row = t.add_row().cells
                
# #                 # Left Column: Project Name (White BG, Centered Vertically & Horizontally)
# #                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
# #                 # Right Column: Details
# #                 set_cell_content(doc_row[1], "", is_header=False)
# #                 process_project_content(doc_row[1], cells[1])
    
# #     doc.save(docx_file_path)

# # # === HELPER FUNCTIONS ===

# # def set_run_background(run, color_hex):
# #     """Applies shading (background color) to a specific text run."""
# #     rPr = run._r.get_or_add_rPr()
# #     shd = OxmlElement('w:shd')
# #     shd.set(qn('w:val'), 'clear')
# #     shd.set(qn('w:color'), 'auto')
# #     shd.set(qn('w:fill'), color_hex)
# #     rPr.append(shd)

# # def set_table_borders(table):
# #     tbl = table._tbl
# #     tblPr = tbl.tblPr
# #     tblBorders = OxmlElement('w:tblBorders')
# #     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
# #         border = OxmlElement(f'w:{border_name}')
# #         border.set(qn('w:val'), 'single')
# #         border.set(qn('w:sz'), '4')
# #         border.set(qn('w:space'), '0')
# #         border.set(qn('w:color'), '000000')
# #         tblBorders.append(border)
# #     tblPr.append(tblBorders)

# # def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
# #     cell.text = ""
# #     paragraph = cell.paragraphs[0]
    
# #     # === CHANGED: Increased Padding to Pt(8) to Match Height ===
# #     paragraph.paragraph_format.space_before = Pt(8)
# #     paragraph.paragraph_format.space_after = Pt(8)
    
# #     # Horizontal Alignment
# #     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    
# #     # Vertical Alignment
# #     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
# #     if is_header:
# #         set_cell_background(cell, THEME_COLOR)
# #         run = paragraph.add_run(text)
# #         run.font.name = 'Arial'
# #         run.font.color.rgb = RGBColor(255, 255, 255)
# #         run.font.bold = True
# #         run.font.size = Pt(font_size)
# #     else:
# #         if text:
# #             run = paragraph.add_run(text)
# #             run.font.name = 'Arial'
# #             run.font.size = Pt(font_size)
# #             run.font.italic = italic
# #             run.font.bold = bold

# # def set_cell_background(cell, color_hex):
# #     shading_elm = OxmlElement('w:shd')
# #     shading_elm.set(qn('w:fill'), color_hex)
# #     cell._tc.get_or_add_tcPr().append(shading_elm)

# # def process_html_content(cell, html_element, text_color=None):
# #     text = html_element.get_text(strip=True)
# #     if text:
# #         # Use EXISTING paragraph
# #         para = cell.paragraphs[0]
        
# #         # === CHANGED: Ensure padding is applied here too ===
# #         para.paragraph_format.space_before = Pt(8)
# #         para.paragraph_format.space_after = Pt(8)
        
# #         run = para.add_run(text)
# #         run.font.name = 'Arial'
# #         run.font.size = Pt(12)
# #         if text_color == "white":
# #             run.font.color.rgb = RGBColor(255, 255, 255)
# #             run.font.bold = True

# # def process_project_content(cell, html_element):
# #     cell.text = "" 
# #     first_element = True
    
# #     for element in html_element.children:
# #         if isinstance(element, Tag):
# #             if element.name == 'div':
# #                 if first_element:
# #                     para = cell.paragraphs[0]
# #                     first_element = False
# #                 else:
# #                     para = cell.add_paragraph()
                
# #                 para.paragraph_format.space_after = Pt(2)
# #                 strong = element.find('strong')
# #                 if strong:
# #                     label = strong.get_text(strip=True)
# #                     run = para.add_run(label + " ")
# #                     run.font.name = 'Arial'
# #                     run.bold = True
                    
# #                     val = element.get_text().replace(label, '', 1).strip()
# #                     if val: 
# #                         r2 = para.add_run(val)
# #                         r2.font.name = 'Arial'
# #                         r2.font.size = Pt(12)
# #                 else:
# #                     run = para.add_run(element.get_text(strip=True))
# #                     run.font.name = 'Arial'
# #                     run.font.size = Pt(12)
# #                     if "Project Description" in run.text: 
# #                         run.bold = True
# #                         para.paragraph_format.space_before = Pt(6)

# #             elif element.name == 'ol':
# #                 for i, li in enumerate(element.find_all('li'), 1):
# #                     if first_element:
# #                         para = cell.paragraphs[0]
# #                         first_element = False
# #                     else:
# #                         para = cell.add_paragraph()
                        
# #                     para.paragraph_format.left_indent = Inches(0.25)
# #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# #                     run = para.add_run(f"{i}. " + li.get_text(strip=True))
# #                     run.font.name = 'Arial'
# #                     run.font.size = Pt(12)

# #             elif element.name == 'ul':
# #                 for li in element.find_all('li'):
# #                     if first_element:
# #                         para = cell.paragraphs[0]
# #                         first_element = False
# #                     else:
# #                         para = cell.add_paragraph()

# #                     para.paragraph_format.left_indent = Inches(0.25)
# #                     para.paragraph_format.first_line_indent = Inches(-0.25)
# #                     run = para.add_run("• " + li.get_text(strip=True))
# #                     run.font.name = 'Arial'
# #                     run.font.size = Pt(12)

# #             elif element.name == 'p':
# #                 if first_element:
# #                     para = cell.paragraphs[0]
# #                     first_element = False
# #                 else:
# #                     para = cell.add_paragraph()
                    
# #                 para.paragraph_format.space_before = Pt(3)
# #                 run = para.add_run(element.get_text(strip=True))
# #                 run.font.name = 'Arial'
# #                 run.font.size = Pt(12)

# # def convert_salesforce_resume(html_file_path, docx_file_path):
# #     try:
# #         html_to_docx(html_file_path, docx_file_path)
# #     except Exception as e:
# #         print(f"Conversion failed: {e}")


# #badge


# from docx import Document
# from docx.shared import Inches, Pt, RGBColor
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from bs4 import BeautifulSoup, Tag
# import os

# # === CONFIGURATION ===
# # Word shading fill expects hex without '#'
# THEME_COLOR = "1C4587"  # Dark Navy Blue

# # PATH to your badges folder
# BADGE_DIR = "/home/ca/Projects/resume_converter/src/certificates_badges/"

# # MAPPING: Keyword in Resume -> Image Filename
# CERT_BADGE_MAP = {
#     "administrator": "admin_badge.png",
#     "platform developer": "pd1_badge.png",
#     "app builder": "app_builder_badge.png",
#     "consultant": "consultant_badge.png",
#     "omnistudio": "omni.png"  # <--- Added OmniStudio mapping
# }

# def html_to_docx(html_file_path, docx_file_path):
#     with open(html_file_path, 'r', encoding='utf-8') as file:
#         html_content = file.read()
    
#     soup = BeautifulSoup(html_content, 'html.parser')
#     doc = Document()
    
#     # Set Margins
#     for section in doc.sections:
#         section.top_margin = Inches(0.5)
#         section.bottom_margin = Inches(0.5)
#         section.left_margin = Inches(0.5)
#         section.right_margin = Inches(0.5)
    
#     # === 1. EXTRACT CERTIFICATIONS TEXT ===
#     cert_text = ""
#     rows = soup.find_all('tr')
#     for row in rows:
#         cells = row.find_all('td')
#         if len(cells) > 0 and "Certifications" in cells[0].get_text():
#             if len(cells) > 1:
#                 cert_text = cells[1].get_text(strip=True).lower()
#             break

#     # === 2. HEADER: BADGES (Left) + LOGO (Right) ===
#     header_table = doc.add_table(rows=1, cols=2)
#     header_table.autofit = False
#     header_table.columns[0].width = Inches(4.0)
#     header_table.columns[1].width = Inches(3.5)
    
#     # -- A. Add Badges to Left Cell --
#     left_cell = header_table.cell(0, 0)
#     left_para = left_cell.paragraphs[0]
#     left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
#     badges_added = False
#     for keyword, image_file in CERT_BADGE_MAP.items():
#         if keyword in cert_text:
#             full_badge_path = os.path.join(BADGE_DIR, image_file)
            
#             if os.path.exists(full_badge_path):
#                 run = left_para.add_run()
#                 run.add_picture(full_badge_path, width=Inches(0.85))
#                 run.add_text("  ") # Spacer
#                 badges_added = True
#             else:
#                 print(f"⚠️ Warning: Badge image not found at {full_badge_path}")
    
#     # -- B. Add Company Logo to Right Cell --
#     right_cell = header_table.cell(0, 1)
#     right_para = right_cell.paragraphs[0]
#     right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
#     logo_path = "/home/ca/Projects/resume_converter/download.png"
#     if os.path.exists(logo_path):
#         run = right_para.add_run()
#         run.add_picture(logo_path, width=Inches(2.5))
#     elif os.path.exists("download.png"):
#         run = right_para.add_run()
#         run.add_picture("download.png", width=Inches(2.5))

#     doc.add_paragraph().paragraph_format.space_after = Pt(6)

#     # === 3. MAIN INFO TABLE ===
#     main_table = soup.find('table', {'class': 'resume-table'})
#     if main_table:
#         t = doc.add_table(rows=0, cols=2)
#         t.style = 'Table Grid'
#         t.columns[0].width = Inches(1.875)
#         t.columns[1].width = Inches(5.625)
#         set_table_borders(t)
        
#         for row in main_table.find_all('tr'):
#             cells = row.find_all('td')
#             if len(cells) == 2:
#                 doc_row = t.add_row().cells
#                 left_text = cells[0].get_text(strip=True)
                
#                 # LEFT CELL
#                 set_cell_content(doc_row[0], left_text, is_header=True, center_align=False)
                
#                 # RIGHT CELL Logic
#                 is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
#                 if is_name_cell:
#                     # Name: Blue BG, Centered
#                     set_cell_content(doc_row[1], "", is_header=True, center_align=True) 
#                     process_html_content(doc_row[1], cells[1], text_color="white")
#                 else:
#                     # Standard: White BG, Centered
#                     set_cell_content(doc_row[1], "", is_header=False, center_align=True)
#                     process_html_content(doc_row[1], cells[1])

#     # === 4. SEPARATE PROJECT HEADER ===
#     proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
#     if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
#     if proj_div:
#         p = doc.add_paragraph()
#         p.alignment = WD_ALIGN_PARAGRAPH.LEFT
#         p.paragraph_format.space_before = Pt(12)
#         p.paragraph_format.space_after = Pt(4)
        
#         run = p.add_run(proj_div.get_text(strip=True))
#         run.font.name = 'Arial'
#         run.bold = True
#         run.font.size = Pt(14)
#         run.font.color.rgb = RGBColor(255, 255, 255)
#         set_run_background(run, THEME_COLOR)

#     # === 5. PROJECTS TABLE ===
#     proj_table = soup.find('table', {'class': 'projects-table'})
#     if proj_table:
#         t = doc.add_table(rows=0, cols=2)
#         t.style = 'Table Grid'
#         t.columns[0].width = Inches(1.875)
#         t.columns[1].width = Inches(5.625)
#         set_table_borders(t)
        
#         for row in proj_table.find_all('tr'):
#             cells = row.find_all('td')
#             if len(cells) == 2:
#                 doc_row = t.add_row().cells
                
#                 # Left Column: Project Name
#                 set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
#                 # Right Column: Details
#                 set_cell_content(doc_row[1], "", is_header=False)
#                 process_project_content(doc_row[1], cells[1])
    
#     doc.save(docx_file_path)

# # === HELPER FUNCTIONS ===

# def set_run_background(run, color_hex):
#     """Applies shading (background color) to a specific text run."""
#     rPr = run._r.get_or_add_rPr()
#     fill = (color_hex or "").replace("#", "")
#     shd = OxmlElement('w:shd')
#     shd.set(qn('w:val'), 'clear')
#     shd.set(qn('w:color'), 'auto')
#     shd.set(qn('w:fill'), fill)
#     rPr.append(shd)

# def set_table_borders(table):
#     tbl = table._tbl
#     tblPr = tbl.tblPr
#     tblBorders = OxmlElement('w:tblBorders')
#     for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
#         border = OxmlElement(f'w:{border_name}')
#         border.set(qn('w:val'), 'single')
#         border.set(qn('w:sz'), '4')
#         border.set(qn('w:space'), '0')
#         border.set(qn('w:color'), '000000')
#         tblBorders.append(border)
#     tblPr.append(tblBorders)

# def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
#     cell.text = ""
#     paragraph = cell.paragraphs[0]
    
#     # Increase Padding for height (Deepak Singh style)
#     paragraph.paragraph_format.space_before = Pt(8)
#     paragraph.paragraph_format.space_after = Pt(8)
    
#     paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
#     cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
#     if is_header:
#         set_cell_background(cell, THEME_COLOR)
#         run = paragraph.add_run(text)
#         run.font.name = 'Arial'
#         run.font.color.rgb = RGBColor(255, 255, 255)
#         run.font.bold = True
#         run.font.size = Pt(font_size)
#     else:
#         if text:
#             run = paragraph.add_run(text)
#             run.font.name = 'Arial'
#             run.font.size = Pt(font_size)
#             run.font.italic = italic
#             run.font.bold = bold

# def set_cell_background(cell, color_hex):
#     shading_elm = OxmlElement('w:shd')
#     fill = (color_hex or "").replace("#", "")
#     shading_elm.set(qn('w:val'), 'clear')
#     shading_elm.set(qn('w:color'), 'auto')
#     shading_elm.set(qn('w:fill'), fill)
#     cell._tc.get_or_add_tcPr().append(shading_elm)

# def process_html_content(cell, html_element, text_color=None):
#     text = html_element.get_text(strip=True)
#     if text:
#         para = cell.paragraphs[0]
#         para.paragraph_format.space_before = Pt(8)
#         para.paragraph_format.space_after = Pt(8)
        
#         run = para.add_run(text)
#         run.font.name = 'Arial'
#         run.font.size = Pt(12)
#         if text_color == "white":
#             run.font.color.rgb = RGBColor(255, 255, 255)
#             run.font.bold = True

# def process_project_content(cell, html_element):
#     cell.text = "" 
#     first_element = True
    
#     for element in html_element.children:
#         if isinstance(element, Tag):
#             if element.name == 'div':
#                 if first_element:
#                     para = cell.paragraphs[0]
#                     first_element = False
#                 else:
#                     para = cell.add_paragraph()
                
#                 para.paragraph_format.space_after = Pt(2)
#                 strong = element.find('strong')
#                 if strong:
#                     label = strong.get_text(strip=True)
#                     run = para.add_run(label + " ")
#                     run.font.name = 'Arial'
#                     run.bold = True
                    
#                     val = element.get_text().replace(label, '', 1).strip()
#                     if val: 
#                         r2 = para.add_run(val)
#                         r2.font.name = 'Arial'
#                         r2.font.size = Pt(12)
#                 else:
                    
#                     text_val = element.get_text(strip=True)
#                     run = para.add_run(text_val)
#                     run.font.name = 'Arial'
#                     run.font.size = Pt(12)

#                     # Make heading bold + spacing
#                     if "Project Description" in text_val:
#                         run.bold = True
#                         para.paragraph_format.space_before = Pt(8)
#                         para.paragraph_format.space_after = Pt(2)


#             elif element.name == 'ol':
#                 for i, li in enumerate(element.find_all('li'), 1):
#                     if first_element:
#                         para = cell.paragraphs[0]
#                         first_element = False
#                     else:
#                         para = cell.add_paragraph()
                        
#                     # Avoid negative indents: some online DOCX viewers render them
#                     # outside table cells (numbers appear far-left).
#                     para.paragraph_format.left_indent = Inches(0.35)
#                     para.paragraph_format.first_line_indent = Inches(0)
#                     try:
#                         para.paragraph_format.tab_stops.clear_all()
#                         para.paragraph_format.tab_stops.add_tab_stop(Inches(0.55))
#                     except Exception:
#                         pass

#                     run = para.add_run(f"{i}.\t" + li.get_text(strip=True))
#                     run.font.name = 'Arial'
#                     run.font.size = Pt(12)

#             elif element.name == 'ul':
#                 for li in element.find_all('li'):
#                     if first_element:
#                         para = cell.paragraphs[0]
#                         first_element = False
#                     else:
#                         para = cell.add_paragraph()

#                     para.paragraph_format.left_indent = Inches(0.35)
#                     para.paragraph_format.first_line_indent = Inches(0)
#                     try:
#                         para.paragraph_format.tab_stops.clear_all()
#                         para.paragraph_format.tab_stops.add_tab_stop(Inches(0.55))
#                     except Exception:
#                         pass

#                     run = para.add_run("•\t" + li.get_text(strip=True))
#                     run.font.name = 'Arial'
#                     run.font.size = Pt(12)

#             elif element.name == 'p':
#                 if first_element:
#                     para = cell.paragraphs[0]
#                     first_element = False
#                 else:
#                     para = cell.add_paragraph()
                    
#                 para.paragraph_format.space_before = Pt(3)
#                 run = para.add_run(element.get_text(strip=True))
#                 run.font.name = 'Arial'
#                 run.font.size = Pt(12)

# def convert_salesforce_resume(html_file_path, docx_file_path):
#     try:
#         html_to_docx(html_file_path, docx_file_path)
#     except Exception as e:
#         print(f"Conversion failed: {e}")


from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, Tag
import os

# === CONFIGURATION ===
# Word shading fill expects hex without '#'
THEME_COLOR = "1C4587"  # Dark Navy Blue

# PATH to your badges folder
BADGE_DIR = "/home/ca/Projects/resume_converter/src/certificates_badges/"

# MAPPING: Keyword in Resume -> Image Filename
CERT_BADGE_MAP = {
    "administrator": "admin_badge.png",
    "platform developer": "pd1_badge.png",
    "app builder": "app_builder_badge.png",
    "consultant": "consultant_badge.png",
    "omnistudio": "omni.png" 
}

def html_to_docx(html_file_path, docx_file_path):
    with open(html_file_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # === 1. EXTRACT CERTIFICATIONS TEXT ===
    cert_text = ""
    rows = soup.find_all('tr')
    for row in rows:
        cells = row.find_all('td')
        if len(cells) > 0 and "Certifications" in cells[0].get_text():
            if len(cells) > 1:
                cert_text = cells[1].get_text(strip=True).lower()
            break

    # === 2. HEADER: BADGES (Left) + LOGO (Right) ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.0)
    header_table.columns[1].width = Inches(3.5)
    
    # -- A. Add Badges to Left Cell --
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    badges_added = False
    for keyword, image_file in CERT_BADGE_MAP.items():
        if keyword in cert_text:
            full_badge_path = os.path.join(BADGE_DIR, image_file)
            
            if os.path.exists(full_badge_path):
                run = left_para.add_run()
                run.add_picture(full_badge_path, width=Inches(0.85))
                run.add_text("  ") # Spacer
                badges_added = True
            else:
                print(f"⚠️ Warning: Badge image not found at {full_badge_path}")
    
    # -- B. Add Company Logo to Right Cell --
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    logo_path = "/home/ca/Projects/resume_converter/download.png"
    if os.path.exists(logo_path):
        run = right_para.add_run()
        run.add_picture(logo_path, width=Inches(2.5))
    elif os.path.exists("download.png"):
        run = right_para.add_run()
        run.add_picture("download.png", width=Inches(2.5))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # === 3. MAIN INFO TABLE ===
    main_table = soup.find('table', {'class': 'resume-table'})
    if main_table:
        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'
        t.columns[0].width = Inches(1.875)
        t.columns[1].width = Inches(5.625)
        set_table_borders(t)
        
        for row in main_table.find_all('tr'):
            cells = row.find_all('td')
            if len(cells) == 2:
                doc_row = t.add_row().cells
                left_text = cells[0].get_text(strip=True)
                
                # LEFT CELL
                set_cell_content(doc_row[0], left_text, is_header=True, center_align=False)
                
                # RIGHT CELL Logic
                is_name_cell = "name-value-cell" in cells[1].get('class', []) or left_text == "Name"
                
                if is_name_cell:
                    # Name: Blue BG, Centered
                    set_cell_content(doc_row[1], "", is_header=True, center_align=True) 
                    process_html_content(doc_row[1], cells[1], text_color="white")
                else:
                    # Standard: White BG, Centered
                    set_cell_content(doc_row[1], "", is_header=False, center_align=True)
                    process_html_content(doc_row[1], cells[1])

    # === 4. SEPARATE PROJECT HEADER ===
    proj_div = soup.find('div', {'class': 'project-title-bar-text-only'})
    if not proj_div: proj_div = soup.find('div', {'class': 'project-title-bar'})
    
    if proj_div:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run(proj_div.get_text(strip=True))
        run.font.name = 'Arial'
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_run_background(run, THEME_COLOR)

    # === 5. PROJECTS TABLE ===
    proj_table = soup.find('table', {'class': 'projects-table'})
    if proj_table:
        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'
        t.columns[0].width = Inches(1.875)
        t.columns[1].width = Inches(5.625)
        set_table_borders(t)
        
        for row in proj_table.find_all('tr'):
            cells = row.find_all('td')
            if len(cells) == 2:
                doc_row = t.add_row().cells
                
                # Left Column: Project Name
                set_cell_content(doc_row[0], cells[0].get_text(strip=True), is_header=False, center_align=True, italic=True, bold=True)
                
                # Right Column: Details
                set_cell_content(doc_row[1], "", is_header=False)
                process_project_content(doc_row[1], cells[1])
    
    doc.save(docx_file_path)

# === HELPER FUNCTIONS ===

def set_run_background(run, color_hex):
    """Applies shading (background color) to a specific text run."""
    rPr = run._r.get_or_add_rPr()
    fill = (color_hex or "").replace("#", "")
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    rPr.append(shd)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_content(cell, text, is_header=False, center_align=False, font_size=12, italic=False, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    
    # Increase Padding for height (Deepak Singh style)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center_align else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Note: Setting vertical_alignment adds a w:vAlign element to tcPr.
    # Our background color must come BEFORE w:vAlign in the XML for Teams compatibility.
    
    if is_header:
        set_cell_background(cell, THEME_COLOR)
        run = paragraph.add_run(text)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(font_size)
    else:
        if text:
            run = paragraph.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(font_size)
            run.font.italic = italic
            run.font.bold = bold

def set_cell_background(cell, color_hex):
    """
    Sets cell background color.
    CRITICAL FIX for Teams/Word Online: Ensure <w:shd> is inserted BEFORE <w:vAlign>.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    fill = (color_hex or "").replace("#", "")
    
    # Define shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    
    # Remove any existing shading to prevent duplication
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
        
    # FIX: Insert <w:shd> before <w:vAlign> or other property tags that must come after it.
    # Common tags that follow 'shd' in OOXML schema: noWrap, tcMar, textDirection, tcFitText, vAlign, hideMark
    
    inserted = False
    for tag in ['w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark']:
        element = tcPr.find(qn(tag))
        if element is not None:
            element.addprevious(shd)
            inserted = True
            break
            
    if not inserted:
        # If none of the later tags exist, safe to append
        tcPr.append(shd)

def process_html_content(cell, html_element, text_color=None):
    text = html_element.get_text(strip=True)
    if text:
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        if text_color == "white":
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

def process_project_content(cell, html_element):
    cell.text = "" 
    first_element = True
    
    for element in html_element.children:
        if isinstance(element, Tag):
            if element.name == 'div':
                if first_element:
                    para = cell.paragraphs[0]
                    first_element = False
                else:
                    para = cell.add_paragraph()
                
                para.paragraph_format.space_after = Pt(2)
                strong = element.find('strong')
                if strong:
                    label = strong.get_text(strip=True)
                    run = para.add_run(label + " ")
                    run.font.name = 'Arial'
                    run.bold = True
                    
                    val = element.get_text().replace(label, '', 1).strip()
                    if val: 
                        r2 = para.add_run(val)
                        r2.font.name = 'Arial'
                        r2.font.size = Pt(12)
                else:
                    text_val = element.get_text(strip=True)
                    run = para.add_run(text_val)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

                    # Make heading bold + spacing
                    if "Project Description" in text_val:
                        run.bold = True
                        para.paragraph_format.space_before = Pt(8)
                        para.paragraph_format.space_after = Pt(2)

            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li'), 1):
                    if first_element:
                        para = cell.paragraphs[0]
                        first_element = False
                    else:
                        para = cell.add_paragraph()
                        
                    # Handle Indentation manually for best compatibility
                    para.paragraph_format.left_indent = Inches(0.35)
                    para.paragraph_format.first_line_indent = Inches(0)
                    try:
                        para.paragraph_format.tab_stops.clear_all()
                        para.paragraph_format.tab_stops.add_tab_stop(Inches(0.55))
                    except Exception:
                        pass

                    run = para.add_run(f"{i}.\t" + li.get_text(strip=True))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            elif element.name == 'ul':
                for li in element.find_all('li'):
                    if first_element:
                        para = cell.paragraphs[0]
                        first_element = False
                    else:
                        para = cell.add_paragraph()

                    para.paragraph_format.left_indent = Inches(0.35)
                    para.paragraph_format.first_line_indent = Inches(0)
                    try:
                        para.paragraph_format.tab_stops.clear_all()
                        para.paragraph_format.tab_stops.add_tab_stop(Inches(0.55))
                    except Exception:
                        pass

                    run = para.add_run("•\t" + li.get_text(strip=True))
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            elif element.name == 'p':
                if first_element:
                    para = cell.paragraphs[0]
                    first_element = False
                else:
                    para = cell.add_paragraph()
                    
                para.paragraph_format.space_before = Pt(3)
                run = para.add_run(element.get_text(strip=True))
                run.font.name = 'Arial'
                run.font.size = Pt(12)

def convert_salesforce_resume(html_file_path, docx_file_path):
    try:
        html_to_docx(html_file_path, docx_file_path)
    except Exception as e:
        print(f"Conversion failed: {e}")